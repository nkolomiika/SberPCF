from __future__ import annotations

import asyncio
import json
import re
import secrets
from datetime import UTC, date, datetime, timedelta
from io import BytesIO
from pathlib import Path
from urllib.parse import parse_qsl, urlparse, urlsplit
from uuid import UUID, uuid4
from xml.sax.saxutils import escape

import magic
import yaml
from cvss import CVSS4
from docx import Document
from docx.shared import Inches
from fastapi import UploadFile
from PIL import Image as PillowImage, ImageOps, UnidentifiedImageError
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as PlatypusImage, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from sqlalchemy import Select, and_, delete, func, or_, select, update
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy.orm.attributes import set_committed_value

from app.audit_store import audit_store
from app.config import get_settings
from app.enums import AssetType, NotificationType, ProjectStatus, Severity, UserRole
from app.exceptions import ConflictError, ForbiddenError, NotFoundError, UnauthorizedError, ValidationError
from app.models import (
    AuditLog,
    Comment,
    CommentMention,
    Endpoint,
    File,
    Host,
    MailJob,
    Notification,
    Port,
    Project,
    ProjectFolder,
    ProjectMember,
    RefreshToken,
    Service,
    User,
    Vulnerability,
    VulnerabilityAsset,
)
from app.mail import build_temporary_password_email
from app.messaging import publish_mail_job
from app.schemas import (
    CommentOut,
    ImportResult,
    MentionOut,
    NotificationContext,
    NotificationOut,
    OpenApiImportResult,
    PcfImportPayload,
)
from app.security import (
    create_access_token,
    create_refresh_token,
    decode_token,
    hash_password,
    hash_refresh_token,
    verify_password,
)
from app.storage.minio_client import MinioStorage
from app.ws_manager import ws_manager

settings = get_settings()
DEBUG_LOG_PATH = Path("/workspace/debug-755228.log")


def _debug_log(hypothesis_id: str, location: str, message: str, data: dict) -> None:
    payload = {
        "sessionId": "755228",
        "runId": "workflow-title-debug",
        "hypothesisId": hypothesis_id,
        "location": location,
        "message": message,
        "data": data,
        "timestamp": int(datetime.now(UTC).timestamp() * 1000),
    }
    try:
        with DEBUG_LOG_PATH.open("a", encoding="utf-8") as handle:
            handle.write(json.dumps(payload, ensure_ascii=False) + "\n")
    except Exception:
        pass
MENTION_RE = re.compile(r"@([a-zA-Z0-9_.-]{1,100})")
MAX_FILE_SIZE = 50 * 1024 * 1024
MAX_OPENAPI_IMPORT_BYTES = 2 * 1024 * 1024
MAX_OPENAPI_PATHS = 2000
ALLOWED_MIME_TYPES = {
    "image/png",
    "image/jpeg",
    "image/gif",
    "image/webp",
    "text/plain",
    "application/pdf",
    "application/xml",
    "application/json",
    "application/zip",
    "application/x-tar",
    "application/gzip",
}


class AuditService:
    """Сервис записи журнала действий."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def log(
        self,
        action: str,
        *,
        user_id: UUID | None = None,
        entity_type: str | None = None,
        entity_id: UUID | None = None,
        details: dict | None = None,
        ip_address: str | None = None,
    ) -> None:
        """Создаёт запись аудита."""
        username: str | None = None
        if user_id:
            username = await self.db.scalar(select(User.username).where(User.id == user_id))
        created_at = datetime.now(UTC)
        self.db.add(
            AuditLog(
                user_id=user_id,
                action=action,
                entity_type=entity_type,
                entity_id=entity_id,
                details=details,
                ip_address=ip_address,
                created_at=created_at,
            )
        )
        await self.db.commit()
        await audit_store.insert(
            user_id=user_id,
            username=username,
            action=action,
            entity_type=entity_type,
            entity_id=entity_id,
            details=details,
            ip_address=ip_address,
            created_at=created_at,
        )


class AuthService:
    """Сервис аутентификации и управления JWT-cookie."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.audit = AuditService(db)

    async def login(self, username: str, password: str, ip_address: str | None = None) -> tuple[str, str, User]:
        """Проверяет учётные данные и создаёт пару токенов."""
        user = await self.db.scalar(select(User).where(User.username == username))
        if not user or not verify_password(password, user.password_hash):
            raise UnauthorizedError("Неверный логин или пароль")
        if not user.is_active:
            raise UnauthorizedError("Пользователь деактивирован")

        access_token = create_access_token(user.id)
        refresh_token = create_refresh_token(user.id)
        refresh_hash = hash_refresh_token(refresh_token)
        expires_at = datetime.now(UTC) + timedelta(days=settings.jwt_refresh_token_expire_days)
        self.db.add(RefreshToken(user_id=user.id, token_hash=refresh_hash, expires_at=expires_at))
        await self.db.commit()

        await self.audit.log("LOGIN", user_id=user.id, ip_address=ip_address)
        return access_token, refresh_token, user

    async def refresh(self, refresh_token: str | None, ip_address: str | None = None) -> tuple[str, bool]:
        """Обновляет access-токен по валидному refresh-cookie."""
        if not refresh_token:
            raise UnauthorizedError("Refresh token отсутствует")
        user_id = decode_token(refresh_token, expected_type="refresh")
        token_hash = hash_refresh_token(refresh_token)
        token_entry = await self.db.scalar(
            select(RefreshToken).where(
                and_(
                    RefreshToken.token_hash == token_hash,
                    RefreshToken.user_id == user_id,
                    RefreshToken.revoked_at.is_(None),
                    RefreshToken.expires_at > datetime.now(UTC),
                )
            )
        )
        if not token_entry:
            raise UnauthorizedError("Refresh token недействителен")
        user = await self.db.scalar(select(User).where(User.id == user_id))
        if not user or not user.is_active:
            raise UnauthorizedError("Пользователь деактивирован")
        return create_access_token(user_id), user.must_change_password

    async def logout(self, user_id: UUID, ip_address: str | None = None) -> None:
        """Отзывает все активные refresh-токены пользователя."""
        await self.db.execute(
            update(RefreshToken)
            .where(and_(RefreshToken.user_id == user_id, RefreshToken.revoked_at.is_(None)))
            .values(revoked_at=datetime.now(UTC))
        )
        await self.db.commit()
        await self.audit.log("LOGOUT", user_id=user_id, ip_address=ip_address)


class UserService:
    """Сервис управления пользователями."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.audit = AuditService(db)
        self.storage = MinioStorage()
        self.storage.ensure_bucket()

    @staticmethod
    def _normalized_admin_email() -> str:
        configured_email = settings.initial_admin_email
        if "@" in configured_email and "." in configured_email.split("@", 1)[1]:
            return configured_email
        return "admin@example.com"

    @staticmethod
    def _normalize_tags(tags: list[str] | None) -> list[str]:
        normalized: list[str] = []
        seen: set[str] = set()
        for tag in tags or []:
            cleaned = tag.strip()
            if not cleaned:
                continue
            lowered = cleaned.lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            normalized.append(cleaned)
        return normalized

    @staticmethod
    def _generate_temporary_password() -> str:
        return secrets.token_urlsafe(12)

    @staticmethod
    def _ensure_mail_delivery_enabled() -> None:
        if not settings.mail_enabled:
            raise ValidationError("Отправка email отключена. Включите SMTP/mail worker или задайте пароль вручную.")

    @staticmethod
    def ensure_can_view_avatar(requester: User, target_user_id: UUID) -> None:
        if requester.role != UserRole.ADMIN and requester.id != target_user_id:
            raise ForbiddenError("Недостаточно прав для просмотра чужого аватара")

    async def _ensure_unique_identity(
        self,
        *,
        username: str | None = None,
        email: str | None = None,
        exclude_user_id: UUID | None = None,
    ) -> None:
        clauses = []
        if username:
            clauses.append(User.username == username)
        if email:
            clauses.append(User.email == email)
        if not clauses:
            return
        query = select(User).where(or_(*clauses))
        if exclude_user_id:
            query = query.where(User.id != exclude_user_id)
        exists = await self.db.scalar(query)
        if exists:
            raise ConflictError("username или email уже заняты")

    async def _enqueue_temporary_password_mail(
        self,
        *,
        user: User,
        temporary_password: str,
        actor_id: UUID | None,
    ) -> MailJob:
        subject, body = build_temporary_password_email(username=user.full_name or user.username, temporary_password=temporary_password)
        job = MailJob(
            user_id=user.id,
            created_by=actor_id,
            recipient_email=user.email,
            subject=subject,
            template="temporary_password",
            payload={"username": user.username, "temporary_password": temporary_password, "body": body},
            status="pending",
        )
        self.db.add(job)
        await self.db.flush()
        return job

    async def _publish_mail_job(self, job: MailJob) -> None:
        try:
            await publish_mail_job(job.id)
        except Exception as exc:
            job.status = "pending"
            job.last_error = str(exc)
        else:
            job.status = "queued"
            job.published_at = datetime.now(UTC)
        await self.db.commit()

    async def get_user_profile(self, user_id: UUID) -> User:
        return await self.get_user(user_id)

    async def bootstrap_admin(self) -> None:
        """Создаёт стартового администратора при пустой таблице users."""
        normalized_email = self._normalized_admin_email()
        existing_admin = await self.db.scalar(
            select(User).where(or_(User.username == settings.initial_admin_username, User.email == normalized_email))
        )
        if existing_admin:
            if existing_admin.username == settings.initial_admin_username and existing_admin.email != normalized_email:
                existing_admin.email = normalized_email
                await self.db.commit()
            return
        total_users = await self.db.scalar(select(func.count()).select_from(User))
        if total_users and total_users > 0:
            admin_user = await self.db.scalar(select(User).where(User.username == settings.initial_admin_username))
            if admin_user and ("@" not in admin_user.email or "." not in admin_user.email.split("@", 1)[1]):
                admin_user.email = normalized_email
                await self.db.commit()
            return
        admin = User(
            username=settings.initial_admin_username,
            email=normalized_email,
            full_name="Administrator",
            tags=["admin"],
            password_hash=hash_password(settings.initial_admin_password),
            password_changed_at=datetime.now(UTC),
            role=UserRole.ADMIN,
            is_active=True,
        )
        self.db.add(admin)
        await self.db.commit()

    async def list_users(self, page: int, size: int) -> tuple[list[User], int]:
        """Возвращает список пользователей с пагинацией."""
        total = await self.db.scalar(select(func.count()).select_from(User)) or 0
        items = (
            await self.db.scalars(select(User).order_by(User.created_at.desc()).offset((page - 1) * size).limit(size))
        ).all()
        return list(items), total

    async def get_user(self, user_id: UUID) -> User:
        """Возвращает пользователя по идентификатору."""
        user = await self.db.scalar(select(User).where(User.id == user_id))
        if not user:
            raise NotFoundError("Пользователь не найден")
        return user

    async def create_user(self, payload: dict, actor_id: UUID, ip_address: str | None = None) -> User:
        """Создаёт нового пользователя."""
        await self._ensure_unique_identity(username=payload["username"], email=payload["email"])
        send_invite_email = bool(payload.get("send_invite_email"))
        temporary_password = payload.get("password")
        if send_invite_email:
            self._ensure_mail_delivery_enabled()
        if not temporary_password and not send_invite_email:
            raise ValidationError("Нужно указать пароль или включить отправку приглашения на email")
        generated_password = False
        if not temporary_password:
            temporary_password = self._generate_temporary_password()
            generated_password = True
        must_change_password = send_invite_email or generated_password
        user = User(
            username=payload["username"],
            email=payload["email"],
            full_name=payload.get("full_name"),
            tags=self._normalize_tags(payload.get("tags")),
            password_hash=hash_password(temporary_password),
            must_change_password=must_change_password,
            password_changed_at=None if must_change_password else datetime.now(UTC),
            role=payload.get("role", UserRole.PENTESTER),
            is_active=True,
        )
        self.db.add(user)
        mail_job: MailJob | None = None
        if send_invite_email:
            await self.db.flush()
            mail_job = await self._enqueue_temporary_password_mail(user=user, temporary_password=temporary_password, actor_id=actor_id)
        await self.db.commit()
        await self.db.refresh(user)
        await self.audit.log(
            "CREATE",
            user_id=actor_id,
            entity_type="user",
            entity_id=user.id,
            details={"email": user.email, "invite_sent": send_invite_email, "must_change_password": must_change_password},
            ip_address=ip_address,
        )
        if mail_job:
            await self._publish_mail_job(mail_job)
        return user

    async def update_user(self, user_id: UUID, payload: dict, actor_id: UUID, ip_address: str | None = None) -> User:
        """Обновляет данные пользователя."""
        user = await self.get_user(user_id)
        submitted_email = payload.pop("email", None)
        if submitted_email is not None and submitted_email != user.email:
            raise ValidationError("Администратор не может менять email пользователя. Пользователь должен изменить его сам.")
        username = payload.get("username", user.username)
        await self._ensure_unique_identity(username=username, email=user.email, exclude_user_id=user.id)
        if "tags" in payload:
            payload["tags"] = self._normalize_tags(payload.get("tags"))
        for key, value in payload.items():
            if value is not None:
                setattr(user, key, value)
        await self.db.commit()
        await self.db.refresh(user)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="user", entity_id=user.id, ip_address=ip_address)
        return user

    async def update_own_profile(self, user_id: UUID, payload: dict, ip_address: str | None = None) -> User:
        user = await self.get_user(user_id)
        username = payload.get("username", user.username)
        email = payload.get("email", user.email)
        await self._ensure_unique_identity(username=username, email=email, exclude_user_id=user.id)
        if "tags" in payload:
            payload["tags"] = self._normalize_tags(payload.get("tags"))
        for key, value in payload.items():
            if value is not None:
                setattr(user, key, value)
        await self.db.commit()
        await self.db.refresh(user)
        await self.audit.log("UPDATE", user_id=user_id, entity_type="user_profile", entity_id=user.id, ip_address=ip_address)
        return user

    async def delete_user(self, user_id: UUID, actor_id: UUID, ip_address: str | None = None) -> None:
        """Удаляет пользователя."""
        if user_id == actor_id:
            raise ValidationError("Нельзя удалить самого себя")
        user = await self.get_user(user_id)
        await self.db.delete(user)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="user", entity_id=user_id, ip_address=ip_address)

    async def reset_password(self, user_id: UUID, actor_id: UUID, ip_address: str | None = None) -> User:
        """Сбрасывает пароль пользователя, выставляет временный пароль и отправляет его по почте."""
        self._ensure_mail_delivery_enabled()
        user = await self.get_user(user_id)
        temporary_password = self._generate_temporary_password()
        user.password_hash = hash_password(temporary_password)
        user.must_change_password = True
        user.password_changed_at = None
        await self.db.execute(
            update(RefreshToken)
            .where(and_(RefreshToken.user_id == user.id, RefreshToken.revoked_at.is_(None)))
            .values(revoked_at=datetime.now(UTC))
        )
        mail_job = await self._enqueue_temporary_password_mail(user=user, temporary_password=temporary_password, actor_id=actor_id)
        await self.db.commit()
        await self.audit.log(
            "UPDATE",
            user_id=actor_id,
            entity_type="user_password_reset",
            entity_id=user_id,
            details={"email": user.email, "must_change_password": True},
            ip_address=ip_address,
        )
        await self._publish_mail_job(mail_job)
        return user

    async def change_own_password(
        self,
        user_id: UUID,
        *,
        current_password: str,
        new_password: str,
        ip_address: str | None = None,
    ) -> User:
        user = await self.get_user(user_id)
        if not verify_password(current_password, user.password_hash):
            raise UnauthorizedError("Текущий пароль указан неверно")
        user.password_hash = hash_password(new_password)
        user.must_change_password = False
        user.password_changed_at = datetime.now(UTC)
        await self.db.execute(
            update(RefreshToken)
            .where(and_(RefreshToken.user_id == user.id, RefreshToken.revoked_at.is_(None)))
            .values(revoked_at=datetime.now(UTC))
        )
        await self.db.commit()
        await self.audit.log("UPDATE", user_id=user.id, entity_type="user_password", entity_id=user.id, ip_address=ip_address)
        await self.db.refresh(user)
        return user

    async def force_change_password(self, user_id: UUID, *, new_password: str, ip_address: str | None = None) -> User:
        user = await self.get_user(user_id)
        if not user.must_change_password:
            raise ForbiddenError("Принудительная смена пароля не требуется")
        user.password_hash = hash_password(new_password)
        user.must_change_password = False
        user.password_changed_at = datetime.now(UTC)
        await self.db.execute(
            update(RefreshToken)
            .where(and_(RefreshToken.user_id == user.id, RefreshToken.revoked_at.is_(None)))
            .values(revoked_at=datetime.now(UTC))
        )
        await self.db.commit()
        await self.audit.log(
            "UPDATE",
            user_id=user.id,
            entity_type="user_force_password_change",
            entity_id=user.id,
            ip_address=ip_address,
        )
        await self.db.refresh(user)
        return user

    async def upload_avatar(self, user_id: UUID, upload: UploadFile, ip_address: str | None = None) -> User:
        user = await self.get_user(user_id)
        content = await upload.read()
        if len(content) > 5 * 1024 * 1024:
            raise ValidationError("Размер аватара превышает 5 МБ")
        mime_type = magic.from_buffer(content, mime=True)
        if mime_type not in {"image/png", "image/jpeg", "image/webp", "image/gif"}:
            raise ValidationError("Аватар должен быть изображением PNG, JPEG, WEBP или GIF")
        if user.avatar_minio_key:
            await asyncio.to_thread(self.storage.delete, user.avatar_minio_key)
        object_key = await asyncio.to_thread(self.storage.upload_bytes, content, mime_type, upload.filename or "avatar.bin")
        user.avatar_minio_bucket = settings.minio_bucket_name
        user.avatar_minio_key = object_key
        user.avatar_content_type = mime_type
        user.avatar_uploaded_at = datetime.now(UTC)
        await self.db.commit()
        await self.db.refresh(user)
        await self.audit.log("FILE_UPLOAD", user_id=user.id, entity_type="user_avatar", entity_id=user.id, ip_address=ip_address)
        return user

    async def download_avatar(self, user_id: UUID) -> tuple[User, bytes]:
        user = await self.get_user(user_id)
        if not user.avatar_minio_key:
            raise NotFoundError("Аватар не найден")
        blob = await asyncio.to_thread(self.storage.download_bytes, user.avatar_minio_key)
        return user, blob


class ProjectService:
    """Сервис управления проектами и участниками."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.audit = AuditService(db)

    @staticmethod
    def _validate_project_dates(start_date: date | None, end_date: date | None) -> None:
        if start_date and end_date and end_date < start_date:
            raise ValidationError("Дата окончания проекта не может быть раньше даты начала")

    @staticmethod
    def _normalize_project_folder(folder: str | None) -> str:
        normalized = (folder or "").strip()
        return normalized

    @staticmethod
    def _normalize_folder_segment(name: str) -> str:
        normalized = name.strip().strip("/")
        if not normalized:
            raise ValidationError("Название папки не может быть пустым")
        if "/" in normalized:
            raise ValidationError("Название папки не должно содержать '/'")
        return normalized

    async def list_projects(self, current_user: User, page: int, size: int, status: str | None = None) -> tuple[list[Project], int]:
        """Возвращает список проектов с учётом доступа."""
        query: Select = select(Project)
        if status:
            query = query.where(Project.status == status)
        if current_user.role != UserRole.ADMIN:
            query = query.join(ProjectMember, ProjectMember.project_id == Project.id).where(ProjectMember.user_id == current_user.id)
        count_query = select(func.count()).select_from(query.subquery())
        total = await self.db.scalar(count_query) or 0
        items = (await self.db.scalars(query.order_by(Project.created_at.desc()).offset((page - 1) * size).limit(size))).all()
        return list(items), total

    async def get_project(self, project_id: UUID) -> Project:
        """Возвращает проект по ID."""
        project = await self.db.scalar(select(Project).where(Project.id == project_id))
        if not project:
            raise NotFoundError("Проект не найден")
        return project

    async def create_project(self, payload: dict, actor_id: UUID, ip_address: str | None = None) -> Project:
        """Создаёт новый проект."""
        self._validate_project_dates(payload.get("start_date"), payload.get("end_date"))
        payload["folder"] = self._normalize_project_folder(payload.get("folder"))
        project = Project(**payload, created_by=actor_id)
        self.db.add(project)
        await self.db.commit()
        await self.db.refresh(project)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="project", entity_id=project.id, ip_address=ip_address)
        await ws_manager.broadcast_projects_index(
            {"event": "created", "entity": "project", "project_id": str(project.id), "data": {"id": str(project.id)}}
        )
        return project

    async def update_project(self, project_id: UUID, payload: dict, actor_id: UUID, ip_address: str | None = None) -> Project:
        """Обновляет проект."""
        project = await self.get_project(project_id)
        if "folder" in payload:
            payload["folder"] = self._normalize_project_folder(payload.get("folder"))
        next_start_date = payload.get("start_date", project.start_date)
        next_end_date = payload.get("end_date", project.end_date)
        self._validate_project_dates(next_start_date, next_end_date)
        old_status = project.status
        next_status = payload.get("status")
        if next_status is not None:
            if next_status == ProjectStatus.ACTIVE:
                payload["timeline_frozen_at"] = None
            elif old_status == ProjectStatus.ACTIVE:
                payload["timeline_frozen_at"] = datetime.now(UTC)
            elif "timeline_frozen_at" not in payload:
                payload["timeline_frozen_at"] = project.timeline_frozen_at
        for key, value in payload.items():
            if value is not None:
                setattr(project, key, value)
            elif key == "timeline_frozen_at":
                setattr(project, key, None)
        await self.db.commit()
        await self.db.refresh(project)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="project", entity_id=project.id, ip_address=ip_address)
        await ws_manager.broadcast(project.id, {"event": "updated", "entity": "project", "project_id": str(project.id), "data": {"id": str(project.id)}})
        await ws_manager.broadcast_projects_index(
            {"event": "updated", "entity": "project", "project_id": str(project.id), "data": {"id": str(project.id)}}
        )
        if payload.get("status") and payload["status"] != old_status:
            await self.audit.log(
                "STATUS_CHANGE",
                user_id=actor_id,
                entity_type="project",
                entity_id=project.id,
                details={"old_status": old_status.value, "new_status": project.status.value},
                ip_address=ip_address,
            )
        return project

    async def delete_project(self, project_id: UUID, actor_id: UUID, ip_address: str | None = None) -> None:
        """Удаляет проект и связанные сущности."""
        project = await self.get_project(project_id)
        await self.db.delete(project)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="project", entity_id=project_id, ip_address=ip_address)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "project", "project_id": str(project_id), "data": {"id": str(project_id)}})
        await ws_manager.broadcast_projects_index(
            {"event": "deleted", "entity": "project", "project_id": str(project_id), "data": {"id": str(project_id)}}
        )

    async def list_members(self, project_id: UUID) -> list[dict]:
        """Возвращает список участников проекта."""
        rows = (
            await self.db.execute(
                select(ProjectMember, User)
                .join(User, User.id == ProjectMember.user_id)
                .where(ProjectMember.project_id == project_id)
                .order_by(ProjectMember.added_at.desc())
            )
        ).all()
        return [
            {
                "user_id": user.id,
                "username": user.username,
                "email": user.email,
                "role": user.role,
                "added_at": member.added_at,
            }
            for member, user in rows
        ]

    async def add_member(self, project_id: UUID, user_id: UUID, actor_id: UUID, ip_address: str | None = None) -> dict:
        """Добавляет пользователя в участники проекта."""
        await self.get_project(project_id)
        user = await self.db.scalar(select(User).where(User.id == user_id))
        if not user:
            raise NotFoundError("Пользователь не найден")
        exists = await self.db.scalar(
            select(ProjectMember).where(and_(ProjectMember.project_id == project_id, ProjectMember.user_id == user_id))
        )
        if exists:
            raise ConflictError("Пользователь уже участник проекта")
        member = ProjectMember(project_id=project_id, user_id=user_id)
        self.db.add(member)
        await self.db.commit()
        await self.db.refresh(member)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="project_member", entity_id=member.id, ip_address=ip_address)
        return {"user_id": user.id, "username": user.username, "added_at": member.added_at}

    async def remove_member(self, project_id: UUID, user_id: UUID, actor_id: UUID, ip_address: str | None = None) -> None:
        """Удаляет пользователя из участников проекта."""
        member = await self.db.scalar(
            select(ProjectMember).where(and_(ProjectMember.project_id == project_id, ProjectMember.user_id == user_id))
        )
        if not member:
            raise NotFoundError("Участник не найден")
        await self.db.delete(member)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="project_member", entity_id=member.id, ip_address=ip_address)

    async def list_folders(self) -> list[ProjectFolder]:
        """Возвращает список папок проектов."""
        rows = await self.db.scalars(select(ProjectFolder).order_by(ProjectFolder.path.asc()))
        return list(rows.all())

    async def create_folder(self, name: str, parent_id: UUID | None, actor_id: UUID, ip_address: str | None = None) -> ProjectFolder:
        """Создаёт новую папку проекта (в т.ч. вложенную)."""
        segment = self._normalize_folder_segment(name)
        parent: ProjectFolder | None = None
        if parent_id:
            parent = await self.db.scalar(select(ProjectFolder).where(ProjectFolder.id == parent_id))
            if not parent:
                raise NotFoundError("Родительская папка не найдена")
        path = f"{parent.path}/{segment}" if parent else segment
        duplicate = await self.db.scalar(select(ProjectFolder).where(ProjectFolder.path == path))
        if duplicate:
            raise ConflictError("Папка с таким путём уже существует")
        folder = ProjectFolder(name=segment, path=path, parent_id=parent_id, created_by=actor_id)
        self.db.add(folder)
        await self.db.commit()
        await self.db.refresh(folder)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="project_folder", entity_id=folder.id, ip_address=ip_address)
        await ws_manager.broadcast_projects_index(
            {"event": "created", "entity": "project_folder", "data": {"id": str(folder.id), "path": folder.path}}
        )
        return folder

    async def move_folder(
        self, folder_id: UUID, new_parent_id: UUID | None, actor_id: UUID, ip_address: str | None = None
    ) -> ProjectFolder:
        """Перемещает папку и пересчитывает пути дочерних папок и проектов."""
        folder = await self.db.scalar(select(ProjectFolder).where(ProjectFolder.id == folder_id))
        if not folder:
            raise NotFoundError("Папка не найдена")
        if new_parent_id == folder.id:
            raise ValidationError("Нельзя переместить папку в саму себя")

        new_parent: ProjectFolder | None = None
        if new_parent_id:
            new_parent = await self.db.scalar(select(ProjectFolder).where(ProjectFolder.id == new_parent_id))
            if not new_parent:
                raise NotFoundError("Родительская папка не найдена")
            if new_parent.path == folder.path or new_parent.path.startswith(f"{folder.path}/"):
                raise ValidationError("Нельзя переместить папку в её дочернюю папку")

        duplicate = await self.db.scalar(
            select(ProjectFolder).where(
                and_(
                    ProjectFolder.id != folder.id,
                    ProjectFolder.parent_id == new_parent_id,
                    ProjectFolder.name == folder.name,
                )
            )
        )
        if duplicate:
            raise ConflictError("В целевой папке уже есть папка с таким именем")

        old_path = folder.path
        new_path = f"{new_parent.path}/{folder.name}" if new_parent else folder.name
        if new_path == old_path and folder.parent_id == new_parent_id:
            return folder

        subtree_folders = (
            await self.db.scalars(
                select(ProjectFolder).where(or_(ProjectFolder.path == old_path, ProjectFolder.path.like(f"{old_path}/%")))
            )
        ).all()
        subtree_projects = (
            await self.db.scalars(select(Project).where(or_(Project.folder == old_path, Project.folder.like(f"{old_path}/%"))))
        ).all()

        for subtree_folder in subtree_folders:
            suffix = subtree_folder.path[len(old_path) :]
            subtree_folder.path = f"{new_path}{suffix}"
            if subtree_folder.id == folder.id:
                subtree_folder.parent_id = new_parent_id

        for project in subtree_projects:
            suffix = project.folder[len(old_path) :]
            project.folder = f"{new_path}{suffix}"

        await self.db.commit()
        await self.db.refresh(folder)
        await self.audit.log(
            "UPDATE",
            user_id=actor_id,
            entity_type="project_folder",
            entity_id=folder.id,
            details={"from": old_path, "to": new_path},
            ip_address=ip_address,
        )
        await ws_manager.broadcast_projects_index(
            {
                "event": "updated",
                "entity": "project_folder",
                "data": {"id": str(folder.id), "from": old_path, "to": new_path},
            }
        )
        return folder


class AssetService:
    """Сервис CRUD для host/port/service/endpoint."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.audit = AuditService(db)

    @staticmethod
    def _normalize_endpoint_query_params(raw_params: list[dict] | None) -> list[dict]:
        normalized: list[dict] = []
        for item in raw_params or []:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or "").strip()
            if not name:
                continue
            value = item.get("value")
            description = item.get("description")
            normalized.append(
                {
                    "name": name,
                    "value": None if value is None else str(value),
                    "required": bool(item.get("required", False)),
                    "description": None if description in (None, "") else str(description).strip(),
                }
            )
        return normalized

    _HTTP_HEADER_DROP = frozenset(
        {
            "referer",
            "referrer",
            "connection",
            "accept-encoding",
            "accept-language",
            "sec-fetch-dest",
            "sec-fetch-mode",
            "sec-fetch-site",
            "sec-fetch-user",
            "sec-ch-ua",
            "sec-ch-ua-mobile",
            "sec-ch-ua-platform",
            "user-agent",
            "host",
        }
    )

    _METHODS_WITHOUT_BODY = frozenset({"GET", "HEAD", "OPTIONS", "DELETE"})
    _UUID_PATH_SEGMENT_RE = re.compile(
        r"^[0-9a-fA-F]{8}-[0-9a-fA-F]{4}-[1-5][0-9a-fA-F]{3}-[89abAB][0-9a-fA-F]{3}-[0-9a-fA-F]{12}$"
    )

    @staticmethod
    def _normalize_endpoint_path(path_value: str | None) -> str | None:
        if path_value is None:
            return None
        path_only, _ = AssetService._parse_request_target_token(str(path_value))
        normalized = re.sub(r"/{2,}", "/", (path_only or "/").strip() or "/")
        normalized = normalized if normalized.startswith("/") else f"/{normalized}"
        if normalized != "/" and normalized.endswith("/"):
            normalized = normalized[:-1]
        segments = [
            "{UUID}" if AssetService._UUID_PATH_SEGMENT_RE.fullmatch(segment) else segment
            for segment in normalized.split("/")
            if segment
        ]
        return f"/{'/'.join(segments)}" if segments else "/"

    @staticmethod
    def _normalize_endpoint_headers(raw_headers: list[dict] | None) -> list[dict]:
        normalized: list[dict] = []
        for item in raw_headers or []:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or "").strip()
            if not name:
                continue
            value = item.get("value")
            normalized.append({"name": name, "value": "" if value is None else str(value)})
        return normalized

    @staticmethod
    def _sanitize_parsed_header_pairs(pairs: list[tuple[str, str]]) -> list[dict]:
        """Strip noisy/sensitive headers; Cookie/Authorization placeholders."""
        out: list[dict] = []
        had_cookie = False
        for name, value in pairs:
            ln = name.strip().lower()
            if ln in AssetService._HTTP_HEADER_DROP:
                continue
            if ln == "cookie":
                had_cookie = True
                continue
            if ln == "authorization":
                out.append({"name": "Authorization", "value": "{YOUR_CREDENTIALS_HERE}"})
                continue
            if ln == "content-type":
                continue
            out.append({"name": name.strip(), "value": value.strip()})
        if had_cookie:
            out.append({"name": "Cookie", "value": "{YOUR_TOKENS_HERE}"})
        return out

    @staticmethod
    def _sanitize_stored_header_items(raw: list[dict] | None) -> list[dict]:
        pairs: list[tuple[str, str]] = []
        for item in raw or []:
            if not isinstance(item, dict):
                continue
            name = str(item.get("name") or "").strip()
            if not name:
                continue
            value = str(item.get("value") or "")
            pairs.append((name, value))
        return AssetService._sanitize_parsed_header_pairs(pairs)

    @staticmethod
    def _parse_request_target_token(path_token: str) -> tuple[str, list[tuple[str, str]]]:
        token = path_token.strip()
        if token.startswith("http://") or token.startswith("https://"):
            u = urlparse(token)
            return u.path or "/", list(parse_qsl(u.query, keep_blank_values=True))
        fake = "http://stub.local" + (token if token.startswith("/") else "/" + token)
        u = urlparse(fake)
        return u.path or "/", list(parse_qsl(u.query, keep_blank_values=True))

    @staticmethod
    def _apply_structured_request_payload(payload: dict) -> dict:
        if "path" in payload and payload.get("path") is not None:
            payload["path"] = AssetService._normalize_endpoint_path(payload.get("path"))
        if "query_params" in payload:
            payload["query_params"] = AssetService._normalize_endpoint_query_params(payload.get("query_params"))
        if "request_headers" in payload and payload.get("request_headers") is not None:
            normalized = AssetService._normalize_endpoint_headers(payload.get("request_headers"))
            payload["request_headers"] = AssetService._sanitize_stored_header_items(normalized)
        if "request_body" in payload and payload.get("request_body") is not None:
            request_body = str(payload.get("request_body", ""))
            payload["request_body"] = request_body if request_body.strip() else None
        if "request_content_type" in payload and payload.get("request_content_type") is not None:
            content_type = str(payload.get("request_content_type", "")).strip()
            payload["request_content_type"] = content_type or None
        return payload

    @staticmethod
    def _apply_raw_request_payload(payload: dict) -> dict:
        request_raw = payload.get("request_raw")
        if not request_raw:
            payload.pop("request_raw", None)
            return payload
        raw_request = str(request_raw).replace("\r", "")
        parts = re.split(r"\n\s*\n", raw_request, maxsplit=1)
        header_block = parts[0]
        body = parts[1] if len(parts) > 1 else ""
        header_lines = header_block.split("\n")
        if not header_lines or not header_lines[0].strip():
            raise ValidationError("request_raw пустой")
        request_line = header_lines[0].split()
        if len(request_line) < 3:
            raise ValidationError("request_raw должен содержать request line вида 'METHOD /path HTTP/1.1'")
        method, path_value, http_part = request_line[0].upper(), request_line[1], request_line[2].upper()
        allowed_methods = {"GET", "POST", "PUT", "PATCH", "DELETE", "HEAD", "OPTIONS"}
        if method not in allowed_methods:
            raise ValidationError("request_raw содержит неподдерживаемый HTTP-метод")
        if not http_part.startswith("HTTP/"):
            raise ValidationError("request_raw должен содержать HTTP-версию в request line")
        path_only, query_pairs = AssetService._parse_request_target_token(path_value)
        payload["method"] = method
        payload["path"] = AssetService._normalize_endpoint_path(path_only)
        if "query_params" not in payload or not payload.get("query_params"):
            payload["query_params"] = [
                {"name": key, "value": value, "required": False, "description": None} for key, value in query_pairs
            ]
        header_pairs: list[tuple[str, str]] = []
        for line in header_lines[1:]:
            line = line.strip()
            if not line or ":" not in line:
                continue
            h_name, _, h_val = line.partition(":")
            header_pairs.append((h_name.strip(), h_val))
        content_type = None
        for name, value in header_pairs:
            if name.lower() == "content-type":
                content_type = value.strip() or None
                break
        if ("request_body" not in payload or payload.get("request_body") is None) and body.strip():
            payload["request_body"] = body
        if ("request_content_type" not in payload or payload.get("request_content_type") is None) and content_type:
            payload["request_content_type"] = content_type
        if "request_headers" not in payload or not payload.get("request_headers"):
            payload["request_headers"] = AssetService._sanitize_parsed_header_pairs(header_pairs)
        if method in AssetService._METHODS_WITHOUT_BODY:
            payload["request_body"] = None
            payload["request_content_type"] = None
        payload.pop("request_raw", None)
        return payload

    @staticmethod
    def _merge_endpoint_fields(endpoint: Endpoint, endpoint_payload: dict) -> None:
        if not endpoint.description and endpoint_payload.get("description"):
            endpoint.description = endpoint_payload["description"]
        if not endpoint.query_params and endpoint_payload.get("query_params"):
            endpoint.query_params = endpoint_payload["query_params"]
        if not endpoint.request_body and endpoint_payload.get("request_body"):
            endpoint.request_body = endpoint_payload["request_body"]
        if not endpoint.request_content_type and endpoint_payload.get("request_content_type"):
            endpoint.request_content_type = endpoint_payload["request_content_type"]
        if not endpoint.request_headers and endpoint_payload.get("request_headers"):
            endpoint.request_headers = endpoint_payload["request_headers"]

    async def _get_host(self, project_id: UUID, host_id: UUID) -> Host:
        host = await self.db.scalar(select(Host).where(and_(Host.id == host_id, Host.project_id == project_id)))
        if not host:
            raise NotFoundError("Хост не найден")
        return host

    async def _get_port(self, host_id: UUID, port_id: UUID) -> Port:
        port = await self.db.scalar(select(Port).where(and_(Port.id == port_id, Port.host_id == host_id)))
        if not port:
            raise NotFoundError("Порт не найден")
        return port

    async def list_hosts(self, project_id: UUID, page: int, size: int, status: str | None) -> tuple[list[Host], int]:
        query = select(Host).where(Host.project_id == project_id)
        if status:
            query = query.where(Host.status == status)
        total = await self.db.scalar(select(func.count()).select_from(query.subquery())) or 0
        items = (await self.db.scalars(query.order_by(Host.created_at.desc()).offset((page - 1) * size).limit(size))).all()
        return list(items), total

    async def create_host(self, project_id: UUID, payload: dict, actor_id: UUID) -> Host:
        host = Host(project_id=project_id, **payload)
        self.db.add(host)
        await self.db.commit()
        await self.db.refresh(host)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="host", entity_id=host.id)
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "host", "project_id": str(project_id), "data": {"id": str(host.id)}})
        return host

    async def get_host(self, project_id: UUID, host_id: UUID) -> dict:
        host = await self._get_host(project_id, host_id)
        ports = (await self.db.scalars(select(Port).where(Port.host_id == host.id))).all()
        endpoints = (await self.db.scalars(select(Endpoint).where(Endpoint.host_id == host.id))).all()
        return {
            "id": host.id,
            "project_id": host.project_id,
            "ip_address": host.ip_address,
            "hostname": host.hostname,
            "status": host.status,
            "notes": host.notes,
            "created_at": host.created_at,
            "updated_at": host.updated_at,
            "ports": [
                {
                    "id": port.id,
                    "host_id": port.host_id,
                    "port_number": port.port_number,
                    "protocol": port.protocol,
                    "state": port.state,
                    "created_at": port.created_at,
                    "updated_at": port.updated_at,
                }
                for port in ports
            ],
            "endpoints": [
                {
                    "id": endpoint.id,
                    "host_id": endpoint.host_id,
                    "path": endpoint.path,
                    "method": endpoint.method,
                    "description": endpoint.description,
                    "query_params": endpoint.query_params or [],
                    "request_body": endpoint.request_body,
                    "request_content_type": endpoint.request_content_type,
                    "request_headers": endpoint.request_headers or [],
                    "created_at": endpoint.created_at,
                    "updated_at": endpoint.updated_at,
                }
                for endpoint in endpoints
            ],
        }

    async def update_host(self, project_id: UUID, host_id: UUID, payload: dict, actor_id: UUID) -> Host:
        host = await self._get_host(project_id, host_id)
        for key, value in payload.items():
            if value is not None:
                setattr(host, key, value)
        await self.db.commit()
        await self.db.refresh(host)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="host", entity_id=host.id)
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "host", "project_id": str(project_id), "data": {"id": str(host.id)}})
        return host

    async def delete_host(self, project_id: UUID, host_id: UUID, actor_id: UUID) -> None:
        host = await self._get_host(project_id, host_id)
        await self.db.delete(host)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="host", entity_id=host.id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "host", "project_id": str(project_id), "data": {"id": str(host.id)}})

    async def list_ports(self, host_id: UUID) -> list[Port]:
        return list((await self.db.scalars(select(Port).where(Port.host_id == host_id).order_by(Port.port_number))).all())

    async def create_port(self, project_id: UUID, host_id: UUID, payload: dict, actor_id: UUID) -> Port:
        await self._get_host(project_id, host_id)
        duplicate = await self.db.scalar(
            select(Port).where(
                and_(
                    Port.host_id == host_id,
                    Port.port_number == payload["port_number"],
                    Port.protocol == payload["protocol"],
                )
            )
        )
        if duplicate:
            raise ConflictError("Порт с таким номером и протоколом уже существует")
        port = Port(host_id=host_id, **payload)
        self.db.add(port)
        await self.db.commit()
        await self.db.refresh(port)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="port", entity_id=port.id)
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "port", "project_id": str(project_id), "data": {"id": str(port.id)}})
        return port

    async def get_port(self, host_id: UUID, port_id: UUID) -> Port:
        return await self._get_port(host_id, port_id)

    async def update_port(self, project_id: UUID, host_id: UUID, port_id: UUID, payload: dict, actor_id: UUID) -> Port:
        port = await self._get_port(host_id, port_id)
        next_port_number = payload.get("port_number", port.port_number)
        next_protocol = payload.get("protocol", port.protocol)
        duplicate = await self.db.scalar(
            select(Port).where(
                and_(
                    Port.host_id == host_id,
                    Port.id != port.id,
                    Port.port_number == next_port_number,
                    Port.protocol == next_protocol,
                )
            )
        )
        if duplicate:
            raise ConflictError("Порт с таким номером и протоколом уже существует")
        for key, value in payload.items():
            if value is not None:
                setattr(port, key, value)
        await self.db.commit()
        await self.db.refresh(port)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="port", entity_id=port.id)
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "port", "project_id": str(project_id), "data": {"id": str(port.id)}})
        return port

    async def delete_port(self, project_id: UUID, host_id: UUID, port_id: UUID, actor_id: UUID) -> None:
        port = await self._get_port(host_id, port_id)
        await self.db.delete(port)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="port", entity_id=port.id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "port", "project_id": str(project_id), "data": {"id": str(port.id)}})

    async def list_services(self, port_id: UUID) -> list[Service]:
        return list((await self.db.scalars(select(Service).where(Service.port_id == port_id).order_by(Service.created_at.desc()))).all())

    async def create_service(self, project_id: UUID, host_id: UUID, port_id: UUID, payload: dict, actor_id: UUID) -> Service:
        await self._get_host(project_id, host_id)
        await self._get_port(host_id, port_id)
        service = Service(port_id=port_id, **payload)
        self.db.add(service)
        await self.db.commit()
        await self.db.refresh(service)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="service", entity_id=service.id)
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "service", "project_id": str(project_id), "data": {"id": str(service.id)}})
        return service

    async def update_service(
        self, project_id: UUID, host_id: UUID, port_id: UUID, service_id: UUID, payload: dict, actor_id: UUID
    ) -> Service:
        await self._get_host(project_id, host_id)
        await self._get_port(host_id, port_id)
        service = await self.db.scalar(select(Service).where(and_(Service.id == service_id, Service.port_id == port_id)))
        if not service:
            raise NotFoundError("Сервис не найден")
        for key, value in payload.items():
            if value is not None:
                setattr(service, key, value)
        await self.db.commit()
        await self.db.refresh(service)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="service", entity_id=service.id)
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "service", "project_id": str(project_id), "data": {"id": str(service.id)}})
        return service

    async def delete_service(self, project_id: UUID, port_id: UUID, service_id: UUID, actor_id: UUID) -> None:
        service = await self.db.scalar(select(Service).where(and_(Service.id == service_id, Service.port_id == port_id)))
        if not service:
            raise NotFoundError("Сервис не найден")
        await self.db.delete(service)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="service", entity_id=service.id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "service", "project_id": str(project_id), "data": {"id": str(service.id)}})

    async def list_endpoints(self, host_id: UUID) -> list[Endpoint]:
        return list((await self.db.scalars(select(Endpoint).where(Endpoint.host_id == host_id).order_by(Endpoint.created_at.desc()))).all())

    async def create_endpoint(self, project_id: UUID, host_id: UUID, payload: dict, actor_id: UUID) -> Endpoint:
        await self._get_host(project_id, host_id)
        payload = self._apply_structured_request_payload(self._apply_raw_request_payload(dict(payload)))
        if payload.get("request_headers") is None:
            payload["request_headers"] = []
        duplicate = await self.db.scalar(
            select(Endpoint).where(
                and_(
                    Endpoint.host_id == host_id,
                    Endpoint.path == payload.get("path"),
                    Endpoint.method == payload.get("method"),
                )
            )
        )
        if duplicate:
            self._merge_endpoint_fields(duplicate, payload)
            await self.db.commit()
            await self.db.refresh(duplicate)
            await self.audit.log("UPDATE", user_id=actor_id, entity_type="endpoint", entity_id=duplicate.id)
            await ws_manager.broadcast(project_id, {"event": "updated", "entity": "endpoint", "project_id": str(project_id), "data": {"id": str(duplicate.id)}})
            return duplicate
        endpoint = Endpoint(host_id=host_id, **payload)
        self.db.add(endpoint)
        await self.db.commit()
        await self.db.refresh(endpoint)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="endpoint", entity_id=endpoint.id)
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "endpoint", "project_id": str(project_id), "data": {"id": str(endpoint.id)}})
        return endpoint

    async def update_endpoint(self, project_id: UUID, host_id: UUID, endpoint_id: UUID, payload: dict, actor_id: UUID) -> Endpoint:
        await self._get_host(project_id, host_id)
        payload = self._apply_structured_request_payload(self._apply_raw_request_payload(dict(payload)))
        endpoint = await self.db.scalar(select(Endpoint).where(and_(Endpoint.id == endpoint_id, Endpoint.host_id == host_id)))
        if not endpoint:
            raise NotFoundError("Endpoint не найден")
        next_path = payload.get("path") if payload.get("path") is not None else endpoint.path
        next_method = payload.get("method") if payload.get("method") is not None else endpoint.method
        duplicate = await self.db.scalar(
            select(Endpoint).where(
                and_(
                    Endpoint.host_id == host_id,
                    Endpoint.id != endpoint_id,
                    Endpoint.path == next_path,
                    Endpoint.method == next_method,
                )
            )
        )
        if duplicate:
            raise ConflictError("Эндпоинт с таким методом и path уже существует")
        for key, value in payload.items():
            if key in {"description", "request_body", "request_content_type", "query_params", "request_headers"} or value is not None:
                setattr(endpoint, key, value)
        await self.db.commit()
        await self.db.refresh(endpoint)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="endpoint", entity_id=endpoint.id)
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "endpoint", "project_id": str(project_id), "data": {"id": str(endpoint.id)}})
        return endpoint

    async def delete_endpoint(self, project_id: UUID, host_id: UUID, endpoint_id: UUID, actor_id: UUID) -> None:
        endpoint = await self.db.scalar(select(Endpoint).where(and_(Endpoint.id == endpoint_id, Endpoint.host_id == host_id)))
        if not endpoint:
            raise NotFoundError("Endpoint не найден")
        await self.db.delete(endpoint)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="endpoint", entity_id=endpoint.id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "endpoint", "project_id": str(project_id), "data": {"id": str(endpoint.id)}})


class VulnerabilityService:
    """Сервис управления уязвимостями и привязками активов."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.audit = AuditService(db)

    async def _get_vuln(self, project_id: UUID, vuln_id: UUID) -> Vulnerability:
        vuln = await self.db.scalar(select(Vulnerability).where(and_(Vulnerability.id == vuln_id, Vulnerability.project_id == project_id)))
        if not vuln:
            raise NotFoundError("Уязвимость не найдена")
        self._hydrate_workflow_steps(vuln)
        return vuln

    @staticmethod
    def _hydrate_workflow_steps(vuln: Vulnerability) -> None:
        if vuln.workflow_steps is not None:
            return
        hydrated_steps: list[dict]
        if vuln.steps_to_reproduce:
            hydrated_steps = [
                {
                    "id": str(uuid4()),
                    "description": vuln.steps_to_reproduce,
                    "image_file_ids": [],
                }
            ]
        else:
            hydrated_steps = []
        if hasattr(vuln, "_sa_instance_state"):
            set_committed_value(vuln, "workflow_steps", hydrated_steps)
        else:
            vuln.workflow_steps = hydrated_steps

    @staticmethod
    def _normalize_workflow_steps(steps: list[dict] | None) -> list[dict]:
        normalized: list[dict] = []
        raw_steps = steps or []
        # #region agent log
        _debug_log(
            "H3",
            "backend/app/services.py:_normalize_workflow_steps:start",
            "Normalizing workflow steps",
            {
                "raw_count": len(raw_steps),
                "raw_with_legacy_title_count": sum(1 for step in raw_steps if str(step.get("title", "")).strip()),
                "raw_has_description_count": sum(1 for step in raw_steps if str(step.get("description", "")).strip()),
            },
        )
        # #endregion
        for raw_step in raw_steps:
            description = str(raw_step.get("description", "")).strip()
            endpoint_request_raw = str(raw_step.get("endpoint_request_raw", "")).strip()
            endpoint_id = raw_step.get("endpoint_id")
            image_file_ids = [str(file_id) for file_id in raw_step.get("image_file_ids", []) if file_id]
            if not description and not image_file_ids and not endpoint_id and not endpoint_request_raw:
                continue
            normalized.append(
                {
                    "id": str(raw_step.get("id") or uuid4()),
                    "description": description or None,
                    "image_file_ids": image_file_ids,
                    "endpoint_id": str(endpoint_id) if endpoint_id else None,
                    "endpoint_request_raw": endpoint_request_raw or None,
                }
            )
        # #region agent log
        _debug_log(
            "H3",
            "backend/app/services.py:_normalize_workflow_steps:end",
            "Workflow steps normalized",
            {
                "normalized_count": len(normalized),
                "normalized_with_title_count": sum(1 for step in normalized if str(step.get("title", "")).strip()),
            },
        )
        # #endregion
        return normalized

    @staticmethod
    def _workflow_steps_to_text(steps: list[dict]) -> str | None:
        if not steps:
            return None
        blocks: list[str] = []
        for index, step in enumerate(steps, start=1):
            description = str(step.get("description", "")).strip()
            block = f"{index}. Этап {index}"
            if description:
                block = f"{block}\n{description}"
            if step.get("endpoint_id") or step.get("endpoint_request_raw"):
                block = f"{block}\n[Endpoint: привязан]"
            if step.get("image_file_ids"):
                block = f"{block}\n[Изображений: {len(step['image_file_ids'])}]"
            blocks.append(block)
        return "\n\n".join(blocks)

    @staticmethod
    def _normalize_cvss_vector(version: str | None, vector: str | None) -> str | None:
        version_value = getattr(version, "value", version)
        if not version or not vector:
            return None
        raw = vector.strip()
        if not raw:
            return None
        normalized = re.sub(r"^CVSS:\d\.\d/", f"CVSS:{version_value}/", raw)
        if not normalized.startswith("CVSS:"):
            normalized = f"CVSS:{version_value}/{normalized.lstrip('/')}"
        return normalized

    @staticmethod
    def _calculate_cvss_score(version: str | None, vector: str | None) -> tuple[str | None, float | None]:
        version_value = getattr(version, "value", version)
        normalized_vector = VulnerabilityService._normalize_cvss_vector(version, vector)
        if not normalized_vector:
            return None, None
        try:
            if version_value == "4.0":
                score = float(CVSS4(normalized_vector).scores()[0])
            else:
                raise ValidationError("Поддерживается только CVSS 4.0")
        except Exception as exc:
            raise ValidationError(f"Некорректный CVSS вектор: {exc}") from exc
        return normalized_vector, score

    @staticmethod
    def _apply_calculated_cvss_fields(payload: dict, *, current_version: str | None = None, current_vector: str | None = None) -> None:
        has_explicit_score = payload.get("cvss_score") is not None
        next_version = getattr(payload.get("cvss_version", current_version), "value", payload.get("cvss_version", current_version))
        next_vector = payload.get("cvss_vector", current_vector)

        if "cvss_vector" in payload and not payload.get("cvss_vector"):
            if has_explicit_score:
                raise ValidationError("CVSS score рассчитывается автоматически и требует корректный CVSS 4.0 вектор")
            payload["cvss_version"] = None
            payload["cvss_vector"] = None
            payload["cvss_score"] = None
            return

        if next_vector:
            if not next_version:
                raise ValidationError("Для расчёта CVSS укажите версию 4.0 и корректный вектор")
            normalized_vector, calculated_score = VulnerabilityService._calculate_cvss_score(next_version, next_vector)
            payload["cvss_version"] = next_version
            payload["cvss_vector"] = normalized_vector
            payload["cvss_score"] = calculated_score
            return

        if has_explicit_score:
            raise ValidationError("CVSS score рассчитывается автоматически и требует корректный CVSS 4.0 вектор")

        if "cvss_version" in payload and payload.get("cvss_version"):
            raise ValidationError("Для расчёта CVSS укажите корректный CVSS 4.0 вектор")

    @staticmethod
    def _severity_from_cvss_score(score: float | None) -> Severity:
        if score is None:
            return Severity.INFO
        if score >= 9.0:
            return Severity.CRITICAL
        if score >= 7.0:
            return Severity.HIGH
        if score >= 4.0:
            return Severity.MEDIUM
        if score > 0:
            return Severity.LOW
        return Severity.INFO

    async def _resolve_workflow_step_endpoints(self, project_id: UUID, host_id: UUID, steps: list[dict]) -> list[dict]:
        resolved_steps: list[dict] = []
        for step in steps:
            next_step = dict(step)
            endpoint_id_raw = next_step.get("endpoint_id")
            endpoint_request_raw = next_step.get("endpoint_request_raw")
            endpoint_id = UUID(str(endpoint_id_raw)) if endpoint_id_raw else None
            if endpoint_id:
                endpoint = await self.db.scalar(
                    select(Endpoint).where(and_(Endpoint.id == endpoint_id, Endpoint.host_id == host_id))
                )
                if not endpoint:
                    raise ValidationError("Выбранный endpoint не принадлежит текущему хосту")
                next_step["endpoint_id"] = str(endpoint.id)
            if endpoint_request_raw:
                endpoint_payload = AssetService._apply_structured_request_payload(
                    AssetService._apply_raw_request_payload({"request_raw": endpoint_request_raw})
                )
                existing_endpoint = await self.db.scalar(
                    select(Endpoint).where(
                        and_(
                            Endpoint.host_id == host_id,
                            Endpoint.path == endpoint_payload["path"],
                            Endpoint.method == endpoint_payload.get("method"),
                        )
                    )
                )
                endpoint = existing_endpoint
                if endpoint is None:
                    if endpoint_payload.get("request_headers") is None:
                        endpoint_payload["request_headers"] = []
                    endpoint = Endpoint(host_id=host_id, **endpoint_payload)
                    self.db.add(endpoint)
                    await self.db.flush()
                next_step["endpoint_id"] = str(endpoint.id)
                next_step["endpoint_request_raw"] = str(endpoint_request_raw).strip()
            resolved_steps.append(next_step)
        return resolved_steps

    async def _get_primary_host_id(self, vuln_id: UUID) -> UUID:
        host_id = await self.db.scalar(
            select(VulnerabilityAsset.asset_id)
            .where(and_(VulnerabilityAsset.vulnerability_id == vuln_id, VulnerabilityAsset.asset_type == AssetType.HOST))
            .limit(1)
        )
        if not host_id:
            raise ValidationError("Уязвимость должна быть привязана хотя бы к одному хосту")
        return host_id

    async def _validate_workflow_step_images(self, vulnerability_id: UUID | None, steps: list[dict] | None) -> None:
        image_ids = {
            UUID(str(file_id))
            for step in steps or []
            for file_id in step.get("image_file_ids", [])
            if file_id
        }
        if not image_ids:
            return
        if vulnerability_id is None:
            raise ValidationError("Нельзя указывать image_file_ids до загрузки файлов уязвимости")
        existing_ids = set(
            (
                await self.db.scalars(
                    select(File.id).where(and_(File.vulnerability_id == vulnerability_id, File.id.in_(list(image_ids))))
                )
            ).all()
        )
        missing_ids = sorted(str(file_id) for file_id in image_ids - existing_ids)
        if missing_ids:
            raise ValidationError(
                f"workflow_steps.image_file_ids содержат файлы, не принадлежащие уязвимости: {', '.join(missing_ids)}"
            )

    async def list(self, project_id: UUID, page: int, size: int, severity: str | None, status: str | None) -> tuple[list[Vulnerability], int]:
        query = select(Vulnerability).where(Vulnerability.project_id == project_id)
        if severity:
            query = query.where(Vulnerability.severity == severity)
        if status:
            query = query.where(Vulnerability.status == status)
        total = await self.db.scalar(select(func.count()).select_from(query.subquery())) or 0
        items = (await self.db.scalars(query.order_by(Vulnerability.created_at.desc()).offset((page - 1) * size).limit(size))).all()
        for item in items:
            self._hydrate_workflow_steps(item)
        return list(items), total

    async def list_for_host(
        self, project_id: UUID, host_id: UUID, page: int, size: int, severity: str | None, status: str | None
    ) -> tuple[list[Vulnerability], int]:
        host_exists = await self.db.scalar(select(Host.id).where(and_(Host.id == host_id, Host.project_id == project_id)))
        if not host_exists:
            raise NotFoundError("Хост не найден")

        query = (
            select(Vulnerability)
            .join(VulnerabilityAsset, VulnerabilityAsset.vulnerability_id == Vulnerability.id)
            .where(
                and_(
                    Vulnerability.project_id == project_id,
                    VulnerabilityAsset.asset_type == AssetType.HOST,
                    VulnerabilityAsset.asset_id == host_id,
                )
            )
        )
        if severity:
            query = query.where(Vulnerability.severity == severity)
        if status:
            query = query.where(Vulnerability.status == status)
        total = await self.db.scalar(select(func.count()).select_from(query.subquery())) or 0
        items = (await self.db.scalars(query.order_by(Vulnerability.created_at.desc()).offset((page - 1) * size).limit(size))).all()
        for item in items:
            self._hydrate_workflow_steps(item)
        return list(items), total

    async def create(self, project_id: UUID, payload: dict, actor_id: UUID) -> Vulnerability:
        host_id = payload.pop("host_id", None)
        if not host_id:
            raise ValidationError("Уязвимость должна быть привязана к конкретному хосту")
        await self._assert_asset_in_project(project_id, AssetType.HOST, host_id)
        self._apply_calculated_cvss_fields(payload)
        if payload.get("cvss_score") is not None:
            payload["severity"] = self._severity_from_cvss_score(payload.get("cvss_score"))
        if "workflow_steps" in payload:
            payload["workflow_steps"] = self._normalize_workflow_steps(payload.get("workflow_steps"))
            payload["workflow_steps"] = await self._resolve_workflow_step_endpoints(project_id, host_id, payload["workflow_steps"])
            await self._validate_workflow_step_images(None, payload["workflow_steps"])
            payload["steps_to_reproduce"] = self._workflow_steps_to_text(payload["workflow_steps"])
        vuln = Vulnerability(project_id=project_id, created_by=actor_id, **payload)
        self.db.add(vuln)
        await self.db.flush()
        self.db.add(VulnerabilityAsset(vulnerability_id=vuln.id, asset_type=AssetType.HOST, asset_id=host_id))
        await self.db.commit()
        await self.db.refresh(vuln)
        await self.audit.log(
            "CREATE",
            user_id=actor_id,
            entity_type="vulnerability",
            entity_id=vuln.id,
            details={"host_id": str(host_id)},
        )
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "vulnerability", "project_id": str(project_id), "data": {"id": str(vuln.id)}})
        return vuln

    async def get(self, project_id: UUID, vuln_id: UUID) -> dict:
        vuln = await self._get_vuln(project_id, vuln_id)
        links = (await self.db.scalars(select(VulnerabilityAsset).where(VulnerabilityAsset.vulnerability_id == vuln.id))).all()
        files = (await self.db.scalars(select(File).where(File.vulnerability_id == vuln.id))).all()
        comments_count = await self.db.scalar(select(func.count()).select_from(Comment).where(Comment.vulnerability_id == vuln.id)) or 0
        return {"vulnerability": vuln, "assets": list(links), "files": list(files), "comments_count": comments_count}

    async def update(self, project_id: UUID, vuln_id: UUID, payload: dict, actor_id: UUID) -> Vulnerability:
        vuln = await self._get_vuln(project_id, vuln_id)
        self._apply_calculated_cvss_fields(
            payload,
            current_version=vuln.cvss_version.value if vuln.cvss_version else None,
            current_vector=vuln.cvss_vector,
        )
        next_score = payload.get("cvss_score", vuln.cvss_score)
        if next_score is not None:
            payload["severity"] = self._severity_from_cvss_score(next_score)
        if "workflow_steps" in payload:
            payload["workflow_steps"] = self._normalize_workflow_steps(payload.get("workflow_steps"))
            host_id = await self._get_primary_host_id(vuln.id)
            payload["workflow_steps"] = await self._resolve_workflow_step_endpoints(project_id, host_id, payload["workflow_steps"])
            await self._validate_workflow_step_images(vuln.id, payload["workflow_steps"])
            payload["steps_to_reproduce"] = self._workflow_steps_to_text(payload["workflow_steps"])
        clearable_fields = {"description", "cvss_version", "cvss_score", "cvss_vector", "cwe_id", "impact", "recommendations"}
        for key, value in payload.items():
            if key == "steps_to_reproduce" and "workflow_steps" in payload:
                setattr(vuln, key, value)
            elif key in clearable_fields or value is not None:
                setattr(vuln, key, value)
        await self.db.commit()
        await self.db.refresh(vuln)
        await self.audit.log("UPDATE", user_id=actor_id, entity_type="vulnerability", entity_id=vuln.id)
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "vulnerability", "project_id": str(project_id), "data": {"id": str(vuln.id)}})
        return vuln

    async def patch_status(self, project_id: UUID, vuln_id: UUID, status: str, actor_id: UUID) -> Vulnerability:
        vuln = await self._get_vuln(project_id, vuln_id)
        old_status = vuln.status
        vuln.status = status
        await self.db.commit()
        await self.db.refresh(vuln)
        await self.audit.log(
            "STATUS_CHANGE",
            user_id=actor_id,
            entity_type="vulnerability",
            entity_id=vuln.id,
            details={"old_status": old_status.value, "new_status": vuln.status.value},
        )
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "vulnerability", "project_id": str(project_id), "data": {"id": str(vuln.id)}})
        return vuln

    async def delete(self, project_id: UUID, vuln_id: UUID, actor_id: UUID) -> None:
        vuln = await self._get_vuln(project_id, vuln_id)
        await self.db.delete(vuln)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="vulnerability", entity_id=vuln_id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "vulnerability", "project_id": str(project_id), "data": {"id": str(vuln_id)}})

    async def list_assets(self, project_id: UUID, vuln_id: UUID) -> list[VulnerabilityAsset]:
        await self._get_vuln(project_id, vuln_id)
        return list((await self.db.scalars(select(VulnerabilityAsset).where(VulnerabilityAsset.vulnerability_id == vuln_id))).all())

    async def _assert_asset_in_project(self, project_id: UUID, asset_type: AssetType, asset_id: UUID) -> None:
        if asset_type == AssetType.HOST:
            found = await self.db.scalar(select(Host.id).where(and_(Host.id == asset_id, Host.project_id == project_id)))
        elif asset_type == AssetType.PORT:
            found = await self.db.scalar(
                select(Port.id).join(Host, Port.host_id == Host.id).where(and_(Port.id == asset_id, Host.project_id == project_id))
            )
        elif asset_type == AssetType.SERVICE:
            found = await self.db.scalar(
                select(Service.id)
                .join(Port, Service.port_id == Port.id)
                .join(Host, Port.host_id == Host.id)
                .where(and_(Service.id == asset_id, Host.project_id == project_id))
            )
        else:
            found = await self.db.scalar(
                select(Endpoint.id).join(Host, Endpoint.host_id == Host.id).where(and_(Endpoint.id == asset_id, Host.project_id == project_id))
            )
        if not found:
            raise ValidationError("Актив не найден или принадлежит другому проекту")

    async def add_asset(self, project_id: UUID, vuln_id: UUID, asset_type: AssetType, asset_id: UUID, actor_id: UUID) -> VulnerabilityAsset:
        await self._get_vuln(project_id, vuln_id)
        await self._assert_asset_in_project(project_id, asset_type, asset_id)
        duplicate = await self.db.scalar(
            select(VulnerabilityAsset).where(
                and_(
                    VulnerabilityAsset.vulnerability_id == vuln_id,
                    VulnerabilityAsset.asset_type == asset_type,
                    VulnerabilityAsset.asset_id == asset_id,
                )
            )
        )
        if duplicate:
            raise ConflictError("Связь уже существует")
        link = VulnerabilityAsset(vulnerability_id=vuln_id, asset_type=asset_type, asset_id=asset_id)
        self.db.add(link)
        await self.db.commit()
        await self.db.refresh(link)
        await self.audit.log("CREATE", user_id=actor_id, entity_type="vulnerability_asset", entity_id=link.id)
        await ws_manager.broadcast(
            project_id,
            {
                "event": "updated",
                "entity": "vulnerability",
                "project_id": str(project_id),
                "data": {"id": str(vuln_id), "asset_link_id": str(link.id)},
            },
        )
        return link

    async def delete_asset(self, project_id: UUID, vuln_id: UUID, link_id: UUID, actor_id: UUID) -> None:
        await self._get_vuln(project_id, vuln_id)
        link = await self.db.scalar(
            select(VulnerabilityAsset).where(and_(VulnerabilityAsset.id == link_id, VulnerabilityAsset.vulnerability_id == vuln_id))
        )
        if not link:
            raise NotFoundError("Связь не найдена")
        if link.asset_type == AssetType.HOST:
            host_links_count = await self.db.scalar(
                select(func.count())
                .select_from(VulnerabilityAsset)
                .where(and_(VulnerabilityAsset.vulnerability_id == vuln_id, VulnerabilityAsset.asset_type == AssetType.HOST))
            )
            if (host_links_count or 0) <= 1:
                raise ValidationError("Уязвимость должна оставаться привязанной хотя бы к одному хосту")
        await self.db.delete(link)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor_id, entity_type="vulnerability_asset", entity_id=link_id)
        await ws_manager.broadcast(
            project_id,
            {
                "event": "updated",
                "entity": "vulnerability",
                "project_id": str(project_id),
                "data": {"id": str(vuln_id), "asset_link_id": str(link_id)},
            },
        )


class FileService:
    """Сервис загрузки и выдачи файлов уязвимостей."""

    def __init__(self, db: AsyncSession, storage: MinioStorage | None = None) -> None:
        self.db = db
        self.audit = AuditService(db)
        self.storage = storage or MinioStorage()
        self.storage.ensure_bucket()

    async def _ensure_vuln(self, project_id: UUID, vuln_id: UUID) -> Vulnerability:
        vuln = await self.db.scalar(select(Vulnerability).where(and_(Vulnerability.id == vuln_id, Vulnerability.project_id == project_id)))
        if not vuln:
            raise NotFoundError("Уязвимость не найдена")
        return vuln

    async def list(self, project_id: UUID, vuln_id: UUID) -> list[File]:
        await self._ensure_vuln(project_id, vuln_id)
        return list((await self.db.scalars(select(File).where(File.vulnerability_id == vuln_id).order_by(File.uploaded_at.desc()))).all())

    async def upload(self, project_id: UUID, vuln_id: UUID, upload: UploadFile, actor_id: UUID) -> File:
        await self._ensure_vuln(project_id, vuln_id)
        content = await upload.read()
        if len(content) > MAX_FILE_SIZE:
            raise ValidationError("Размер файла превышает 50 МБ")
        mime_type = magic.from_buffer(content, mime=True)
        if mime_type not in ALLOWED_MIME_TYPES:
            raise ValidationError("Неподдерживаемый тип файла")
        object_key = await asyncio.to_thread(self.storage.upload_bytes, content, mime_type, upload.filename or "file.bin")
        file_meta = File(
            vulnerability_id=vuln_id,
            original_name=upload.filename or "file.bin",
            content_type=mime_type,
            size_bytes=len(content),
            minio_bucket=settings.minio_bucket_name,
            minio_key=object_key,
            uploaded_by=actor_id,
        )
        self.db.add(file_meta)
        await self.db.commit()
        await self.db.refresh(file_meta)
        await self.audit.log("FILE_UPLOAD", user_id=actor_id, entity_type="file", entity_id=file_meta.id)
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "file", "project_id": str(project_id), "data": {"id": str(file_meta.id)}})
        return file_meta

    async def download(self, file_id: UUID, current_user: User) -> tuple[File, bytes]:
        file_meta = await self.db.scalar(select(File).where(File.id == file_id))
        if not file_meta:
            raise NotFoundError("Файл не найден")
        vuln = await self.db.scalar(select(Vulnerability).where(Vulnerability.id == file_meta.vulnerability_id))
        if not vuln:
            raise NotFoundError("Уязвимость не найдена")
        if current_user.role != UserRole.ADMIN:
            membership = await self.db.scalar(
                select(ProjectMember).where(
                    and_(ProjectMember.project_id == vuln.project_id, ProjectMember.user_id == current_user.id)
                )
            )
            if not membership:
                raise ForbiddenError("Нет доступа к файлу")
        blob = await asyncio.to_thread(self.storage.download_bytes, file_meta.minio_key)
        return file_meta, blob

    async def delete(self, project_id: UUID, vuln_id: UUID, file_id: UUID, actor_id: UUID) -> None:
        await self._ensure_vuln(project_id, vuln_id)
        file_meta = await self.db.scalar(
            select(File).where(and_(File.id == file_id, File.vulnerability_id == vuln_id))
        )
        if not file_meta:
            raise NotFoundError("Файл не найден")
        await self.db.delete(file_meta)
        await self.db.commit()
        await asyncio.to_thread(self.storage.delete, file_meta.minio_key)
        await self.audit.log("FILE_DELETE", user_id=actor_id, entity_type="file", entity_id=file_id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "file", "project_id": str(project_id), "data": {"id": str(file_id)}})


class CommentService:
    """Сервис комментариев и @упоминаний."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.audit = AuditService(db)

    async def _ensure_vuln(self, project_id: UUID, vuln_id: UUID) -> Vulnerability:
        vuln = await self.db.scalar(select(Vulnerability).where(and_(Vulnerability.id == vuln_id, Vulnerability.project_id == project_id)))
        if not vuln:
            raise NotFoundError("Уязвимость не найдена")
        return vuln

    async def _extract_mentions(self, project_id: UUID, content: str) -> list[User]:
        usernames = set(MENTION_RE.findall(content))
        if not usernames:
            return []
        rows = (
            await self.db.scalars(
                select(User)
                .join(ProjectMember, ProjectMember.user_id == User.id)
                .where(and_(ProjectMember.project_id == project_id, User.username.in_(list(usernames))))
            )
        ).all()
        return list(rows)

    async def list(self, vuln_id: UUID, page: int, size: int) -> tuple[list[CommentOut], int]:
        total = await self.db.scalar(select(func.count()).select_from(Comment).where(Comment.vulnerability_id == vuln_id)) or 0
        rows = (
            await self.db.execute(
                select(Comment, User)
                .join(User, User.id == Comment.user_id)
                .where(Comment.vulnerability_id == vuln_id)
                .order_by(Comment.created_at.asc())
                .offset((page - 1) * size)
                .limit(size)
            )
        ).all()
        result: list[CommentOut] = []
        for comment, user in rows:
            mention_rows = (
                await self.db.execute(
                    select(CommentMention, User)
                    .join(User, User.id == CommentMention.user_id)
                    .where(CommentMention.comment_id == comment.id)
                )
            ).all()
            mentions = [MentionOut(user_id=u.id, username=u.username) for _, u in mention_rows]
            result.append(
                CommentOut(
                    id=comment.id,
                    vulnerability_id=comment.vulnerability_id,
                    user_id=comment.user_id,
                    username=user.username,
                    avatar_url=user.avatar_url,
                    content=comment.content,
                    mentions=mentions,
                    created_at=comment.created_at,
                    updated_at=comment.updated_at,
                )
            )
        return result, total

    async def create(self, project_id: UUID, vuln_id: UUID, content: str, actor: User) -> CommentOut:
        vuln = await self._ensure_vuln(project_id, vuln_id)
        comment = Comment(vulnerability_id=vuln.id, user_id=actor.id, content=content)
        self.db.add(comment)
        await self.db.flush()
        mentioned_users = await self._extract_mentions(project_id, content)
        mention_models: list[MentionOut] = []
        for user in mentioned_users:
            self.db.add(CommentMention(comment_id=comment.id, user_id=user.id))
            notification = Notification(user_id=user.id, type=NotificationType.MENTION, comment_id=comment.id, is_read=False)
            self.db.add(notification)
            mention_models.append(MentionOut(user_id=user.id, username=user.username))
        await self.db.commit()
        await self.db.refresh(comment)
        for user in mentioned_users:
            await ws_manager.notify_user(
                user.id,
                {
                    "event": "notification",
                    "entity": "notification",
                    "data": {
                        "type": NotificationType.MENTION.value,
                        "is_read": False,
                        "comment_id": str(comment.id),
                        "vulnerability_id": str(vuln.id),
                        "vulnerability_title": vuln.title,
                        "project_id": str(project_id),
                        "commenter_username": actor.username,
                    },
                },
            )
        await self.audit.log("CREATE", user_id=actor.id, entity_type="comment", entity_id=comment.id)
        await ws_manager.broadcast(project_id, {"event": "created", "entity": "comment", "project_id": str(project_id), "data": {"id": str(comment.id)}})
        return CommentOut(
            id=comment.id,
            vulnerability_id=comment.vulnerability_id,
            user_id=comment.user_id,
            username=actor.username,
            avatar_url=actor.avatar_url,
            content=comment.content,
            mentions=mention_models,
            created_at=comment.created_at,
            updated_at=comment.updated_at,
        )

    async def update(self, project_id: UUID, vuln_id: UUID, comment_id: UUID, content: str, actor: User) -> CommentOut:
        await self._ensure_vuln(project_id, vuln_id)
        comment = await self.db.scalar(
            select(Comment).where(and_(Comment.id == comment_id, Comment.vulnerability_id == vuln_id))
        )
        if not comment:
            raise NotFoundError("Комментарий не найден")
        if comment.user_id != actor.id:
            raise ForbiddenError("Можно редактировать только свой комментарий")
        comment.content = content
        await self.db.execute(delete(CommentMention).where(CommentMention.comment_id == comment.id))
        mentioned_users = await self._extract_mentions(project_id, content)
        mention_models: list[MentionOut] = []
        for user in mentioned_users:
            self.db.add(CommentMention(comment_id=comment.id, user_id=user.id))
            mention_models.append(MentionOut(user_id=user.id, username=user.username))
        await self.db.commit()
        await self.db.refresh(comment)
        await self.audit.log("UPDATE", user_id=actor.id, entity_type="comment", entity_id=comment.id)
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "comment", "project_id": str(project_id), "data": {"id": str(comment.id)}})
        return CommentOut(
            id=comment.id,
            vulnerability_id=comment.vulnerability_id,
            user_id=comment.user_id,
            username=actor.username,
            avatar_url=actor.avatar_url,
            content=comment.content,
            mentions=mention_models,
            created_at=comment.created_at,
            updated_at=comment.updated_at,
        )

    async def delete(self, project_id: UUID, vuln_id: UUID, comment_id: UUID, actor: User) -> None:
        await self._ensure_vuln(project_id, vuln_id)
        comment = await self.db.scalar(
            select(Comment).where(and_(Comment.id == comment_id, Comment.vulnerability_id == vuln_id))
        )
        if not comment:
            raise NotFoundError("Комментарий не найден")
        if comment.user_id != actor.id:
            raise ForbiddenError("Можно удалить только свой комментарий")
        await self.db.delete(comment)
        await self.db.commit()
        await self.audit.log("DELETE", user_id=actor.id, entity_type="comment", entity_id=comment.id)
        await ws_manager.broadcast(project_id, {"event": "deleted", "entity": "comment", "project_id": str(project_id), "data": {"id": str(comment.id)}})


class NotificationService:
    """Сервис уведомлений пользователя."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def list(self, user_id: UUID, page: int, size: int, is_read: bool | None) -> tuple[list[NotificationOut], int]:
        query = select(Notification).where(Notification.user_id == user_id)
        if is_read is not None:
            query = query.where(Notification.is_read == is_read)
        total = await self.db.scalar(select(func.count()).select_from(query.subquery())) or 0
        notifications = (
            await self.db.scalars(query.order_by(Notification.created_at.desc()).offset((page - 1) * size).limit(size))
        ).all()
        result: list[NotificationOut] = []
        for item in notifications:
            context = None
            if item.comment_id:
                host_id = await self.db.scalar(
                    select(VulnerabilityAsset.asset_id)
                    .join(Vulnerability, Vulnerability.id == VulnerabilityAsset.vulnerability_id)
                    .where(
                        and_(
                            VulnerabilityAsset.asset_type == AssetType.HOST,
                            Vulnerability.id
                            == select(Comment.vulnerability_id).where(Comment.id == item.comment_id).scalar_subquery(),
                        )
                    )
                    .limit(1)
                )
                row = await self.db.execute(
                    select(Comment, Vulnerability, User)
                    .join(Vulnerability, Vulnerability.id == Comment.vulnerability_id)
                    .join(User, User.id == Comment.user_id)
                    .where(Comment.id == item.comment_id)
                )
                data = row.first()
                if data:
                    comment, vuln, commenter = data
                    context = NotificationContext(
                        vulnerability_id=vuln.id,
                        vulnerability_title=vuln.title,
                        project_id=vuln.project_id,
                        host_id=host_id,
                        commenter_username=commenter.username,
                    )
            result.append(
                NotificationOut(
                    id=item.id,
                    type=item.type.value,
                    comment_id=item.comment_id,
                    is_read=item.is_read,
                    created_at=item.created_at,
                    context=context,
                )
            )
        return result, total

    async def unread_count(self, user_id: UUID) -> int:
        return await self.db.scalar(
            select(func.count()).select_from(Notification).where(and_(Notification.user_id == user_id, Notification.is_read.is_(False)))
        ) or 0

    async def mark_read(self, notification_id: UUID, user_id: UUID) -> Notification:
        notification = await self.db.scalar(
            select(Notification).where(and_(Notification.id == notification_id, Notification.user_id == user_id))
        )
        if not notification:
            raise NotFoundError("Уведомление не найдено")
        notification.is_read = True
        await self.db.commit()
        await self.db.refresh(notification)
        return notification

    async def mark_all_read(self, user_id: UUID) -> None:
        await self.db.execute(update(Notification).where(Notification.user_id == user_id).values(is_read=True))
        await self.db.commit()


class ImportService:
    """Сервис импорта инфраструктуры из JSON-формата PCF."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db

    async def _find_matching_host(self, project_id: UUID, host_data) -> Host | None:
        clauses = []
        if host_data.ip_address:
            clauses.append(Host.ip_address == host_data.ip_address)
        if host_data.hostname:
            clauses.append(Host.hostname == host_data.hostname)
        if not clauses:
            return None
        matches = list(
            (
                await self.db.scalars(
                    select(Host).where(and_(Host.project_id == project_id, or_(*clauses)))
                )
            ).all()
        )
        if not matches:
            return None
        exact_matches = [
            host for host in matches if host.ip_address == host_data.ip_address and host.hostname == host_data.hostname
        ]
        if exact_matches:
            return exact_matches[0]
        if len(matches) == 1:
            return matches[0]
        raise ValidationError(
            f"Найдено несколько host для ip/hostname ({host_data.ip_address or '-'} / {host_data.hostname or '-'})"
        )

    @staticmethod
    def _merge_host_fields(host: Host, host_data) -> None:
        if not host.ip_address and host_data.ip_address:
            host.ip_address = host_data.ip_address
        if not host.hostname and host_data.hostname:
            host.hostname = host_data.hostname
        if host.status.value == "unknown" and host_data.status.value != "unknown":
            host.status = host_data.status
        if not host.notes and host_data.notes:
            host.notes = host_data.notes

    @staticmethod
    def _merge_service_fields(service: Service, service_data) -> None:
        if not service.version and service_data.version:
            service.version = service_data.version
        if not service.banner and service_data.banner:
            service.banner = service_data.banner

    @staticmethod
    def _merge_endpoint_fields(endpoint: Endpoint, endpoint_payload: dict) -> None:
        if not endpoint.description and endpoint_payload.get("description"):
            endpoint.description = endpoint_payload["description"]
        if not endpoint.query_params and endpoint_payload.get("query_params"):
            endpoint.query_params = endpoint_payload["query_params"]
        if not endpoint.request_body and endpoint_payload.get("request_body"):
            endpoint.request_body = endpoint_payload["request_body"]
        if not endpoint.request_content_type and endpoint_payload.get("request_content_type"):
            endpoint.request_content_type = endpoint_payload["request_content_type"]
        if not endpoint.request_headers and endpoint_payload.get("request_headers"):
            endpoint.request_headers = endpoint_payload["request_headers"]

    @staticmethod
    def _decode_text_payload(payload: bytes) -> str:
        try:
            return payload.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise ValidationError("Файл должен быть в кодировке UTF-8") from exc

    @staticmethod
    def _validate_openapi_payload(payload: bytes) -> None:
        if not payload:
            raise ValidationError("Swagger/OpenAPI файл пуст")
        if len(payload) > MAX_OPENAPI_IMPORT_BYTES:
            raise ValidationError("Swagger/OpenAPI файл превышает 2 МБ")

    @staticmethod
    def _parse_relaxed_openapi_scalar(raw_value: str) -> object:
        value = raw_value.strip().rstrip(",").strip()
        if not value:
            return ""
        lowered = value.lower()
        if lowered == "true":
            return True
        if lowered == "false":
            return False
        if re.fullmatch(r"-?\d+", value):
            try:
                return int(value)
            except ValueError:
                return value
        return value

    @classmethod
    def _parse_relaxed_openapi_object(cls, lines: list[str], index: int) -> tuple[dict, int]:
        result: dict = {}
        while index < len(lines):
            line = lines[index].strip()
            if not line:
                index += 1
                continue
            normalized = line[:-1].rstrip() if line.endswith(",") else line
            if normalized == "}":
                return result, index + 1
            if normalized.endswith("{"):
                key = normalized[:-1].strip()
                nested, index = cls._parse_relaxed_openapi_object(lines, index + 1)
                result[key] = nested
                continue
            if normalized.endswith("["):
                key = normalized[:-1].strip()
                nested, index = cls._parse_relaxed_openapi_array(lines, index + 1)
                result[key] = nested
                continue
            parts = normalized.split(None, 1)
            key = parts[0].strip()
            raw_value = parts[1] if len(parts) > 1 else ""
            if raw_value == "{}":
                result[key] = {}
            elif raw_value == "[]":
                result[key] = []
            else:
                result[key] = cls._parse_relaxed_openapi_scalar(raw_value)
            index += 1
        return result, index

    @classmethod
    def _parse_relaxed_openapi_array(cls, lines: list[str], index: int) -> tuple[list[object], int]:
        result: list[object] = []
        while index < len(lines):
            line = lines[index].strip()
            if not line:
                index += 1
                continue
            normalized = line[:-1].rstrip() if line.endswith(",") else line
            if normalized == "]":
                return result, index + 1
            if normalized == "{":
                nested, index = cls._parse_relaxed_openapi_object(lines, index + 1)
                result.append(nested)
                continue
            if normalized == "[":
                nested, index = cls._parse_relaxed_openapi_array(lines, index + 1)
                result.append(nested)
                continue
            result.append(cls._parse_relaxed_openapi_scalar(normalized))
            index += 1
        return result, index

    @staticmethod
    def _normalize_relaxed_openapi_mime(value: str) -> str:
        text = value.strip()
        if "/" in text:
            return text
        for prefix in ("application", "multipart", "text", "image", "audio", "video"):
            if text.startswith(prefix) and len(text) > len(prefix):
                return f"{prefix}/{text[len(prefix):]}"
        return text

    @staticmethod
    def _normalize_relaxed_openapi_ref(value: str) -> str:
        text = value.strip()
        for prefix in ("definitions", "parameters", "responses", "securityDefinitions"):
            compact = f"#{prefix}"
            if text.startswith(compact) and not text.startswith(f"#/{prefix}/"):
                tail = text[len(compact) :].strip("/")
                if tail:
                    return f"#/{prefix}/{tail}"
        return text

    @classmethod
    def _normalize_relaxed_openapi_path_key(cls, raw_path: str, path_item: object) -> str:
        text = str(raw_path or "").strip()
        if not text:
            return "/"
        if text.startswith("/"):
            return text
        first_tag = ""
        if isinstance(path_item, dict):
            for operation in path_item.values():
                if not isinstance(operation, dict):
                    continue
                tags = operation.get("tags")
                if isinstance(tags, list) and tags:
                    candidate = str(tags[0] or "").strip().strip("/")
                    if candidate:
                        first_tag = candidate
                        break
        if first_tag and text.lower().startswith(first_tag.lower()):
            tail = text[len(first_tag) :].lstrip("/")
            text = f"/{first_tag}/{tail}" if tail else f"/{first_tag}"
        else:
            text = f"/{text.lstrip('/')}"
        text = re.sub(r"(?<!/)\{", "/{", text)
        text = re.sub(r"\}(?!/|$)", "}/", text)
        text = re.sub(r"/{2,}", "/", text)
        return text

    @classmethod
    def _normalize_relaxed_openapi_document(cls, value: object, key: str | None = None) -> object:
        if isinstance(value, dict):
            normalized = {item_key: cls._normalize_relaxed_openapi_document(item_value, item_key) for item_key, item_value in value.items()}
            if "basePath" in normalized and isinstance(normalized["basePath"], str):
                base_path = normalized["basePath"].strip()
                normalized["basePath"] = f"/{base_path.lstrip('/')}" if base_path else ""
            if "paths" in normalized and isinstance(normalized["paths"], dict):
                normalized["paths"] = {
                    cls._normalize_relaxed_openapi_path_key(path_key, path_item): path_item
                    for path_key, path_item in normalized["paths"].items()
                }
            return normalized
        if isinstance(value, list):
            return [cls._normalize_relaxed_openapi_document(item, key) for item in value]
        if isinstance(value, str):
            text = value.strip()
            if key == "$ref":
                return cls._normalize_relaxed_openapi_ref(text)
            if key in {"consumes", "produces"}:
                return cls._normalize_relaxed_openapi_mime(text)
            return text
        return value

    @classmethod
    def _load_relaxed_openapi_document(cls, raw_text: str) -> dict:
        lines = raw_text.splitlines()
        index = 0
        while index < len(lines) and not lines[index].strip():
            index += 1
        if index >= len(lines) or lines[index].strip() != "{":
            raise ValidationError("Swagger/OpenAPI файл должен быть валидным JSON или YAML")
        parsed, next_index = cls._parse_relaxed_openapi_object(lines, index + 1)
        trailing = [line.strip() for line in lines[next_index:] if line.strip()]
        if trailing:
            raise ValidationError("Swagger/OpenAPI файл должен быть валидным JSON или YAML")
        normalized = cls._normalize_relaxed_openapi_document(parsed)
        if not isinstance(normalized, dict):
            raise ValidationError("Swagger/OpenAPI документ должен быть объектом")
        return normalized

    @classmethod
    def _load_json_or_yaml_document(cls, raw_text: str) -> dict:
        if not raw_text.strip():
            raise ValidationError("Swagger/OpenAPI файл пуст")
        try:
            parsed = json.loads(raw_text)
        except Exception:
            try:
                parsed = yaml.safe_load(raw_text)
            except Exception as exc:
                try:
                    parsed = cls._load_relaxed_openapi_document(raw_text)
                except Exception:
                    raise ValidationError("Swagger/OpenAPI файл должен быть валидным JSON или YAML") from exc
        if not isinstance(parsed, dict):
            raise ValidationError("Swagger/OpenAPI документ должен быть объектом")
        return parsed

    @staticmethod
    def _resolve_json_pointer(document: dict, ref: str) -> object:
        if not ref.startswith("#/"):
            raise ValidationError("Поддерживаются только локальные $ref вида '#/...'")
        current: object = document
        for raw_token in ref[2:].split("/"):
            token = raw_token.replace("~1", "/").replace("~0", "~")
            if not isinstance(current, dict) or token not in current:
                raise ValidationError(f"Swagger/OpenAPI содержит битую ссылку $ref: {ref}")
            current = current[token]
        return current

    @classmethod
    def _resolve_openapi_ref(cls, document: dict, value: object, depth: int = 0) -> object:
        if depth > 20:
            raise ValidationError("Swagger/OpenAPI содержит слишком глубокую цепочку $ref")
        if isinstance(value, dict) and "$ref" in value:
            ref = value.get("$ref")
            if not isinstance(ref, str):
                raise ValidationError("Некорректный $ref в Swagger/OpenAPI документе")
            resolved = cls._resolve_json_pointer(document, ref)
            if isinstance(resolved, dict):
                merged = dict(resolved)
                merged.update({key: item for key, item in value.items() if key != "$ref"})
                return cls._resolve_openapi_ref(document, merged, depth + 1)
            return cls._resolve_openapi_ref(document, resolved, depth + 1)
        return value

    @staticmethod
    def _normalize_openapi_path(path_value: str, prefix: str) -> str:
        combined = f"{prefix.rstrip('/')}/{path_value.lstrip('/')}" if prefix else path_value
        return AssetService._normalize_endpoint_path(combined) or "/"

    @staticmethod
    def _extract_openapi_host(document: dict) -> str | None:
        swagger_host = str(document.get("host") or "").strip()
        if swagger_host:
            return re.sub(r"^https?://", "", swagger_host, flags=re.IGNORECASE).split("/", 1)[0]
        servers = document.get("servers")
        if not isinstance(servers, list):
            return None
        for server in servers:
            if not isinstance(server, dict):
                continue
            raw_url = str(server.get("url") or "").strip()
            if not raw_url or "://" not in raw_url:
                continue
            parsed = urlsplit(raw_url)
            if parsed.netloc:
                return parsed.netloc
        return None

    @staticmethod
    def _extract_openapi_path_prefix(document: dict) -> str:
        base_path = str(document.get("basePath") or "").strip()
        if base_path:
            return base_path
        servers = document.get("servers")
        if not isinstance(servers, list):
            return ""
        for server in servers:
            if not isinstance(server, dict):
                continue
            raw_url = str(server.get("url") or "").strip()
            if not raw_url:
                continue
            parsed = urlsplit(raw_url if "://" in raw_url else f"https://placeholder.local{raw_url}")
            if parsed.path and parsed.path != "/":
                return parsed.path
        return ""

    @staticmethod
    def _serialize_openapi_example(value: object) -> str | None:
        if value is None:
            return None
        if isinstance(value, bool):
            return "true" if value else "false"
        if isinstance(value, str):
            text = value.strip()
            return text or None
        if isinstance(value, (dict, list)):
            return json.dumps(value, ensure_ascii=False, indent=2)
        return str(value)

    @classmethod
    def _build_openapi_example_from_schema(cls, document: dict, schema: object, depth: int = 0) -> object:
        if schema is None or depth > 8:
            return None
        resolved = cls._resolve_openapi_ref(document, schema)
        if not isinstance(resolved, dict):
            return None
        if "example" in resolved:
            return resolved.get("example")
        if "default" in resolved:
            return resolved.get("default")
        enum_values = resolved.get("enum")
        if isinstance(enum_values, list) and enum_values:
            return enum_values[0]
        schema_type = str(resolved.get("type") or "").lower()
        if schema_type == "object" or isinstance(resolved.get("properties"), dict):
            properties = resolved.get("properties") if isinstance(resolved.get("properties"), dict) else {}
            example: dict = {}
            for prop_name, prop_schema in properties.items():
                example[prop_name] = cls._build_openapi_example_from_schema(document, prop_schema, depth + 1)
            return example
        if schema_type == "array":
            item_example = cls._build_openapi_example_from_schema(document, resolved.get("items"), depth + 1)
            return [item_example] if item_example is not None else []
        if schema_type in {"integer", "number"}:
            return 0
        if schema_type == "boolean":
            return False
        if schema_type == "file":
            return "<binary>"
        if schema_type == "string" or not schema_type:
            fmt = str(resolved.get("format") or "").lower()
            if fmt == "date-time":
                return "2024-01-01T00:00:00Z"
            if fmt == "date":
                return "2024-01-01"
            if fmt == "uuid":
                return "00000000-0000-0000-0000-000000000000"
            if fmt == "email":
                return "user@example.com"
            return "string"
        return None

    @classmethod
    def _extract_openapi_query_params(cls, document: dict, path_item: dict, operation: dict) -> list[dict]:
        collected: list[dict] = []
        seen: set[str] = set()
        for source in (path_item.get("parameters"), operation.get("parameters")):
            if not isinstance(source, list):
                continue
            for item in source:
                resolved = cls._resolve_openapi_ref(document, item)
                if not isinstance(resolved, dict) or str(resolved.get("in") or "").lower() != "query":
                    continue
                name = str(resolved.get("name") or "").strip()
                if not name or name in seen:
                    continue
                seen.add(name)
                example_value: object | None = None
                if "example" in resolved:
                    example_value = resolved.get("example")
                elif "default" in resolved:
                    example_value = resolved.get("default")
                elif isinstance(resolved.get("enum"), list) and resolved["enum"]:
                    example_value = resolved["enum"][0]
                else:
                    schema_for_value = resolved.get("schema") if isinstance(resolved.get("schema"), dict) else resolved
                    built = cls._build_openapi_example_from_schema(document, schema_for_value)
                    if built is not None and not isinstance(built, (dict, list)):
                        example_value = built
                collected.append(
                    {
                        "name": name,
                        "value": cls._serialize_openapi_example(example_value),
                        "required": bool(resolved.get("required")),
                        "description": str(resolved.get("description") or "").strip() or None,
                    }
                )
        return collected

    @classmethod
    def _collect_openapi_form_params(cls, document: dict, path_item: dict, operation: dict) -> list[dict]:
        body_param: dict | None = None
        form_params: list[dict] = []
        for source in (path_item.get("parameters"), operation.get("parameters")):
            if not isinstance(source, list):
                continue
            for item in source:
                resolved = cls._resolve_openapi_ref(document, item)
                if not isinstance(resolved, dict):
                    continue
                in_value = str(resolved.get("in") or "").lower()
                if in_value == "body" and body_param is None:
                    body_param = resolved
                elif in_value == "formdata":
                    form_params.append(resolved)
        return [{"body": body_param, "form": form_params}]

    @classmethod
    def _extract_openapi_request_details(
        cls, document: dict, path_item: dict, operation: dict
    ) -> tuple[str | None, str | None]:
        request_body = operation.get("requestBody")
        if request_body:
            resolved = cls._resolve_openapi_ref(document, request_body)
            if isinstance(resolved, dict):
                content = resolved.get("content")
                if isinstance(content, dict) and content:
                    for content_type, content_schema in content.items():
                        normalized_type = str(content_type or "").strip() or None
                        if not isinstance(content_schema, dict):
                            return normalized_type, None
                        if "example" in content_schema:
                            return normalized_type, cls._serialize_openapi_example(content_schema.get("example"))
                        examples = content_schema.get("examples")
                        if isinstance(examples, dict):
                            for example in examples.values():
                                if isinstance(example, dict) and "value" in example:
                                    return normalized_type, cls._serialize_openapi_example(example.get("value"))
                        schema = content_schema.get("schema")
                        if schema is not None:
                            built = cls._build_openapi_example_from_schema(document, schema)
                            if built is not None:
                                return normalized_type, cls._serialize_openapi_example(built)
                        return normalized_type, None
        swagger_consumes = operation.get("consumes") if isinstance(operation.get("consumes"), list) else document.get("consumes")
        consumes_list = swagger_consumes if isinstance(swagger_consumes, list) else []
        preferred_type = next(
            (str(item).strip() for item in consumes_list if isinstance(item, str) and str(item).strip()),
            None,
        )
        collected = cls._collect_openapi_form_params(document, path_item, operation)
        body_param = collected[0]["body"] if collected else None
        form_params = collected[0]["form"] if collected else []
        if body_param:
            content_type = preferred_type or "application/json"
            if "example" in body_param:
                return content_type, cls._serialize_openapi_example(body_param.get("example"))
            built = cls._build_openapi_example_from_schema(document, body_param.get("schema"))
            if built is None:
                return content_type, None
            return content_type, cls._serialize_openapi_example(built)
        if form_params:
            content_type = preferred_type or "application/x-www-form-urlencoded"
            pairs: list[tuple[str, str]] = []
            for param in form_params:
                name = str(param.get("name") or "").strip()
                if not name:
                    continue
                schema_for_value = param.get("schema") if isinstance(param.get("schema"), dict) else param
                built = cls._build_openapi_example_from_schema(document, schema_for_value)
                value_text = cls._serialize_openapi_example(built) or ""
                pairs.append((name, value_text))
            if not pairs:
                return content_type, None
            if content_type.startswith("multipart"):
                body_text = "\n".join(f"{name}={value}" for name, value in pairs)
            else:
                body_text = "&".join(f"{name}={value}" for name, value in pairs)
            return content_type, body_text or None
        return None, None

    async def import_openapi(self, project_id: UUID, host_id: UUID, payload: bytes, actor_id: UUID) -> OpenApiImportResult:
        host = await self.db.scalar(select(Host).where(and_(Host.id == host_id, Host.project_id == project_id)))
        if not host:
            raise NotFoundError("Хост не найден")

        self._validate_openapi_payload(payload)
        raw_text = self._decode_text_payload(payload)
        document = self._load_json_or_yaml_document(raw_text)
        paths = document.get("paths")
        if not isinstance(paths, dict) or not paths:
            raise ValidationError("В Swagger/OpenAPI документе отсутствует объект paths")
        if len(paths) > MAX_OPENAPI_PATHS:
            raise ValidationError("Swagger/OpenAPI документ содержит слишком много paths для импорта")

        spec_host = self._extract_openapi_host(document)
        result = OpenApiImportResult(
            host_id=host.id,
            spec_host=spec_host,
            endpoints_created=0,
            endpoints_skipped=0,
            errors=[],
        )
        current_host_targets = {value.lower() for value in (host.hostname, host.ip_address) if value}
        if spec_host and current_host_targets and spec_host.lower() not in current_host_targets:
            result.errors.append(
                f"В спецификации указан host '{spec_host}', но импорт выполнен в текущий хост '{host.hostname or host.ip_address}'."
            )

        supported_methods = {"get", "post", "put", "patch", "delete", "head", "options"}
        path_prefix = self._extract_openapi_path_prefix(document)

        try:
            for path_value, raw_path_item in paths.items():
                if not isinstance(path_value, str) or not path_value.strip():
                    continue
                path_item = self._resolve_openapi_ref(document, raw_path_item)
                if not isinstance(path_item, dict):
                    result.errors.append(f"Раздел paths['{path_value}'] имеет некорректную структуру и был пропущен.")
                    continue
                for raw_method_name, raw_operation in path_item.items():
                    if not isinstance(raw_method_name, str):
                        continue
                    method_name = raw_method_name.lower()
                    if method_name in {"parameters"} or method_name.startswith("x-"):
                        continue
                    if method_name not in supported_methods:
                        result.errors.append(f"Метод '{raw_method_name}' для '{path_value}' не поддерживается и был пропущен.")
                        continue
                    operation = self._resolve_openapi_ref(document, raw_operation)
                    if not isinstance(operation, dict):
                        result.errors.append(f"Операция '{raw_method_name}' для '{path_value}' имеет некорректную структуру.")
                        continue
                    if bool(operation.get("deprecated")):
                        result.errors.append(
                            f"Операция '{raw_method_name.upper()} {path_value}' помечена как deprecated и была пропущена."
                        )
                        continue
                    request_content_type, request_body = self._extract_openapi_request_details(
                        document, path_item, operation
                    )
                    endpoint_payload = {
                        "path": self._normalize_openapi_path(path_value, path_prefix),
                        "method": method_name.upper(),
                        "description": "\n\n".join(
                            part for part in [str(operation.get("summary") or "").strip(), str(operation.get("description") or "").strip()] if part
                        )
                        or None,
                        "query_params": self._extract_openapi_query_params(document, path_item, operation),
                        "request_body": request_body,
                        "request_content_type": request_content_type,
                        "request_headers": [],
                    }
                    endpoint = await self.db.scalar(
                        select(Endpoint).where(
                            and_(
                                Endpoint.host_id == host.id,
                                Endpoint.path == endpoint_payload["path"],
                                Endpoint.method == endpoint_payload["method"],
                            )
                        )
                    )
                    if endpoint is None:
                        self.db.add(
                            Endpoint(
                                host_id=host.id,
                                path=endpoint_payload["path"],
                                method=endpoint_payload["method"],
                                description=endpoint_payload["description"],
                                query_params=endpoint_payload["query_params"],
                                request_body=endpoint_payload["request_body"],
                                request_content_type=endpoint_payload["request_content_type"],
                                request_headers=[],
                            )
                        )
                        result.endpoints_created += 1
                    else:
                        self._merge_endpoint_fields(endpoint, endpoint_payload)
                        result.endpoints_skipped += 1

            if result.endpoints_created == 0 and result.endpoints_skipped == 0:
                raise ValidationError("В Swagger/OpenAPI документе не найдено методов для импорта")

            await self.db.commit()
        except Exception:
            await self.db.rollback()
            raise

        await AuditService(self.db).log(
            "CREATE",
            user_id=actor_id,
            entity_type="openapi_import",
            details=result.model_dump(mode="json"),
        )
        await ws_manager.broadcast(
            project_id,
            {"event": "updated", "entity": "endpoint", "project_id": str(project_id), "data": {"host_id": str(host_id), "imported": True}},
        )
        return result

    async def export_openapi(self, project_id: UUID, host_id: UUID) -> dict:
        """Собирает OpenAPI 3.0 документ из эндпоинтов хоста."""
        host = await self.db.scalar(select(Host).where(and_(Host.id == host_id, Host.project_id == project_id)))
        if not host:
            raise NotFoundError("Хост не найден")
        endpoints_result = await self.db.scalars(
            select(Endpoint).where(Endpoint.host_id == host.id).order_by(Endpoint.path, Endpoint.method)
        )
        endpoints = endpoints_result.all()
        title = host.hostname or host.ip_address or "API"
        server_url = None
        if host.hostname:
            server_url = f"https://{host.hostname}"
        elif host.ip_address:
            server_url = f"http://{host.ip_address}"
        document: dict = {
            "openapi": "3.0.0",
            "info": {"title": title, "version": "1.0.0"},
            "paths": {},
        }
        if server_url:
            document["servers"] = [{"url": server_url}]
        for endpoint in endpoints:
            path_value = endpoint.path or "/"
            method_value = (endpoint.method.value if endpoint.method else "GET").lower()
            path_item = document["paths"].setdefault(path_value, {})
            operation: dict = {}
            description = (endpoint.description or "").strip()
            if description:
                first_line, _, rest = description.partition("\n\n")
                operation["summary"] = first_line.strip()
                if rest.strip():
                    operation["description"] = rest.strip()
            parameters: list[dict] = []
            for raw_param in endpoint.query_params or []:
                if not isinstance(raw_param, dict):
                    continue
                name = str(raw_param.get("name") or "").strip()
                if not name:
                    continue
                parameter: dict = {
                    "name": name,
                    "in": "query",
                    "required": bool(raw_param.get("required")),
                    "schema": {"type": "string"},
                }
                description_value = str(raw_param.get("description") or "").strip()
                if description_value:
                    parameter["description"] = description_value
                example_value = raw_param.get("value")
                if example_value not in (None, ""):
                    parameter["example"] = example_value
                parameters.append(parameter)
            for raw_header in endpoint.request_headers or []:
                if not isinstance(raw_header, dict):
                    continue
                name = str(raw_header.get("name") or "").strip()
                if not name:
                    continue
                parameter = {
                    "name": name,
                    "in": "header",
                    "required": False,
                    "schema": {"type": "string"},
                }
                example_value = raw_header.get("value")
                if example_value not in (None, ""):
                    parameter["example"] = example_value
                parameters.append(parameter)
            if parameters:
                operation["parameters"] = parameters
            request_body_text = (endpoint.request_body or "").strip()
            if request_body_text:
                content_type = (endpoint.request_content_type or "application/json").strip() or "application/json"
                example_payload: object = request_body_text
                if content_type.endswith("json") or "json" in content_type.lower():
                    try:
                        example_payload = json.loads(request_body_text)
                    except (ValueError, TypeError):
                        example_payload = request_body_text
                operation["requestBody"] = {
                    "required": True,
                    "content": {content_type: {"example": example_payload}},
                }
            operation["responses"] = {"200": {"description": "OK"}}
            path_item[method_value] = operation
        return document

    async def import_json(self, project_id: UUID, payload: bytes, actor_id: UUID) -> ImportResult:
        """Импортирует данные атомарно: при ошибке откатывает всё."""
        try:
            parsed = json.loads(payload.decode("utf-8"))
        except Exception as exc:
            raise ValidationError("Невалидный JSON-файл") from exc
        try:
            import_payload = PcfImportPayload.model_validate(parsed)
        except Exception as exc:
            raise ValidationError(f"JSON импорта не соответствует схеме PCF: {exc}") from exc

        result = ImportResult(hosts_created=0, ports_created=0, services_created=0, endpoints_created=0, errors=[])
        try:
            for host_data in import_payload.hosts:
                host = await self._find_matching_host(project_id, host_data)
                if host is None:
                    host = Host(
                        project_id=project_id,
                        ip_address=host_data.ip_address,
                        hostname=host_data.hostname,
                        status=host_data.status,
                        notes=host_data.notes,
                    )
                    self.db.add(host)
                    await self.db.flush()
                    result.hosts_created += 1
                else:
                    self._merge_host_fields(host, host_data)

                for port_data in host_data.ports:
                    port = await self.db.scalar(
                        select(Port).where(
                            and_(
                                Port.host_id == host.id,
                                Port.port_number == port_data.port_number,
                                Port.protocol == port_data.protocol,
                            )
                        )
                    )
                    if port is None:
                        port = Port(
                            host_id=host.id,
                            port_number=port_data.port_number,
                            protocol=port_data.protocol,
                            state=port_data.state,
                        )
                        self.db.add(port)
                        await self.db.flush()
                        result.ports_created += 1

                    for service_data in port_data.services:
                        service = await self.db.scalar(
                            select(Service).where(
                                and_(
                                    Service.port_id == port.id,
                                    Service.name == service_data.name,
                                )
                            )
                        )
                        if service is None:
                            self.db.add(
                                Service(
                                    port_id=port.id,
                                    name=service_data.name,
                                    version=service_data.version,
                                    banner=service_data.banner,
                                )
                            )
                            result.services_created += 1
                        else:
                            self._merge_service_fields(service, service_data)

                for endpoint_data in host_data.endpoints:
                    endpoint_payload = {
                        "path": endpoint_data.path,
                        "method": endpoint_data.method,
                        "description": endpoint_data.description,
                        "request_raw": endpoint_data.request_raw,
                        "query_params": [item.model_dump() for item in endpoint_data.query_params],
                        "request_body": endpoint_data.request_body,
                        "request_content_type": endpoint_data.request_content_type,
                        "request_headers": [item.model_dump() for item in endpoint_data.request_headers],
                    }
                    endpoint_payload = AssetService._apply_raw_request_payload(endpoint_payload)
                    endpoint_payload = AssetService._apply_structured_request_payload(endpoint_payload)
                    if not endpoint_payload.get("path"):
                        raise ValidationError("Каждый endpoint должен содержать path или request_raw")
                    if endpoint_payload.get("request_headers") is None:
                        endpoint_payload["request_headers"] = []
                    endpoint = await self.db.scalar(
                        select(Endpoint).where(
                            and_(
                                Endpoint.host_id == host.id,
                                Endpoint.path == endpoint_payload["path"],
                                Endpoint.method == endpoint_payload.get("method"),
                            )
                        )
                    )
                    if endpoint is None:
                        self.db.add(
                            Endpoint(
                                host_id=host.id,
                                path=endpoint_payload["path"],
                                method=endpoint_payload.get("method"),
                                description=endpoint_payload.get("description"),
                                query_params=endpoint_payload.get("query_params") or [],
                                request_body=endpoint_payload.get("request_body"),
                                request_content_type=endpoint_payload.get("request_content_type"),
                                request_headers=endpoint_payload.get("request_headers") or [],
                            )
                        )
                        result.endpoints_created += 1
                    else:
                        self._merge_endpoint_fields(endpoint, endpoint_payload)

            await self.db.commit()
        except Exception:
            await self.db.rollback()
            raise

        await AuditService(self.db).log("CREATE", user_id=actor_id, entity_type="import", details=result.model_dump())
        await ws_manager.broadcast(project_id, {"event": "updated", "entity": "host", "project_id": str(project_id), "data": {"imported": True}})
        return result


class ReportService:
    """Сервис генерации отчётов в форматах md/pdf/docx."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.storage = MinioStorage()

    async def _collect_project_data(self, project_id: UUID) -> dict:
        project = await self.db.scalar(select(Project).where(Project.id == project_id))
        if not project:
            raise NotFoundError("Проект не найден")
        members = (
            await self.db.execute(
                select(User.username)
                .join(ProjectMember, ProjectMember.user_id == User.id)
                .where(ProjectMember.project_id == project_id)
            )
        ).all()
        vulnerabilities = (await self.db.scalars(select(Vulnerability).where(Vulnerability.project_id == project_id))).all()
        hosts = (await self.db.scalars(select(Host).where(Host.project_id == project_id))).all()
        ports = (
            await self.db.scalars(select(Port).join(Host, Port.host_id == Host.id).where(Host.project_id == project_id))
        ).all()
        vulnerability_ids = [vulnerability.id for vulnerability in vulnerabilities]
        host_ids = [host.id for host in hosts]
        vulnerability_assets = []
        files = []
        if vulnerability_ids:
            vulnerability_assets = list(
                (
                    await self.db.scalars(
                        select(VulnerabilityAsset).where(VulnerabilityAsset.vulnerability_id.in_(vulnerability_ids))
                    )
                ).all()
            )
            files = list((await self.db.scalars(select(File).where(File.vulnerability_id.in_(vulnerability_ids)))).all())
        return {
            "project": project,
            "members": [m[0] for m in members],
            "vulnerabilities": list(vulnerabilities),
            "hosts": list(hosts),
            "ports": list(ports),
            "host_ids": host_ids,
            "vulnerability_assets": vulnerability_assets,
            "files": files,
        }

    @staticmethod
    def _build_report_indexes(data: dict) -> dict:
        hosts: list[Host] = data["hosts"]
        ports: list[Port] = data["ports"]
        vulnerability_assets: list[VulnerabilityAsset] = data["vulnerability_assets"]
        files: list[File] = data["files"]
        host_by_id = {host.id: host for host in hosts}
        ports_by_host_id: dict[UUID, list[Port]] = {}
        for port in ports:
            ports_by_host_id.setdefault(port.host_id, []).append(port)
        assets_by_vuln_id: dict[UUID, list[VulnerabilityAsset]] = {}
        for asset in vulnerability_assets:
            assets_by_vuln_id.setdefault(asset.vulnerability_id, []).append(asset)
        files_by_vuln_id: dict[UUID, list[File]] = {}
        for file_meta in files:
            files_by_vuln_id.setdefault(file_meta.vulnerability_id, []).append(file_meta)
        files_by_id = {file_meta.id: file_meta for file_meta in files}
        severity_stats: dict[str, int] = {}
        status_stats: dict[str, int] = {}
        for vuln in data["vulnerabilities"]:
            severity_stats[vuln.severity.value] = severity_stats.get(vuln.severity.value, 0) + 1
            status_stats[vuln.status.value] = status_stats.get(vuln.status.value, 0) + 1
        return {
            "host_by_id": host_by_id,
            "ports_by_host_id": ports_by_host_id,
            "assets_by_vuln_id": assets_by_vuln_id,
            "files_by_vuln_id": files_by_vuln_id,
            "files_by_id": files_by_id,
            "severity_stats": severity_stats,
            "status_stats": status_stats,
        }

    @staticmethod
    def _report_text(value: object | None, fallback: str = "-") -> str:
        text = str(value).strip() if value is not None else ""
        return text or fallback

    @staticmethod
    def _report_paragraph(value: object | None, fallback: str = "-") -> str:
        return escape(ReportService._report_text(value, fallback)).replace("\n", "<br/>")

    @staticmethod
    def _host_label(host: Host | None, fallback: object) -> str:
        if not host:
            return str(fallback)
        return host.hostname or host.ip_address or str(host.id)

    @staticmethod
    def _ports_summary(host_ports: list[Port]) -> str:
        return ", ".join(
            f"{port.port_number}/{port.protocol.value if hasattr(port.protocol, 'value') else port.protocol}" for port in host_ports
        ) or "-"

    @staticmethod
    def _linked_assets_for_vuln(vuln: Vulnerability, indexes: dict) -> list[str]:
        linked_assets: list[str] = []
        for asset in indexes["assets_by_vuln_id"].get(vuln.id, []):
            if asset.asset_type == AssetType.HOST:
                linked_assets.append(ReportService._host_label(indexes["host_by_id"].get(asset.asset_id), f"host:{asset.asset_id}"))
            else:
                linked_assets.append(f"{asset.asset_type.value}:{asset.asset_id}")
        return linked_assets

    @staticmethod
    def _resolve_step_files(step: dict, files_by_id: dict[UUID, File]) -> list[File]:
        resolved: list[File] = []
        for raw_file_id in step.get("image_file_ids", []):
            try:
                file_id = UUID(str(raw_file_id))
            except (TypeError, ValueError):
                continue
            file_meta = files_by_id.get(file_id)
            if file_meta:
                resolved.append(file_meta)
        return resolved

    @staticmethod
    def _is_image_file(file_meta: File) -> bool:
        return file_meta.content_type.startswith("image/")

    @staticmethod
    def _normalize_report_image_bytes(image_bytes: bytes) -> bytes | None:
        try:
            with PillowImage.open(BytesIO(image_bytes)) as image:
                normalized = ImageOps.exif_transpose(image)
                output = BytesIO()
                save_image = normalized.convert("RGBA") if normalized.mode in {"P", "LA"} else normalized
                save_format = "PNG" if "A" in save_image.getbands() else "JPEG"
                if save_format == "JPEG":
                    save_image = save_image.convert("RGB")
                    save_image.save(output, format=save_format, quality=90)
                else:
                    save_image.save(output, format=save_format)
                return output.getvalue()
        except (UnidentifiedImageError, OSError, ValueError):
            return None

    async def _download_report_images(self, files: list[File]) -> dict[UUID, bytes]:
        image_bytes: dict[UUID, bytes] = {}
        for file_meta in files:
            if not self._is_image_file(file_meta):
                continue
            try:
                raw_bytes = await asyncio.to_thread(self.storage.download_bytes, file_meta.minio_key)
            except Exception:
                continue
            normalized_bytes = await asyncio.to_thread(self._normalize_report_image_bytes, raw_bytes)
            if normalized_bytes:
                image_bytes[file_meta.id] = normalized_bytes
        return image_bytes

    def _render_markdown_from_data(self, data: dict) -> bytes:
        project: Project = data["project"]
        vulnerabilities: list[Vulnerability] = data["vulnerabilities"]
        hosts: list[Host] = data["hosts"]
        ports: list[Port] = data["ports"]
        indexes = self._build_report_indexes(data)

        lines = [
            f"# Отчёт по проекту {project.name}",
            "",
            "## Общая информация",
            f"- Статус: {project.status.value}",
            f"- Даты: {project.start_date} - {project.end_date}",
            f"- Участники: {', '.join(data['members']) if data['members'] else 'нет'}",
            "",
            "## Статистика уязвимостей по критичности",
        ]
        for key, value in indexes["severity_stats"].items():
            lines.append(f"- {key}: {value}")
        lines += ["", "## Статистика уязвимостей по статусу"]
        for key, value in indexes["status_stats"].items():
            lines.append(f"- {key}: {value}")
        lines += ["", "## Активы", f"- Хосты: {len(hosts)}", f"- Порты: {len(ports)}", ""]
        if hosts:
            lines += ["### Сводка по хостам", "", "| Хост | IP/Hostname | Порты |", "| --- | --- | --- |"]
            for host in hosts:
                lines.append(
                    f"| {self._host_label(host, host.id)} | {host.ip_address or host.hostname or '-'} | {self._ports_summary(indexes['ports_by_host_id'].get(host.id, []))} |"
                )
        lines += ["", "## Уязвимости"]
        for vuln in vulnerabilities:
            VulnerabilityService._hydrate_workflow_steps(vuln)
            attachment_names = [file_meta.original_name for file_meta in indexes["files_by_vuln_id"].get(vuln.id, [])]
            workflow_lines: list[str] = []
            for index, step in enumerate(vuln.workflow_steps or [], start=1):
                title = self._report_text(step.get("title"), f"Этап {index}")
                description = self._report_text(step.get("description"), "")
                workflow_lines.append(f"{index}. {title}")
                if description:
                    workflow_lines.append(f"   {description}")
                image_names = [file_meta.original_name for file_meta in self._resolve_step_files(step, indexes["files_by_id"])]
                if image_names:
                    workflow_lines.append(f"   Скриншоты: {', '.join(image_names)}")
            workflow_text = "\n".join(workflow_lines) if workflow_lines else (vuln.steps_to_reproduce or "-")
            lines += [
                f"### {vuln.title}",
                f"- Severity: {vuln.severity.value}",
                f"- Status: {vuln.status.value}",
                f"- CVSS: {vuln.cvss_version.value if vuln.cvss_version else '-'} {vuln.cvss_score if vuln.cvss_score else '-'}",
                f"- CVSS Vector: {vuln.cvss_vector or '-'}",
                f"- CWE: {vuln.cwe_id or '-'}",
                f"- Активы: {', '.join(self._linked_assets_for_vuln(vuln, indexes)) or '-'}",
                f"- Вложения: {', '.join(attachment_names) if attachment_names else '-'}",
                "",
                "**Описание**",
                vuln.description or "-",
                "",
                "**Шаги воспроизведения**",
                workflow_text,
                "",
                "**Влияние**",
                vuln.impact or "-",
                "",
                "**Рекомендации**",
                vuln.recommendations or "-",
                "",
            ]
        return "\n".join(lines).encode("utf-8")

    @staticmethod
    def _build_pdf(data: dict, indexes: dict, image_bytes_by_id: dict[UUID, bytes]) -> bytes:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        story: list = []
        project: Project = data["project"]
        vulnerabilities: list[Vulnerability] = data["vulnerabilities"]
        hosts: list[Host] = data["hosts"]
        project_dates = f"{project.start_date} - {project.end_date}"

        story.append(Paragraph(ReportService._report_paragraph(f"Отчёт по проекту {project.name}"), styles["Title"]))
        story.append(Spacer(1, 12))
        story.append(Paragraph("Общая информация", styles["Heading2"]))
        story.append(Paragraph(f"<b>Статус:</b> {ReportService._report_paragraph(project.status.value)}", styles["BodyText"]))
        story.append(Paragraph(f"<b>Даты:</b> {ReportService._report_paragraph(project_dates)}", styles["BodyText"]))
        story.append(
            Paragraph(
                f"<b>Участники:</b> {ReportService._report_paragraph(', '.join(data['members']) if data['members'] else 'нет')}",
                styles["BodyText"],
            )
        )
        story.append(Spacer(1, 12))
        story.append(Paragraph("Статистика уязвимостей по критичности", styles["Heading2"]))
        for key, value in indexes["severity_stats"].items():
            story.append(Paragraph(f"• {escape(key)}: {value}", styles["BodyText"]))
        story.append(Spacer(1, 8))
        story.append(Paragraph("Статистика уязвимостей по статусу", styles["Heading2"]))
        for key, value in indexes["status_stats"].items():
            story.append(Paragraph(f"• {escape(key)}: {value}", styles["BodyText"]))
        story.append(Spacer(1, 12))
        if hosts:
            story.append(Paragraph("Сводка по хостам", styles["Heading2"]))
            table_data = [["Хост", "IP/Hostname", "Порты"]]
            for host in hosts:
                table_data.append(
                    [
                        ReportService._host_label(host, host.id),
                        host.ip_address or host.hostname or "-",
                        ReportService._ports_summary(indexes["ports_by_host_id"].get(host.id, [])),
                    ]
                )
            table = Table(table_data, colWidths=[150, 160, 170])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ]
                )
            )
            story.extend([table, Spacer(1, 12)])

        story.append(Paragraph("Уязвимости", styles["Heading2"]))
        for vuln in vulnerabilities:
            VulnerabilityService._hydrate_workflow_steps(vuln)
            cvss_text = f"{vuln.cvss_version.value if vuln.cvss_version else '-'} {vuln.cvss_score if vuln.cvss_score else '-'}"
            story.append(Paragraph(ReportService._report_paragraph(vuln.title), styles["Heading3"]))
            story.append(Paragraph(f"<b>Severity:</b> {escape(vuln.severity.value)}", styles["BodyText"]))
            story.append(Paragraph(f"<b>Status:</b> {escape(vuln.status.value)}", styles["BodyText"]))
            story.append(
                Paragraph(
                    f"<b>CVSS:</b> {ReportService._report_paragraph(cvss_text)}",
                    styles["BodyText"],
                )
            )
            story.append(Paragraph(f"<b>CWE:</b> {ReportService._report_paragraph(vuln.cwe_id)}", styles["BodyText"]))
            story.append(
                Paragraph(
                    f"<b>Активы:</b> {ReportService._report_paragraph(', '.join(ReportService._linked_assets_for_vuln(vuln, indexes)) or '-')}",
                    styles["BodyText"],
                )
            )
            story.append(Paragraph(f"<b>Описание:</b> {ReportService._report_paragraph(vuln.description)}", styles["BodyText"]))
            story.append(Paragraph(f"<b>Влияние:</b> {ReportService._report_paragraph(vuln.impact)}", styles["BodyText"]))
            story.append(Paragraph(f"<b>Рекомендации:</b> {ReportService._report_paragraph(vuln.recommendations)}", styles["BodyText"]))
            story.append(Paragraph("<b>Шаги воспроизведения:</b>", styles["BodyText"]))
            rendered_image_ids: set[UUID] = set()
            if vuln.workflow_steps:
                for index, step in enumerate(vuln.workflow_steps, start=1):
                    story.append(Paragraph(f"{index}. {ReportService._report_paragraph(step.get('title'), f'Этап {index}')}", styles["BodyText"]))
                    if step.get("description"):
                        story.append(Paragraph(ReportService._report_paragraph(step.get("description")), styles["BodyText"]))
                    for file_meta in ReportService._resolve_step_files(step, indexes["files_by_id"]):
                        story.append(Paragraph(f"Скриншот: {ReportService._report_paragraph(file_meta.original_name)}", styles["BodyText"]))
                        image_bytes = image_bytes_by_id.get(file_meta.id)
                        if image_bytes:
                            pdf_image = PlatypusImage(BytesIO(image_bytes))
                            pdf_image._restrictSize(5.5 * inch, 4.5 * inch)
                            story.append(pdf_image)
                            rendered_image_ids.add(file_meta.id)
                        story.append(Spacer(1, 6))
            else:
                story.append(Paragraph(ReportService._report_paragraph(vuln.steps_to_reproduce), styles["BodyText"]))
            additional_images = [
                file_meta
                for file_meta in indexes["files_by_vuln_id"].get(vuln.id, [])
                if ReportService._is_image_file(file_meta) and file_meta.id not in rendered_image_ids
            ]
            if additional_images:
                story.append(Paragraph("<b>Дополнительные скриншоты:</b>", styles["BodyText"]))
                for file_meta in additional_images:
                    story.append(Paragraph(ReportService._report_paragraph(file_meta.original_name), styles["BodyText"]))
                    image_bytes = image_bytes_by_id.get(file_meta.id)
                    if image_bytes:
                        pdf_image = PlatypusImage(BytesIO(image_bytes))
                        pdf_image._restrictSize(5.5 * inch, 4.5 * inch)
                        story.append(pdf_image)
                    story.append(Spacer(1, 6))
            attachment_names = [file_meta.original_name for file_meta in indexes["files_by_vuln_id"].get(vuln.id, []) if not ReportService._is_image_file(file_meta)]
            story.append(Paragraph(f"<b>Прочие вложения:</b> {ReportService._report_paragraph(', '.join(attachment_names) if attachment_names else '-')}", styles["BodyText"]))
            story.append(Spacer(1, 12))
        doc.build(story)
        return buffer.getvalue()

    @staticmethod
    def _build_docx(data: dict, indexes: dict, image_bytes_by_id: dict[UUID, bytes]) -> bytes:
        doc = Document()
        project: Project = data["project"]
        vulnerabilities: list[Vulnerability] = data["vulnerabilities"]
        hosts: list[Host] = data["hosts"]
        doc.add_heading(f"Отчёт по проекту {project.name}", level=0)
        doc.add_heading("Общая информация", level=1)
        doc.add_paragraph(f"Статус: {project.status.value}")
        doc.add_paragraph(f"Даты: {project.start_date} - {project.end_date}")
        doc.add_paragraph(f"Участники: {', '.join(data['members']) if data['members'] else 'нет'}")
        doc.add_heading("Статистика уязвимостей по критичности", level=1)
        for key, value in indexes["severity_stats"].items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")
        doc.add_heading("Статистика уязвимостей по статусу", level=1)
        for key, value in indexes["status_stats"].items():
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")
        if hosts:
            doc.add_heading("Сводка по хостам", level=1)
            table = doc.add_table(rows=1, cols=3)
            table.rows[0].cells[0].text = "Хост"
            table.rows[0].cells[1].text = "IP/Hostname"
            table.rows[0].cells[2].text = "Порты"
            for host in hosts:
                row = table.add_row().cells
                row[0].text = ReportService._host_label(host, host.id)
                row[1].text = host.ip_address or host.hostname or "-"
                row[2].text = ReportService._ports_summary(indexes["ports_by_host_id"].get(host.id, []))
        doc.add_heading("Уязвимости", level=1)
        for vuln in vulnerabilities:
            VulnerabilityService._hydrate_workflow_steps(vuln)
            doc.add_heading(vuln.title, level=2)
            doc.add_paragraph(f"Severity: {vuln.severity.value}")
            doc.add_paragraph(f"Status: {vuln.status.value}")
            doc.add_paragraph(f"CVSS: {vuln.cvss_version.value if vuln.cvss_version else '-'} {vuln.cvss_score if vuln.cvss_score else '-'}")
            doc.add_paragraph(f"CWE: {vuln.cwe_id or '-'}")
            doc.add_paragraph(f"Активы: {', '.join(ReportService._linked_assets_for_vuln(vuln, indexes)) or '-'}")
            doc.add_paragraph(f"Описание: {ReportService._report_text(vuln.description)}")
            doc.add_paragraph(f"Влияние: {ReportService._report_text(vuln.impact)}")
            doc.add_paragraph(f"Рекомендации: {ReportService._report_text(vuln.recommendations)}")
            doc.add_paragraph("Шаги воспроизведения:")
            rendered_image_ids: set[UUID] = set()
            if vuln.workflow_steps:
                for index, step in enumerate(vuln.workflow_steps, start=1):
                    doc.add_paragraph(f"{index}. {ReportService._report_text(step.get('title'), f'Этап {index}')}", style="List Number")
                    if step.get("description"):
                        doc.add_paragraph(ReportService._report_text(step.get("description")))
                    for file_meta in ReportService._resolve_step_files(step, indexes["files_by_id"]):
                        doc.add_paragraph(f"Скриншот: {file_meta.original_name}")
                        image_bytes = image_bytes_by_id.get(file_meta.id)
                        if image_bytes:
                            doc.add_picture(BytesIO(image_bytes), width=Inches(5.5))
                            rendered_image_ids.add(file_meta.id)
            else:
                doc.add_paragraph(ReportService._report_text(vuln.steps_to_reproduce))
            additional_images = [
                file_meta
                for file_meta in indexes["files_by_vuln_id"].get(vuln.id, [])
                if ReportService._is_image_file(file_meta) and file_meta.id not in rendered_image_ids
            ]
            if additional_images:
                doc.add_paragraph("Дополнительные скриншоты:")
                for file_meta in additional_images:
                    doc.add_paragraph(file_meta.original_name)
                    image_bytes = image_bytes_by_id.get(file_meta.id)
                    if image_bytes:
                        doc.add_picture(BytesIO(image_bytes), width=Inches(5.5))
            attachments = [file_meta.original_name for file_meta in indexes["files_by_vuln_id"].get(vuln.id, []) if not ReportService._is_image_file(file_meta)]
            doc.add_paragraph(f"Прочие вложения: {', '.join(attachments) if attachments else '-'}")
        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    async def generate_markdown(self, project_id: UUID) -> bytes:
        data = await self._collect_project_data(project_id)
        return self._render_markdown_from_data(data)

    async def generate(self, project_id: UUID, output_format: str) -> bytes:
        """Генерирует отчёт в запрошенном формате."""
        data = await self._collect_project_data(project_id)
        if output_format == "md":
            return self._render_markdown_from_data(data)
        indexes = self._build_report_indexes(data)
        image_bytes_by_id = await self._download_report_images(data["files"])
        if output_format == "pdf":
            return await asyncio.to_thread(self._build_pdf, data, indexes, image_bytes_by_id)
        if output_format == "docx":
            return await asyncio.to_thread(self._build_docx, data, indexes, image_bytes_by_id)
        raise ValidationError("Неподдерживаемый формат отчёта")
