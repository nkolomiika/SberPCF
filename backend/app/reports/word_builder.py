"""Сборка Word-отчётов СЗИ и ПП по предзагруженным шаблонам.

Шаблоны (`backend/app/reports/templates/*.docx`) содержат фиксированный текст
и пример карточки уязвимости/слабости. Сборщик заполняет известные
плейсхолдеры данными проекта, клонирует пример карточки под каждую
уязвимость/слабость и оставляет всё, что неизвестно, как в шаблоне.
"""
from __future__ import annotations

import copy
import re
from datetime import date as _date
from io import BytesIO
from pathlib import Path
from typing import Iterable, Literal
from uuid import UUID

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from app.enums import AssetType, Severity
from app.models import File, Host, Project, Vulnerability, VulnerabilityAsset

ReportKind = Literal["szi", "pp"]

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates"
TEMPLATE_FILES: dict[str, str] = {
    "szi": "szi_template.docx",
    "pp": "pp_template.docx",
}

SEVERITY_RU: dict[Severity, str] = {
    Severity.CRITICAL: "Критический",
    Severity.HIGH: "Высокий",
    Severity.MEDIUM: "Средний",
    Severity.LOW: "Низкий",
    Severity.INFO: "Слабость/недостаток",
}

# Цвета шрифта для уровня риска и значения CVSS в карточках уязвимостей.
# Бордовый — критический, красный — высокий, жёлтый — средний, зелёный — низкий.
# Слабости/недостатки оставляем без явного цвета (наследуют стиль шаблона).
SEVERITY_COLORS: dict[Severity, RGBColor] = {
    Severity.CRITICAL: RGBColor(0x80, 0x00, 0x00),
    Severity.HIGH: RGBColor(0xC0, 0x00, 0x00),
    Severity.MEDIUM: RGBColor(0xE6, 0xA8, 0x00),
    Severity.LOW: RGBColor(0x00, 0xB0, 0x50),
}

# Шрифт и размеры для всех заполняемых данных (СЗИ-отчёт).
REPORT_FONT_NAME = "SB Sans Display"
REPORT_FONT_SIZE_PT = 11.0
VULN_TABLE_FONT_SIZE_PT = 10.0

# Принудительные размеры шрифта для заголовков H1-H4. Уровни выше 4
# берут размер последнего объявленного уровня (H4) — это редкий случай и
# нужен только как разумный fallback.
HEADING_FONT_SIZE_PT: dict[int, float] = {
    1: 14.0,
    2: 14.0,
    3: 14.0,
    4: 13.5,
}


def _heading_size_pt_for_level(level: int | None) -> float | None:
    """Возвращает требуемый размер шрифта (pt) для заголовка указанного уровня."""
    if level is None:
        return None
    if level in HEADING_FONT_SIZE_PT:
        return HEADING_FONT_SIZE_PT[level]
    if level > max(HEADING_FONT_SIZE_PT):
        return HEADING_FONT_SIZE_PT[max(HEADING_FONT_SIZE_PT)]
    return None

WEAKNESS_SEVERITIES: set[Severity] = {Severity.INFO}

# Маппинг кодов CVSS 4.0 в человекочитаемые значения.
CVSS4_METRIC_LABELS: dict[str, dict[str, str]] = {
    "AV": {"N": "Сетевой (N)", "A": "Смежный (A)", "L": "Локальный (L)", "P": "Физический (P)"},
    "AC": {"H": "Высокая (H)", "L": "Низкая (L)"},
    "AT": {"N": "Отсутствуют (N)", "P": "Существуют (P)"},
    "PR": {"H": "Высокий (H)", "L": "Низкий (L)", "N": "Не требуется (N)"},
    "UI": {"N": "Не требуется (N)", "P": "Пассивное (P)", "A": "Активное (A)"},
    "VC": {"H": "Высокое (H)", "L": "Низкое (L)", "N": "Не оказывает (N)"},
    "VI": {"H": "Высокое (H)", "L": "Низкое (L)", "N": "Не оказывает (N)"},
    "VA": {"H": "Высокое (H)", "L": "Низкое (L)", "N": "Не оказывает (N)"},
    "SC": {"H": "Высокое (H)", "L": "Низкое (L)", "N": "Не оказывает (N)"},
    "SI": {"H": "Высокое (H)", "L": "Низкое (L)", "N": "Не оказывает (N)"},
    "SA": {"H": "Высокое (H)", "L": "Низкое (L)", "N": "Не оказывает (N)"},
}

# Заголовки разделов внутри карточки уязвимости/слабости (после метрик-таблицы).
SECTION_HEADER_FUNC = "Описание уязвимого функционала"
SECTION_HEADER_STEPS = "Шаги для воспроизведения"
SECTION_HEADER_IMPACT = "Возможные последствия эксплуатации"
SECTION_HEADER_RECOMMENDATIONS = "Рекомендации по устранению"
TOP_HEADERS = (
    "Детальное описание уязвимости",
    "Детальное описание недостатка",
)
SECTION_HEADERS_ALL = (
    *TOP_HEADERS,
    SECTION_HEADER_FUNC,
    SECTION_HEADER_STEPS,
    SECTION_HEADER_IMPACT,
    SECTION_HEADER_RECOMMENDATIONS,
)

# Заголовок столбца с количеством в таблице "Выявленные уязвимости и недостатки".
SUMMARY_TABLE_HEADER = "Уровень риска уязвимости"
# Соответствие подписей строк сводной таблицы Severity-уровням.
SUMMARY_ROW_LABEL_TO_SEVERITY: dict[str, Severity] = {
    "критический": Severity.CRITICAL,
    "высокий": Severity.HIGH,
    "средний": Severity.MEDIUM,
    "низкий": Severity.LOW,
    "слабость/недостаток": Severity.INFO,
    "слабость": Severity.INFO,
}


def load_template(kind: ReportKind) -> DocxDocument:
    """Загружает оригинальный шаблон Word из bundled-файлов."""
    path = TEMPLATE_DIR / TEMPLATE_FILES[kind]
    return Document(str(path))


def parse_cvss_vector(vector: str | None) -> dict[str, str]:
    """Разбирает CVSS-вектор и возвращает читабельные значения по каждой метрике."""
    if not vector:
        return {}
    out: dict[str, str] = {}
    for part in vector.strip().split("/"):
        if ":" not in part:
            continue
        key, _, value = part.partition(":")
        key = key.strip().upper()
        value = value.strip().upper()
        if key in CVSS4_METRIC_LABELS:
            out[key] = CVSS4_METRIC_LABELS[key].get(value, value)
    return out


def severity_label(value: Severity | str | None) -> str:
    """Возвращает русскоязычное обозначение уровня риска."""
    if value is None:
        return ""
    if isinstance(value, str):
        try:
            value = Severity(value)
        except ValueError:
            return ""
    return SEVERITY_RU.get(value, "")


def _format_date(value: object) -> str:
    if value is None:
        return ""
    try:
        return value.strftime("%d.%m.%Y")  # type: ignore[union-attr]
    except AttributeError:
        return str(value)


def _clear_paragraph_inline(paragraph: Paragraph) -> object | None:
    """Удаляет все inline-элементы параграфа (runs, hyperlinks, …), сохраняет только pPr.

    Возвращает rPr первого исходного run для повторного использования.
    """
    saved_rpr = None
    first_run = None
    for child in list(paragraph._element):
        tag = child.tag.split("}")[-1]
        if tag == "pPr":
            continue
        if tag == "r" and first_run is None:
            first_run = child
        paragraph._element.remove(child)
    if first_run is not None:
        rpr = first_run.find(qn("w:rPr"))
        if rpr is not None:
            saved_rpr = copy.deepcopy(rpr)
    return saved_rpr


def _apply_run_font(
    run: Run,
    *,
    font_size_pt: float | None = REPORT_FONT_SIZE_PT,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    """Принудительно проставляет SB Sans Display, размер и цвет на run."""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rFonts.set(qn(f"w:{attr}"), REPORT_FONT_NAME)
    if font_size_pt is not None:
        run.font.size = Pt(font_size_pt)
    if color is not None:
        run.font.color.rgb = color
    if bold is not None:
        run.bold = bold


def _set_paragraph_text(
    paragraph: Paragraph,
    text: str,
    *,
    font_size_pt: float | None = REPORT_FONT_SIZE_PT,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    """Полностью заменяет текст параграфа, очищая run/hyperlink-содержимое и сохраняя стиль run.

    По умолчанию принудительно проставляет шрифт SB Sans Display 11 пт.

    Для заголовков (H1-H9) размер шрифта и bold НЕ форсятся — они должны
    наследоваться от paragraph-стиля и сохранённого rPr (иначе H3-заголовок
    «Исследование стенда …» сжимается с ~14 пт до 11 пт). Цвет, если задан,
    применяется в любом случае.
    """
    heading_level = _heading_level(paragraph)
    is_heading = heading_level is not None
    if is_heading:
        # Размер для заголовков задаём принудительно по уровню (H1-H3 = 14 pt,
        # H4 = 13.5 pt). Жирность оставляем за стилем шаблона.
        font_size_pt = _heading_size_pt_for_level(heading_level)
        bold = None
    saved_rpr = _clear_paragraph_inline(paragraph)
    run = paragraph.add_run(text)
    if saved_rpr is not None:
        existing_rpr = run._element.find(qn("w:rPr"))
        if existing_rpr is not None:
            run._element.remove(existing_rpr)
        run._element.insert(0, saved_rpr)
    _apply_run_font(run, font_size_pt=font_size_pt, color=color, bold=bold)
    if is_heading and font_size_pt is not None:
        # Сохранённый rPr может содержать «застывший» <w:sz>/<w:szCs> старого
        # размера — переопределяем явно после восстановления rPr.
        _force_run_font_size_pt(run, font_size_pt)


def _force_run_font_size_pt(run: Run, size_pt: float) -> None:
    """Жёстко проставляет на run-элемент `<w:sz>` и `<w:szCs>` указанного размера."""
    rpr = run._element.get_or_add_rPr()
    sz_value = str(int(round(size_pt * 2)))
    sz = rpr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr.append(sz)
    sz.set(qn("w:val"), sz_value)
    sz_cs = rpr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        rpr.append(sz_cs)
    sz_cs.set(qn("w:val"), sz_value)


def _set_heading_text(paragraph: Paragraph, text: str) -> None:
    """Алиас `_set_paragraph_text` для явной семантики «правим заголовок».

    Внутри `_set_paragraph_text` уже автоматически распознаёт заголовки и
    оставляет размер шрифта/жирность за стилем; этот хелпер существует ради
    читаемости вызывающего кода.
    """
    _set_paragraph_text(paragraph, text)


def _replace_in_paragraph(paragraph: Paragraph, old: str, new: str) -> bool:
    """Заменяет первое вхождение `old` на `new` в тексте параграфа.

    Если `old` содержится целиком в одном run, изменяется только этот run.
    Иначе весь параграф нормализуется в один run (с потерей mixed-форматирования).
    """
    if old not in paragraph.text:
        return False
    for run in paragraph.runs:
        if old in run.text:
            run.text = run.text.replace(old, new, 1)
            _apply_run_font(run)
            return True
    full = paragraph.text.replace(old, new, 1)
    _set_paragraph_text(paragraph, full)
    return True


def _replace_xxx_in_runs(paragraph: Paragraph, replacement: str) -> bool:
    """Заменяет каждое вхождение `XXX` в каждом run параграфа на `replacement`.

    Используется для обложки, где имя продукта вшито в текстовое поле и разбито
    на множество отдельных run-ов (часто `XXX` лежит в одном run, а соседний
    текст — в других). Стиль рана сохраняется, лишь меняется его текст.
    """
    changed = False
    for run in paragraph.runs:
        if "XXX" in run.text:
            run.text = run.text.replace("XXX", replacement)
            _apply_run_font(run)
            changed = True
    return changed


def _set_cell_text(
    cell: _Cell,
    text: str,
    *,
    font_size_pt: float | None = REPORT_FONT_SIZE_PT,
    color: RGBColor | None = None,
    bold: bool | None = None,
) -> None:
    """Заменяет содержимое ячейки, оставляя только один параграф со стилем шаблона."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        new_p = cell.add_paragraph()
        run = new_p.add_run(text)
        _apply_run_font(run, font_size_pt=font_size_pt, color=color, bold=bold)
        return
    _set_paragraph_text(paragraphs[0], text, font_size_pt=font_size_pt, color=color, bold=bold)
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)


def _heading_level(paragraph: Paragraph) -> int | None:
    """Возвращает уровень заголовка (1-9) для параграфа или None для обычного текста."""
    style = paragraph.style
    if style is None:
        return None
    name = (style.name or "").strip()
    match = re.match(r"^(?:Heading|Заголовок)\s*(\d+)$", name, flags=re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


def _iter_body_paragraphs(doc: DocxDocument) -> list[Paragraph]:
    """Возвращает список параграфов верхнего уровня в порядке следования."""
    return list(doc.paragraphs)


def _iter_all_paragraphs(doc: DocxDocument) -> list[Paragraph]:
    """Все `w:p` в документе, включая параграфы внутри текстовых полей (txbxContent).

    Используется для обработки обложки/футеров, где надписи живут не в основном
    потоке, а внутри `<w:drawing>/<wp:anchor>/.../<w:txbxContent>`.
    """
    body = doc.element.body
    paragraphs: list[Paragraph] = []
    for p_element in body.iter(qn("w:p")):
        paragraphs.append(Paragraph(p_element, doc))
    return paragraphs


def _block_elements_until_heading(start_para: Paragraph, max_stop_level: int = 2) -> list:
    """Собирает XML-элементы от стартового параграфа до следующего заголовка ≤ max_stop_level."""
    parent = start_para._parent
    body = start_para._element.getparent()
    children = list(body.iterchildren())
    idx = children.index(start_para._element)
    elements = [start_para._element]
    for child in children[idx + 1:]:
        if child.tag == qn("w:p"):
            tmp = Paragraph(child, parent)
            level = _heading_level(tmp)
            if level is not None and level <= max_stop_level:
                break
        elements.append(child)
    return elements


def _table_for_element(doc: DocxDocument, element) -> Table | None:
    if element.tag != qn("w:tbl"):
        return None
    for table in doc.tables:
        if table._element is element:
            return table
    return None


def _normalize_label(text: str) -> str:
    return re.sub(r"\s+", " ", text or "").strip().lower()


def _hosts_text(hosts: Iterable[Host]) -> str:
    parts: list[str] = []
    for host in hosts:
        label = host.hostname or host.ip_address
        if label:
            parts.append(label)
    return ", ".join(parts)


def _vuln_host_address(vuln: Vulnerability, indexes: dict) -> str:
    host_by_id: dict[UUID, Host] = indexes["host_by_id"]
    for asset in indexes["assets_by_vuln_id"].get(vuln.id, []):
        if asset.asset_type == AssetType.HOST:
            host = host_by_id.get(asset.asset_id)
            if host is not None:
                return host.hostname or host.ip_address or ""
    return ""


def _vuln_card_label(value: str, *, kind: ReportKind, vuln: Vulnerability, indexes: dict, project: Project) -> str:
    """Возвращает значение ячейки метрик-таблицы по русскому названию строки.

    Возвращает пустую строку, если данных нет.
    """
    label = _normalize_label(value)
    metrics = parse_cvss_vector(vuln.cvss_vector)
    if label.startswith("продукт"):
        return f"ПП «{project.name}»" if project.name else ""
    if label.startswith("адрес"):
        return _vuln_host_address(vuln, indexes)
    if label.startswith("cwe"):
        return vuln.cwe_id or ""
    if label.startswith("уровень риска"):
        return severity_label(vuln.severity)
    if label.startswith("cvss"):
        score_str = ""
        if vuln.cvss_score is not None:
            score_str = f"{vuln.cvss_score:g}"
        if score_str and vuln.cvss_vector:
            return f"{score_str} ({vuln.cvss_vector})"
        if score_str:
            return score_str
        return vuln.cvss_vector or ""
    if "вектор атаки" in label:
        return metrics.get("AV", "")
    if "сложность атаки" in label:
        return metrics.get("AC", "")
    if "требования к атаке" in label:
        return metrics.get("AT", "")
    if "уровень привилегий" in label:
        return metrics.get("PR", "")
    if "взаимодействие с пользователем" in label:
        return metrics.get("UI", "")
    if "влияние на конфиденциальность (vc)" in label or label == "влияние на конфиденциальность":
        return metrics.get("VC", "")
    if "влияние на целостность (vi)" in label or label == "влияние на целостность":
        return metrics.get("VI", "")
    if "влияние на доступность (va)" in label or label == "влияние на доступность":
        return metrics.get("VA", "")
    if "влияние на конфиденциальность (sc)" in label:
        return metrics.get("SC", "")
    if "влияние на целостность (si)" in label:
        return metrics.get("SI", "")
    if "влияние на доступность (sa)" in label:
        return metrics.get("SA", "")
    return ""


def _set_cell_label_and_score(
    cell: _Cell,
    *,
    score_text: str,
    vector_text: str,
    color: RGBColor | None,
    font_size_pt: float = VULN_TABLE_FONT_SIZE_PT,
) -> None:
    """Спец-обработчик ячейки CVSS: красит только число, вектор остаётся нейтральным."""
    paragraphs = cell.paragraphs
    if not paragraphs:
        target = cell.add_paragraph()
    else:
        target = paragraphs[0]
        _clear_paragraph_inline(target)
    score_run = target.add_run(score_text)
    _apply_run_font(score_run, font_size_pt=font_size_pt, color=color, bold=True)
    if vector_text:
        sep_run = target.add_run(" (")
        _apply_run_font(sep_run, font_size_pt=font_size_pt)
        vec_run = target.add_run(vector_text)
        _apply_run_font(vec_run, font_size_pt=font_size_pt)
        close_run = target.add_run(")")
        _apply_run_font(close_run, font_size_pt=font_size_pt)
    for p in paragraphs[1:]:
        p._element.getparent().remove(p._element)


# Размеры колонок карточной таблицы уязвимости (из шаблона СЗИ).
# tblGrid (twips) и значения tcW (pct/5000 = 100%).
_SZI_VULN_TABLE_GRID = ("5382", "4382")
_SZI_VULN_TABLE_CELL_PCT = ("2756", "2244")
# Серая заливка для заголовков «Показатели воздействия...» строк (из СЗИ).
_SZI_SECTION_ROW_FILL = "D9D9D9"


def _set_cell_shading(cell: _Cell, fill_hex: str) -> None:
    """Заливает фон ячейки цветом fill_hex (например, D9D9D9 — светло-серый)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)


def _set_cell_width_pct(cell: _Cell, pct_w: str) -> None:
    """Задаёт ширину ячейки в процентах (значение `w` для tcW type=pct)."""
    tc = cell._tc
    tcPr = tc.find(qn("w:tcPr"))
    if tcPr is None:
        tcPr = OxmlElement("w:tcPr")
        tc.insert(0, tcPr)
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), pct_w)
    tcW.set(qn("w:type"), "pct")


def _apply_szi_vuln_table_geometry(table: Table) -> None:
    """Выставляет на ПП-карточке тот же tblGrid/tcW, что и в СЗИ-шаблоне."""
    tbl = table._tbl
    tblGrid = tbl.find(qn("w:tblGrid"))
    if tblGrid is not None:
        cols = tblGrid.findall(qn("w:gridCol"))
        if len(cols) == 2:
            cols[0].set(qn("w:w"), _SZI_VULN_TABLE_GRID[0])
            cols[1].set(qn("w:w"), _SZI_VULN_TABLE_GRID[1])
    for row in table.rows:
        if len(row.cells) >= 2:
            _set_cell_width_pct(row.cells[0], _SZI_VULN_TABLE_CELL_PCT[0])
            _set_cell_width_pct(row.cells[1], _SZI_VULN_TABLE_CELL_PCT[1])


def _normalize_vuln_card_table_to_szi(table: Table) -> None:
    """Приводит карточную таблицу уязвимости (любой шаблон) к СЗИ-разметке.

    Шаблон ПП использует устаревшую CVSS 3.x-структуру: 13 строк без «Требования
    к атаке (AT)», без секций «Показатели воздействия на уязвимую/последующие
    системы» и без суффиксов (AV)/(AC)/(VC)/(SC)/.... Шаблон СЗИ — современная
    CVSS 4.0-разметка из 18 строк.

    Чтобы карточки в обоих отчётах выглядели одинаково (и заполнялись одной
    логикой `_fill_card_table`), переразмечаем PP-таблицу на лету:
      • переименовываем существующие строки (добавляем суффиксы метрик),
      • заменяем «Влияние на другие компоненты системы» на заголовок секции
        «Показатели воздействия на уязвимую систему»,
      • вставляем недостающие строки: AT, заголовок «Показатели воздействия на
        последующие системы» и три SC/SI/SA-метрики,
      • выставляем tblGrid/tcW как в СЗИ (узкая колонка значений, широкая
        колонка названий — иначе у ПП они «перевёрнуты»),
      • заливаем серым (D9D9D9) обе строки-заголовки секций воздействия.

    Функция идемпотентна: если таблица уже соответствует СЗИ-разметке (18 строк
    с правильными названиями), ничего не меняется (заливка/ширины тоже
    идемпотентны).
    """
    rows = table.rows
    if not rows:
        return
    # Считаем таблицу карточкой уязвимости только при наличии «Продукт» + «CVSS».
    labels_norm = [_normalize_label(r.cells[0].text) if r.cells else "" for r in rows]
    if not any(l.startswith("продукт") for l in labels_norm):
        return
    if not any(l.startswith("cvss") for l in labels_norm):
        return

    def _find_row_idx(predicate) -> int | None:
        for i, lbl in enumerate(labels_norm):
            if predicate(lbl):
                return i
        return None

    def _rename(idx: int, new_label: str) -> None:
        if idx is None or idx < 0 or idx >= len(table.rows):
            return
        cell = table.rows[idx].cells[0]
        _set_cell_text(cell, new_label, font_size_pt=VULN_TABLE_FONT_SIZE_PT)
        labels_norm[idx] = _normalize_label(new_label)

    def _clone_row_at(template_idx: int, new_label: str, after_idx: int) -> int:
        """Клонирует строку `template_idx` и вставляет её после `after_idx`.

        Возвращает индекс новой строки.
        """
        template_tr = table.rows[template_idx]._tr
        new_tr = copy.deepcopy(template_tr)
        # Вставляем XML-узел сразу после after_idx-й строки.
        anchor_tr = table.rows[after_idx]._tr
        anchor_tr.addnext(new_tr)
        # Перечитываем list строк (python-docx кеширует через свойство).
        new_idx = after_idx + 1
        # Обновим label и очистим значение (вторую ячейку).
        new_row = table.rows[new_idx]
        if len(new_row.cells) >= 1:
            _set_cell_text(new_row.cells[0], new_label, font_size_pt=VULN_TABLE_FONT_SIZE_PT)
        if len(new_row.cells) >= 2:
            _set_cell_text(new_row.cells[1], "", font_size_pt=VULN_TABLE_FONT_SIZE_PT)
        # Перестроим labels_norm соответствующе.
        labels_norm.insert(new_idx, _normalize_label(new_label))
        return new_idx

    # 1. Простые переименования (добавление суффиксов метрик).
    rename_pairs = [
        (lambda l: l == "вектор атаки", "Вектор атаки (AV)"),
        (lambda l: l == "сложность атаки", "Сложность атаки (AC)"),
        (lambda l: l == "уровень привилегий", "Уровень привилегий (PR)"),
        (lambda l: l == "взаимодействие с пользователем", "Взаимодействие с пользователем (UI)"),
        (lambda l: l == "влияние на другие компоненты системы", "Показатели воздействия на уязвимую систему"),
        (lambda l: l == "влияние на конфиденциальность", "Влияние на конфиденциальность (VC)"),
        (lambda l: l == "влияние на целостность", "Влияние на целостность (VI)"),
        (lambda l: l == "влияние на доступность", "Влияние на доступность (VA)"),
    ]
    for predicate, new_label in rename_pairs:
        idx = _find_row_idx(predicate)
        if idx is not None:
            _rename(idx, new_label)

    # 2. Вставляем «Требования к атаке (AT)» сразу после «Сложность атаки (AC)».
    if _find_row_idx(lambda l: "требования к атаке" in l) is None:
        ac_idx = _find_row_idx(lambda l: l.startswith("сложность атаки"))
        if ac_idx is not None:
            _clone_row_at(ac_idx, "Требования к атаке (AT)", after_idx=ac_idx)

    # 3. Если в шаблоне нет секционного заголовка «...на уязвимую систему», создаём
    #    его клонированием перед строкой «Влияние на конфиденциальность (VC)».
    if _find_row_idx(lambda l: "показатели воздействия на уязвимую систему" in l) is None:
        vc_idx = _find_row_idx(lambda l: l.startswith("влияние на конфиденциальность (vc)"))
        if vc_idx is not None and vc_idx > 0:
            # Клонируем предыдущую обычную строку и кладём как «заголовок секции».
            _clone_row_at(vc_idx - 1, "Показатели воздействия на уязвимую систему", after_idx=vc_idx - 1)

    # 4. Добавляем секцию «Показатели воздействия на последующие системы» + SC/SI/SA.
    if _find_row_idx(lambda l: "показатели воздействия на последующие системы" in l) is None:
        va_idx = _find_row_idx(lambda l: l.startswith("влияние на доступность (va)"))
        if va_idx is not None:
            after = _clone_row_at(va_idx, "Показатели воздействия на последующие системы", after_idx=va_idx)
            after = _clone_row_at(after, "Влияние на конфиденциальность (SC)", after_idx=after)
            after = _clone_row_at(after, "Влияние на целостность (SI)", after_idx=after)
            after = _clone_row_at(after, "Влияние на доступность (SA)", after_idx=after)

    # 5. Геометрия колонок — как в СЗИ-шаблоне (узкая колонка значений).
    _apply_szi_vuln_table_geometry(table)

    # 6. Серая заливка для строк-заголовков секций воздействия.
    for predicate in (
        lambda l: "показатели воздействия на уязвимую систему" in l,
        lambda l: "показатели воздействия на последующие системы" in l,
    ):
        idx = _find_row_idx(predicate)
        if idx is None:
            continue
        for cell in table.rows[idx].cells:
            _set_cell_shading(cell, _SZI_SECTION_ROW_FILL)


def _fill_card_table(table: Table, *, kind: ReportKind, vuln: Vulnerability, indexes: dict, project: Project) -> None:
    # Выравниваем структуру PP-карточки под СЗИ (см. `_normalize_vuln_card_table_to_szi`).
    _normalize_vuln_card_table_to_szi(table)

    severity_color = SEVERITY_COLORS.get(vuln.severity)
    score_str = f"{vuln.cvss_score:g}" if vuln.cvss_score is not None else ""
    vector_str = vuln.cvss_vector or ""

    for row in table.rows:
        if len(row.cells) < 2:
            continue
        label_cell = row.cells[0]
        value_cell = row.cells[1]
        label_norm = _normalize_label(label_cell.text)

        # Принудительно проставляем 10 пт + SB Sans Display и в ячейках-подписях.
        for paragraph in label_cell.paragraphs:
            for run in paragraph.runs:
                _apply_run_font(run, font_size_pt=VULN_TABLE_FONT_SIZE_PT)

        # «Уровень риска»: красим значение под severity.
        if label_norm.startswith("уровень риска"):
            value = severity_label(vuln.severity)
            if value or value_cell.text.strip():
                _set_cell_text(
                    value_cell,
                    value,
                    font_size_pt=VULN_TABLE_FONT_SIZE_PT,
                    color=severity_color,
                    bold=True if severity_color is not None else None,
                )
            continue

        # «CVSS v.4.0»: число — цветное, вектор — нейтральный.
        if label_norm.startswith("cvss"):
            if score_str or vector_str:
                _set_cell_label_and_score(
                    value_cell,
                    score_text=score_str,
                    vector_text=vector_str,
                    color=severity_color,
                )
            elif value_cell.text.strip():
                _set_cell_text(value_cell, "", font_size_pt=VULN_TABLE_FONT_SIZE_PT)
            continue

        new_value = _vuln_card_label(label_cell.text, kind=kind, vuln=vuln, indexes=indexes, project=project)
        if new_value or value_cell.text.strip():
            _set_cell_text(value_cell, new_value, font_size_pt=VULN_TABLE_FONT_SIZE_PT)


def _identify_section(text: str) -> str | None:
    norm = _normalize_label(text)
    for header in SECTION_HEADERS_ALL:
        if norm == _normalize_label(header):
            return header
    return None


_SECTION_CONTENT_DOUBLE_INDENT_TWIPS = 1134
"""Базовый левый отступ для контентных абзацев разделов карточек, в твипсах.

1134 твипс ≈ 2 см. Применяется в `_double_paragraph_left_indent` для текста
*под* жирными подписями разделов («Описание теста:», «Шаги для воспроизведения»
и т. п.). Сами подписи мы не трогаем — их положение задаётся шаблоном.
"""


def _double_paragraph_left_indent(paragraph: Paragraph) -> None:
    """Удваивает (или задаёт по умолчанию) левый отступ параграфа.

    Используется для контентных абзацев разделов карточек уязвимости/теста:
    пользователь попросил, чтобы текст под жирной подписью раздела имел
    «отступ от начала строки два раза». Поэтому:

    * если в `<w:pPr><w:ind>` уже стоит `w:left="N"` — заменяем на `2*N`;
    * если отступа не было — ставим `_SECTION_CONTENT_DOUBLE_INDENT_TWIPS`
      (≈2 см), что визуально соответствует «два уровня отступа» от левого
      края страницы.

    Также сбрасываем `w:firstLine`/`w:hanging`, чтобы Word не «съезжал»
    при отрисовке.
    """
    p_element = paragraph._element
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    ind = pPr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        pPr.append(ind)
    current_left_str = ind.get(qn("w:left"))
    try:
        current_left = int(current_left_str) if current_left_str is not None else 0
    except ValueError:
        current_left = 0
    new_left = max(current_left * 2, _SECTION_CONTENT_DOUBLE_INDENT_TWIPS)
    ind.set(qn("w:left"), str(new_left))
    for attr in ("w:firstLine", "w:hanging", "w:firstLineChars", "w:hangingChars"):
        if ind.get(qn(attr)) is not None:
            del ind.attrib[qn(attr)]


def _fill_card_prose(card_paragraphs: list[Paragraph], vuln: Vulnerability) -> None:
    """Подставляет описание/шаги/влияние/рекомендации в параграфы карточки.

    Структура: заголовки разделов из шаблона сохраняются. Внутри каждого раздела
    шаблонные параграфы заменяются на:
    - данные уязвимости — если поле заполнено;
    - единственный «…»-плейсхолдер — если данных нет.
    """
    sections: dict[str, list[Paragraph]] = {}
    current: str | None = None
    for p in card_paragraphs:
        header = _identify_section(p.text)
        if header is not None:
            current = header
            sections.setdefault(current, [])
            continue
        if current is None:
            continue
        sections[current].append(p)

    section_value: dict[str, str] = {
        SECTION_HEADER_FUNC: (vuln.description or "").strip(),
        SECTION_HEADER_STEPS: _workflow_steps_text(vuln).strip(),
        SECTION_HEADER_IMPACT: (vuln.impact or "").strip(),
        SECTION_HEADER_RECOMMENDATIONS: (vuln.recommendations or "").strip(),
    }

    # Собираем все контентные абзацы карточки, чтобы в конце удвоить им левый
    # отступ — пользователь попросил, чтобы текст *под* жирной подписью имел
    # «отступ от начала строки два раза».
    content_paragraphs: list[Paragraph] = []

    for header, paragraphs in sections.items():
        if not paragraphs:
            continue
        value = section_value.get(header)
        if value is None:
            # Раздел "Детальное описание уязвимости/недостатка" — пропускаем.
            continue
        first = paragraphs[0]
        if not value:
            _set_paragraph_text(first, "…")
            content_paragraphs.append(first)
            for tail in paragraphs[1:]:
                tail._element.getparent().remove(tail._element)
            continue
        lines = [line for line in value.splitlines() if line.strip() != ""] or [value]
        # Шаги шаблона (стиль `a6`) приходят жирным/«листовым» — для шагов
        # воспроизведения принудительно ставим обычное начертание.
        is_steps = header == SECTION_HEADER_STEPS
        _set_paragraph_text(first, lines[0], bold=False if is_steps else None)
        if is_steps:
            _strip_paragraph_numbering(first._element)
        content_paragraphs.append(first)
        anchor = first._element
        for line in lines[1:]:
            new_p = copy.deepcopy(first._element)
            for r in list(new_p.findall(qn("w:r"))):
                new_p.remove(r)
            anchor.addnext(new_p)
            wrapper = Paragraph(new_p, first._parent)
            run = wrapper.add_run(line)
            _apply_run_font(run, bold=False if is_steps else None)
            if is_steps:
                _strip_paragraph_numbering(new_p)
            content_paragraphs.append(wrapper)
            anchor = new_p
        for tail in paragraphs[1:]:
            tail._element.getparent().remove(tail._element)

    for content_p in content_paragraphs:
        _double_paragraph_left_indent(content_p)


def _flatten_step_description(description: str | None) -> str:
    """Сворачивает многострочное описание этапа в одну строку.

    Если этап описан в виде «1. сделать А\n2. сделать Б», в отчёте мы хотим
    получить «1. сделать А 2. сделать Б» внутри одной строки этапа.
    Markdown-ссылки на картинки (`![alt](/api/v1/files/.../download)`) убираем
    из текста — картинка вставляется отдельно.
    """
    if not description:
        return ""
    cleaned = _strip_markdown_images(description)
    parts = [line.strip() for line in cleaned.splitlines() if line.strip()]
    return " ".join(parts)


def _step_one_line(index: int, step: dict) -> str:
    """Формирует одну строку для этапа.

    Заголовок этапа объединяется с описанием (свёрнутым в одну строку):
    - title + description → «N. <title>: <description>»;
    - только title       → «N. <title>»;
    - только description → «N. <description>» (без редундантного «Этап N»);
    - пусто              → «N. Этап N» (запасной вариант, обычно не встречается).

    Если описание уже начинается с собственной нумерации «N. ...», внешний
    префикс «N.» не добавляется, чтобы не получить «1. 1. ...».
    """
    title = (step.get("title") or "").strip()
    description = _flatten_step_description(step.get("description"))
    if title and description:
        return f"{index}. {title}: {description}"
    if title:
        return f"{index}. {title}"
    if description:
        if re.match(r"^\d+[.)]\s", description):
            return description
        return f"{index}. {description}"
    return f"{index}. Этап {index}"


def _workflow_steps_text(vuln: Vulnerability) -> str:
    steps = vuln.workflow_steps or []
    if not steps:
        return _flatten_step_description(vuln.steps_to_reproduce)
    return "\n".join(_step_one_line(idx, step) for idx, step in enumerate(steps, start=1))


def _find_card_anchor(doc: DocxDocument, section_h2_text: str) -> Paragraph | None:
    """Находит первый параграф-карточку (Heading 3 или Heading 4) после H2 раздела."""
    paragraphs = _iter_body_paragraphs(doc)
    norm_target = _normalize_label(section_h2_text)
    in_section = False
    for paragraph in paragraphs:
        level = _heading_level(paragraph)
        norm = _normalize_label(paragraph.text)
        if level == 2 and norm == norm_target:
            in_section = True
            continue
        if in_section:
            if level in (3, 4) and norm.startswith("sa"):
                return paragraph
            if level == 2:
                return None
    return None


def _card_block(start_para: Paragraph) -> list:
    """XML-элементы карточки (заголовок + таблица + проза) до следующего H2/H1."""
    return _block_elements_until_heading(start_para, max_stop_level=2)


def _strip_paragraph_numbering(paragraph_element) -> None:
    """Отключает Word-нумерацию/маркеры для параграфа (через `numId=0`).

    Шаблонный стиль шагов (`a6` и т.п.) задаёт сквозную авто-нумерацию, из-за
    чего перед каждой нашей строкой Word добавляет «4.», «5.», … Чтобы
    оставить только наш собственный префикс «N. <текст>», выставляем на
    параграфе `<w:numPr><w:numId w:val="0"/></w:numPr>` — это явная команда
    Word'у «не нумеровать этот параграф», перекрывающая стиль.
    """
    pPr = paragraph_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph_element.insert(0, pPr)
    numPr = pPr.find(qn("w:numPr"))
    if numPr is None:
        numPr = OxmlElement("w:numPr")
        pPr.append(numPr)
    for child in list(numPr):
        numPr.remove(child)
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "0")
    numPr.append(numId)


def _force_center_paragraph(paragraph_element) -> None:
    """Проставляет `<w:pPr><w:jc w:val="center"/></w:pPr>` напрямую через XML.

    Делается так, чтобы гарантированно сработало даже на свежесозданных `<w:p>`
    без стиля, где `Paragraph.alignment = ...` иногда может не установить jc.
    """
    pPr = paragraph_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        paragraph_element.insert(0, pPr)
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "center")


def _add_picture_after(
    reference_element,
    doc: DocxDocument,
    image_bytes: bytes,
    width_in: float = 5.5,
):
    """Добавляет изображение как параграф сразу после reference_element.

    Важно: картинка ДОЛЖНА создаваться внутри целевого документа (`doc`), иначе
    связь (relationship id) указывает на временный документ и Word показывает
    заглушку «Не удаётся отобразить рисунок». Параграф принудительно выравнивается
    по центру, чтобы картинка не «уезжала» к левому краю. Возвращает новый XML-элемент.
    """
    new_p = OxmlElement("w:p")
    reference_element.addnext(new_p)
    paragraph = Paragraph(new_p, doc)
    run = paragraph.add_run()
    run.add_picture(BytesIO(image_bytes), width=Inches(width_in))
    _force_center_paragraph(new_p)
    return new_p


def _apply_run_highlight(run: Run, value: str = "yellow") -> None:
    """Проставляет на run жёлтую/иную заливку (`<w:highlight>`)."""
    rPr = run._element.get_or_add_rPr()
    highlight = rPr.find(qn("w:highlight"))
    if highlight is None:
        highlight = OxmlElement("w:highlight")
        rPr.append(highlight)
    highlight.set(qn("w:val"), value)


_CAPTION_STYLE_NAMES: tuple[str, ...] = (
    "caption",
    "название объекта",
    "название",
    "подпись",
)


def _find_caption_style_id(doc: DocxDocument) -> str | None:
    """Возвращает styleId параграф-стиля «Caption»/«Название объекта»/«Подпись».

    Это тот самый стиль, который Word применяет к подписи, созданной через
    «Ссылки → Вставить название». Если в шаблоне стиля нет — None, и подпись
    будет жить без явного pStyle (Word всё равно сможет её распознать как
    подпись благодаря полю `SEQ Figure`, но связать «Insert Caption»-форматом
    уже не получится автоматически).
    """
    styles_part = getattr(doc.part, "styles_part", None)
    styles_root = styles_part.element if styles_part is not None else doc.styles.element
    if styles_root is None:
        return None
    for style in styles_root.iter(qn("w:style")):
        if style.get(qn("w:type")) != "paragraph":
            continue
        name_el = style.find(qn("w:name"))
        if name_el is None:
            continue
        name_value = (name_el.get(qn("w:val")) or "").strip().lower()
        if not name_value:
            continue
        if name_value in _CAPTION_STYLE_NAMES:
            return style.get(qn("w:styleId"))
    return None


def _apply_paragraph_style(p_element, style_id: str) -> None:
    """Проставляет `<w:pStyle w:val="<style_id>">` на абзац (создаёт `<w:pPr>` при необходимости)."""
    pPr = p_element.find(qn("w:pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_element.insert(0, pPr)
    pStyle = pPr.find(qn("w:pStyle"))
    if pStyle is None:
        pStyle = OxmlElement("w:pStyle")
        pPr.insert(0, pStyle)
    pStyle.set(qn("w:val"), style_id)


def _add_figure_caption_after(
    reference_element,
    doc: DocxDocument,
    *,
    figure_number: int,
    caption_text: str,
):
    """Вставляет подпись рисунка через стандартную Word-функцию «Вставить название».

    Структура полностью повторяет то, что Word создаёт сам через
    «Ссылки → Вставить название»:

    * абзац оформлен встроенным стилем «Название объекта» / «Caption» (тот же
      `pStyle`, что и при ручной вставке подписи);
    * текст «Рисунок », поле `SEQ Figure \\* ARABIC`, текст « — <описание>»;
    * описание подкрашено жёлтым highlight, как мы и договаривались ранее.

    Поле `SEQ Figure` помечено `w:dirty="true"` — Word пересчитает нумерацию
    при открытии файла (включён `updateFields`) или вручную (Ctrl+A → F9),
    поэтому удаление/добавление подписей автоматически пересортирует номера.

    Возвращает только что созданный XML-элемент `<w:p>`.
    """
    new_p_element = OxmlElement("w:p")
    reference_element.addnext(new_p_element)

    caption_style_id = _find_caption_style_id(doc)
    if caption_style_id:
        _apply_paragraph_style(new_p_element, caption_style_id)

    paragraph = Paragraph(new_p_element, doc)

    prefix_run = paragraph.add_run("Рисунок\u00a0")
    _apply_run_font(prefix_run, font_size_pt=REPORT_FONT_SIZE_PT)
    _apply_run_highlight(prefix_run, "yellow")
    _force_caption_run(prefix_run)

    _append_seq_field(paragraph, sequence_name="Figure", initial_value=figure_number)

    suffix_text = f" — {caption_text}" if caption_text else ""
    if suffix_text:
        suffix_run = paragraph.add_run(suffix_text)
        _apply_run_font(suffix_run, font_size_pt=REPORT_FONT_SIZE_PT)
        _apply_run_highlight(suffix_run, "yellow")
        _force_caption_run(suffix_run)

    _force_center_paragraph(new_p_element)
    return new_p_element


def _append_seq_field(paragraph: Paragraph, *, sequence_name: str, initial_value: int) -> None:
    """Добавляет в параграф поле `SEQ <sequence_name> \\* ARABIC \\* MERGEFORMAT`.

    Особенности:

    * `\\* ARABIC` — арабские цифры в выводе.
    * `\\* MERGEFORMAT` — Word сохраняет ручное форматирование run'а при
      пересчёте, поэтому подпись остаётся не курсивной/чёрной даже после F9.
    * `w:dirty="true"` ставится на каждый `<w:fldChar>` (begin/separate/end);
      благодаря этому Word при открытии (включён `updateFields` в settings.xml)
      и при `Ctrl+A → F9` гарантированно пересобирает значение, что даёт
      корректную сквозную нумерацию даже после ручного удаления/перестановки
      подписей внутри Word.
    * Все run'ы дополнительно нормализуются: жёлтый highlight, шрифт SB Sans,
      снятый курсив и чёрный цвет (через `_force_caption_run_appearance`).
    """
    p_element = paragraph._element

    def _make_field_rpr() -> OxmlElement:
        rpr = OxmlElement("w:rPr")
        _apply_run_highlight_xml(rpr, "yellow")
        _apply_run_font_xml(rpr)
        _force_caption_run_appearance(rpr)
        return rpr

    begin_run = OxmlElement("w:r")
    begin_run.append(_make_field_rpr())
    begin_fld = OxmlElement("w:fldChar")
    begin_fld.set(qn("w:fldCharType"), "begin")
    begin_fld.set(qn("w:dirty"), "true")
    begin_run.append(begin_fld)
    p_element.append(begin_run)

    instr_run = OxmlElement("w:r")
    instr_run.append(_make_field_rpr())
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = f" SEQ {sequence_name} \\* ARABIC \\* MERGEFORMAT "
    instr_run.append(instr_text)
    p_element.append(instr_run)

    sep_run = OxmlElement("w:r")
    sep_run.append(_make_field_rpr())
    sep_fld = OxmlElement("w:fldChar")
    sep_fld.set(qn("w:fldCharType"), "separate")
    sep_fld.set(qn("w:dirty"), "true")
    sep_run.append(sep_fld)
    p_element.append(sep_run)

    value_run = OxmlElement("w:r")
    value_run.append(_make_field_rpr())
    value_text = OxmlElement("w:t")
    value_text.text = str(initial_value)
    value_run.append(value_text)
    p_element.append(value_run)

    end_run = OxmlElement("w:r")
    end_run.append(_make_field_rpr())
    end_fld = OxmlElement("w:fldChar")
    end_fld.set(qn("w:fldCharType"), "end")
    end_fld.set(qn("w:dirty"), "true")
    end_run.append(end_fld)
    p_element.append(end_run)


def _apply_run_highlight_xml(rpr_element, value: str) -> None:
    """Версия `_apply_run_highlight` для прямой работы с уже созданным `<w:rPr>`."""
    highlight = rpr_element.find(qn("w:highlight"))
    if highlight is None:
        highlight = OxmlElement("w:highlight")
        rpr_element.append(highlight)
    highlight.set(qn("w:val"), value)


def _force_caption_run_appearance(rpr_element) -> None:
    """Снимает курсив и красит run в чёрный — перекрывает наследие `caption`-стиля.

    Шаблонный `caption` (styleId=`af7`) задаёт `<w:i/>`, `<w:iCs/>` и тёмно-серый
    цвет `44546A`. Нам же нужен обычный (не курсивный) чёрный текст. Поэтому
    явно прописываем на каждом run'е подписи `<w:i w:val="false"/>`,
    `<w:iCs w:val="false"/>` и `<w:color w:val="000000"/>`. Без этого Word
    наследует курсив/цвет из стиля и подписи рисунков выглядят неправильно.
    """
    italic = rpr_element.find(qn("w:i"))
    if italic is None:
        italic = OxmlElement("w:i")
        rpr_element.append(italic)
    italic.set(qn("w:val"), "false")

    italic_cs = rpr_element.find(qn("w:iCs"))
    if italic_cs is None:
        italic_cs = OxmlElement("w:iCs")
        rpr_element.append(italic_cs)
    italic_cs.set(qn("w:val"), "false")

    color = rpr_element.find(qn("w:color"))
    if color is None:
        color = OxmlElement("w:color")
        rpr_element.append(color)
    color.set(qn("w:val"), "000000")
    # Удаляем привязку к теме, иначе тема перекрывает явный val.
    for theme_attr in ("w:themeColor", "w:themeTint", "w:themeShade"):
        if color.get(qn(theme_attr)) is not None:
            del color.attrib[qn(theme_attr)]


def _force_caption_run(run: Run) -> None:
    """То же, что `_force_caption_run_appearance`, но для готового `Run` python-docx."""
    rpr = run._element.get_or_add_rPr()
    _force_caption_run_appearance(rpr)


def _apply_run_font_xml(rpr_element) -> None:
    """Применяет шрифт `REPORT_FONT_NAME`/`REPORT_FONT_SIZE_PT` к существующему `<w:rPr>`."""
    rfonts = rpr_element.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr_element.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(qn(attr), REPORT_FONT_NAME)
    sz_value = str(int(REPORT_FONT_SIZE_PT * 2))
    sz = rpr_element.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        rpr_element.append(sz)
    sz.set(qn("w:val"), sz_value)
    sz_cs = rpr_element.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        rpr_element.append(sz_cs)
    sz_cs.set(qn("w:val"), sz_value)


# Паттерн markdown-картинки в описании этапа:
# `![любой alt](/api/v1/files/<uuid>/download)` или с абсолютным URL.
_MARKDOWN_IMAGE_RE = re.compile(
    r"!\[[^\]]*\]\(\s*(?:https?://[^)\s]*)?/api/v1/files/([0-9a-fA-F-]{36})/download[^)]*\)"
)


def _extract_file_ids_from_markdown(text: str | None) -> list[UUID]:
    """Возвращает UUID файлов, упомянутых как markdown-картинки в тексте.

    Фронт Markdown-редактор при добавлении картинки в шаг вставляет ссылку вида
    `![image.png](/api/v1/files/<UUID>/download)`. В отчёте нам нужно вывести
    саму картинку, а не ссылку.
    """
    if not text:
        return []
    result: list[UUID] = []
    seen: set[UUID] = set()
    for match in _MARKDOWN_IMAGE_RE.finditer(text):
        try:
            file_id = UUID(match.group(1))
        except ValueError:
            continue
        if file_id in seen:
            continue
        seen.add(file_id)
        result.append(file_id)
    return result


def _strip_markdown_images(text: str | None) -> str:
    """Убирает markdown-вставки картинок из текста, оставляя читаемое описание."""
    if not text:
        return ""
    return _MARKDOWN_IMAGE_RE.sub("", text)


def _step_image_file_ids(step: dict) -> list[UUID]:
    """Все file_id этапа: из `image_file_ids` и из markdown в title/description."""
    ids: list[UUID] = []
    seen: set[UUID] = set()
    for raw_id in step.get("image_file_ids") or []:
        try:
            file_id = UUID(str(raw_id))
        except (TypeError, ValueError):
            continue
        if file_id in seen:
            continue
        seen.add(file_id)
        ids.append(file_id)
    for source in (step.get("title"), step.get("description")):
        for file_id in _extract_file_ids_from_markdown(source):
            if file_id in seen:
                continue
            seen.add(file_id)
            ids.append(file_id)
    return ids


_FIGURE_COUNTER_KEY = "_figure_counter"


def _collect_step_image_ids(steps: list[dict] | None) -> set[UUID]:
    """Собирает множество file_id этапов (`image_file_ids` + markdown-вставки)."""
    ids: set[UUID] = set()
    for step in steps or []:
        ids.update(_step_image_file_ids(step))
    return ids


def _next_figure_number(image_state: dict) -> int:
    """Берёт следующее число «Рисунок N» из общего счётчика построения отчёта."""
    counter = image_state.get(_FIGURE_COUNTER_KEY, 0) + 1
    image_state[_FIGURE_COUNTER_KEY] = counter
    return counter


def _figure_caption_text(*, step: dict | None, file_name: str | None, fallback_index: int) -> str:
    """Возвращает «описание картинки» для подписи."""
    if file_name:
        # Убираем расширение файла, если есть.
        name = re.sub(r"\.[A-Za-z0-9]{1,8}$", "", file_name).strip()
        if name:
            return name
    if step is not None:
        title = (step.get("title") or "").strip()
        if title:
            return title
        description = _flatten_step_description(step.get("description"))
        if description:
            return description[:120]
    return f"Этап {fallback_index}"


def _insert_image_with_caption(
    *,
    anchor_element,
    doc: DocxDocument,
    image_bytes: bytes,
    caption_text: str,
    image_state: dict,
):
    """Вставляет картинку и подпись «Рисунок N — <caption>» (жёлтый highlight) после anchor.

    Сам номер N оформлен как Word-поле `SEQ Figure`, поэтому Word автоматически
    пересчитывает нумерацию при удалении/добавлении подписей (см. `_add_figure_caption_after`).
    Возвращает элемент-подпись (последний вставленный), который должен стать новым anchor.
    """
    pic_element = _add_picture_after(anchor_element, doc, image_bytes)
    figure_number = _next_figure_number(image_state)
    return _add_figure_caption_after(
        pic_element,
        doc,
        figure_number=figure_number,
        caption_text=caption_text,
    )


def _fill_card(
    block_elements: list,
    *,
    kind: ReportKind,
    doc: DocxDocument,
    vuln: Vulnerability,
    sequence: int,
    indexes: dict,
    project: Project,
    image_bytes_by_id: dict[UUID, bytes],
    image_state: dict,
) -> None:
    """Заполняет одну карточку уязвимости/слабости данными `vuln`."""
    parent = doc.part
    paragraphs: list[Paragraph] = []
    table: Table | None = None
    for element in block_elements:
        if element.tag == qn("w:p"):
            paragraphs.append(Paragraph(element, doc))
        elif element.tag == qn("w:tbl") and table is None:
            for t in doc.tables:
                if t._element is element:
                    table = t
                    break
    if not paragraphs:
        return

    # Заголовок карточки: SA-NN: <название уязвимости>.
    heading = paragraphs[0]
    title = (vuln.title or "").strip()
    label = f"SA-{sequence:02d}"
    new_heading = f"{label}: {title}" if title else label
    _set_heading_text(heading, new_heading)

    if table is not None:
        _fill_card_table(table, kind=kind, vuln=vuln, indexes=indexes, project=project)

    # Параграфы после таблицы (или после заголовка, если таблицы нет в карточке).
    after_table_paragraphs: list[Paragraph] = []
    table_seen = table is None
    for element in block_elements[1:]:
        if element.tag == qn("w:tbl"):
            table_seen = True
            continue
        if not table_seen:
            continue
        if element.tag == qn("w:p"):
            after_table_paragraphs.append(Paragraph(element, doc))

    _fill_card_prose(after_table_paragraphs, vuln)
    _attach_workflow_images(
        after_table_paragraphs,
        vuln,
        image_bytes_by_id=image_bytes_by_id,
        files_by_id=indexes.get("files_by_id", {}),
        files_by_vuln_id=indexes.get("files_by_vuln_id", {}),
        image_state=image_state,
        doc=doc,
    )


def _attach_workflow_images(
    card_paragraphs: list[Paragraph],
    vuln: Vulnerability,
    *,
    image_bytes_by_id: dict[UUID, bytes],
    files_by_id: dict[UUID, File],
    files_by_vuln_id: dict[UUID, list[File]],
    image_state: dict,
    doc: DocxDocument,
) -> None:
    """Вставляет картинки в секцию «Шаги для воспроизведения» с подписями «Рисунок N — …».

    В отчёт идут ТОЛЬКО картинки, явно привязанные к этапам уязвимости
    (через `image_file_ids` или markdown в описании этапа). «Свободные»
    файлы-доказательства, оставшиеся в БД после удалений, в отчёт не попадают.
    """
    del files_by_vuln_id  # extras-картинки больше не вставляем — см. docstring
    has_step_images = bool(_collect_step_image_ids(vuln.workflow_steps) & set(image_bytes_by_id))
    if not has_step_images:
        return

    # Находим заголовок «Шаги для воспроизведения» среди ИСХОДНЫХ параграфов карточки.
    # `_fill_card_prose` уже разложил каждый этап в отдельный параграф — собираем
    # их прямо из живого XML (а не по устаревшему списку card_paragraphs).
    steps_header_para: Paragraph | None = None
    for p in card_paragraphs:
        if _identify_section(p.text) == SECTION_HEADER_STEPS:
            steps_header_para = p
            break
    if steps_header_para is None:
        return

    # Параграфы-этапы — все `<w:p>` между заголовком STEPS и следующим
    # известным заголовком секции (IMPACT/RECOMMENDATIONS/…).
    step_paragraph_elements: list = []
    cursor = steps_header_para._element.getnext()
    while cursor is not None:
        if cursor.tag == qn("w:p"):
            text = "".join(t.text or "" for t in cursor.iter(qn("w:t")))
            if _identify_section(text) is not None:
                break
            step_paragraph_elements.append(cursor)
        cursor = cursor.getnext()

    fallback_anchor = (
        step_paragraph_elements[-1] if step_paragraph_elements else steps_header_para._element
    )

    steps = vuln.workflow_steps or []
    for step_index, step in enumerate(steps, start=1):
        file_ids = _step_image_file_ids(step)
        if not file_ids:
            continue
        if step_index - 1 < len(step_paragraph_elements):
            anchor = step_paragraph_elements[step_index - 1]
        else:
            anchor = fallback_anchor
        for file_id in file_ids:
            data = image_bytes_by_id.get(file_id)
            if not data:
                continue
            file_meta = files_by_id.get(file_id)
            caption = _figure_caption_text(
                step=step,
                file_name=file_meta.original_name if file_meta is not None else None,
                fallback_index=step_index,
            )
            anchor = _insert_image_with_caption(
                anchor_element=anchor,
                doc=doc,
                image_bytes=data,
                caption_text=caption,
                image_state=image_state,
            )
        if step_index >= len(step_paragraph_elements):
            fallback_anchor = anchor


def _replicate_card(template_block: list, count: int) -> list[list]:
    """Создаёт `count` копий карточки. Первая копия — оригинал; остальные клонируются и вставляются после."""
    if count <= 0:
        return []
    blocks = [template_block]
    last_element = template_block[-1]
    for _ in range(count - 1):
        cloned: list = []
        for element in template_block:
            new_element = copy.deepcopy(element)
            last_element.addnext(new_element)
            last_element = new_element
            cloned.append(new_element)
        blocks.append(cloned)
    return blocks


def _remove_block(elements: list) -> None:
    for element in elements:
        parent = element.getparent()
        if parent is not None:
            parent.remove(element)


def _split_vulnerabilities(vulnerabilities: list[Vulnerability]) -> tuple[list[Vulnerability], list[Vulnerability]]:
    vulns: list[Vulnerability] = []
    weaknesses: list[Vulnerability] = []
    for vuln in vulnerabilities:
        if vuln.severity in WEAKNESS_SEVERITIES:
            weaknesses.append(vuln)
        else:
            vulns.append(vuln)
    return vulns, weaknesses


def _strip_run_highlight(run: Run) -> None:
    """Снимает `<w:highlight>` с run'а (нужно для шаблонных «жёлтых» ячеек)."""
    rPr = run._element.find(qn("w:rPr"))
    if rPr is None:
        return
    for hl in rPr.findall(qn("w:highlight")):
        rPr.remove(hl)


def _strip_cell_highlight(cell: _Cell) -> None:
    """Снимает желтую (и любую) подсветку со всех run'ов внутри ячейки."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            _strip_run_highlight(run)


def _update_severity_summary_table(doc: DocxDocument, severity_counts: dict[Severity, int]) -> None:
    """Обновляет колонку «Количество» в сводной таблице уязвимостей/недостатков."""
    for table in doc.tables:
        if not table.rows:
            continue
        header_text = table.rows[0].cells[0].text
        if _normalize_label(header_text) != _normalize_label(SUMMARY_TABLE_HEADER):
            continue
        for row in table.rows[1:]:
            if len(row.cells) < 2:
                continue
            label = _normalize_label(row.cells[0].text)
            severity = SUMMARY_ROW_LABEL_TO_SEVERITY.get(label)
            if severity is None:
                continue
            count = severity_counts.get(severity, 0)
            _set_cell_text(row.cells[1], str(count))
            # В шаблоне ячейки с количеством были помечены жёлтой подсветкой,
            # `_set_paragraph_text` сохраняет шаблонный rPr → подсветка
            # переносится на наше число. Снимаем её явно.
            _strip_cell_highlight(row.cells[1])
        return


def _update_project_intro(doc: DocxDocument, project: Project) -> None:
    """Заменяет «Программного Продукта XXX» и интервал работ в «Информация о проекте»."""
    name = (project.name or "").strip()
    start = _format_date(project.start_date)
    end = _format_date(project.end_date)
    pattern_dates_with_dots = re.compile(r"с \d{2}[.\-/]\d{2}[.\-/]\d{2,4} по \d{2}[.\-/]\d{2}[.\-/]\d{2,4}")

    for paragraph in _iter_body_paragraphs(doc):
        text = paragraph.text
        if "анализу защищенности Программного Продукта" not in text and "анализу защищенности Программного" not in text:
            continue
        if name:
            _replace_in_paragraph(paragraph, "Программного Продукта XXX", f"Программного Продукта {name}")
        if start and end:
            replaced = False
            for run in paragraph.runs:
                if pattern_dates_with_dots.search(run.text):
                    run.text = pattern_dates_with_dots.sub(f"с {start} по {end}", run.text, count=1)
                    replaced = True
                    break
            if not replaced and pattern_dates_with_dots.search(paragraph.text):
                new_text = pattern_dates_with_dots.sub(f"с {start} по {end}", paragraph.text, count=1)
                _set_paragraph_text(paragraph, new_text)
        break


def _update_cover_page(doc: DocxDocument, project: Project) -> None:
    """Подставляет название проекта на обложку отчёта.

    Текст обложки лежит внутри `<w:txbxContent>` (текстовых фреймов) и часто
    разнесён по нескольким соседним `<w:p>`-параграфам, причём «XXX» обычно
    лежит в отдельном параграфе. Поэтому ищем каждый `txbxContent`, и если в
    его суммарном тексте присутствует «Программного» или «Продукта», заменяем
    все «XXX» в любых run-ах внутри этого фрейма на имя проекта.
    """
    name = (project.name or "").strip()
    if not name:
        return
    body = doc.element.body
    for txbx in body.iter(qn("w:txbxContent")):
        full_text = "".join(t.text or "" for t in txbx.iter(qn("w:t")))
        if "Продукта" not in full_text or "XXX" not in full_text:
            continue
        # Идём по run-ам в порядке следования. Меняем «XXX» только если
        # предшествующий непустой run заканчивается словом «Продукта»
        # (это титульная строка с именем продукта). Так мы не задеваем
        # «distribVersion: XXX», которое идёт после двоеточия.
        prev_text = ""
        for run_element in txbx.iter(qn("w:r")):
            run_text = "".join(t.text or "" for t in run_element.iter(qn("w:t")))
            if not run_text:
                continue
            stripped_prev = prev_text.rstrip()
            after_product = stripped_prev.endswith("Продукта") or stripped_prev.endswith("Продукт")
            if after_product and "XXX" in run_text:
                for text_element in run_element.iter(qn("w:t")):
                    if text_element.text and "XXX" in text_element.text:
                        text_element.text = text_element.text.replace("XXX", name)
            prev_text = run_text


_COVER_DATE_RE = re.compile(r"Москва,\s*\d{1,2}\.\d{4}")


def _update_cover_date(doc: DocxDocument, project: Project) -> None:
    """Заменяет «Москва, MM.YYYY» на обложке отчёта актуальной датой проекта.

    Дата выбирается по приоритету: `project.end_date` → `project.start_date` →
    текущая дата. Формат — `MM.YYYY` (две цифры месяца, точка, четыре цифры
    года). Текст обложки лежит внутри `<w:txbxContent>` и часто разнесён по
    нескольким `<w:t>` в одном параграфе, поэтому работаем на уровне параграфа:
    склеиваем все его текст-элементы, делаем замену на полной строке и пишем
    результат в первый `<w:t>`, очищая остальные.
    """
    target_date = project.end_date or project.start_date or _date.today()
    try:
        formatted = target_date.strftime("%m.%Y")
    except AttributeError:
        return
    new_text = f"Москва, {formatted}"

    body = doc.element.body
    for txbx in body.iter(qn("w:txbxContent")):
        txbx_text = "".join(t.text or "" for t in txbx.iter(qn("w:t")))
        if not _COVER_DATE_RE.search(txbx_text):
            continue
        for paragraph_element in txbx.iter(qn("w:p")):
            text_elements = list(paragraph_element.iter(qn("w:t")))
            if not text_elements:
                continue
            full_text = "".join(t.text or "" for t in text_elements)
            if not _COVER_DATE_RE.search(full_text):
                continue
            replaced = _COVER_DATE_RE.sub(new_text, full_text)
            if replaced == full_text:
                continue
            text_elements[0].text = replaced
            for tail in text_elements[1:]:
                tail.text = ""


def _enable_auto_field_update(doc: DocxDocument) -> None:
    """Добавляет `<w:updateFields w:val="true"/>` в `word/settings.xml`.

    Это стандартная директива OOXML, которая просит Word при открытии документа
    автоматически обновить все поля — TOC, перекрёстные ссылки, PAGEREF и т.д.
    Word покажет предупреждение «Update fields» (Yes по-умолчанию у части
    локалей), но даже если пользователь нажимает Cancel, оглавление перестроится
    на основе фактических заголовков. Без этой директивы оглавление остаётся
    «застывшим» — таким, каким было сохранено в шаблоне.
    """
    settings_part = getattr(doc, "settings", None)
    settings_element = getattr(settings_part, "element", None)
    if settings_element is None:
        return
    existing = settings_element.find(qn("w:updateFields"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    # OOXML рекомендует ставить элемент в начало `<w:settings>`.
    settings_element.insert(0, update_fields)


# Имя бук-марки, которую мы создаём для каждого заголовка под наш TOC. Префикс
# отличается от стандартного `_Toc...`, чтобы не пересекаться с шаблонными.
_TOC_AUTO_BOOKMARK_PREFIX = "_TocAuto_"


def _find_outer_toc_field(doc: DocxDocument):
    """Возвращает (begin, separate, end) узлы внешнего TOC-поля или `None`.

    Идёт плоским потоком `fldChar`/`instrText` по всему телу документа,
    учитывая вложенные поля (PAGEREF внутри TOC и т.п.).
    """
    body = doc.element.body
    fld_char_tag = qn("w:fldChar")
    instr_text_tag = qn("w:instrText")
    fld_type_attr = qn("w:fldCharType")

    flat = [el for el in body.iter() if el.tag in (fld_char_tag, instr_text_tag)]
    stack: list[dict] = []
    for n in flat:
        if n.tag == fld_char_tag:
            ftype = n.get(fld_type_attr)
            if ftype == "begin":
                stack.append({"begin": n, "instr": [], "separate": None})
            elif ftype == "separate" and stack:
                stack[-1]["separate"] = n
            elif ftype == "end" and stack:
                frame = stack.pop()
                instr_full = "".join(frame["instr"]).strip().upper()
                if instr_full.startswith("TOC") and frame["separate"] is not None:
                    return frame["begin"], frame["separate"], n
        elif n.tag == instr_text_tag and stack:
            stack[-1]["instr"].append(n.text or "")
    return None


def _find_toc_style_ids_by_name(doc: DocxDocument) -> dict[int, str]:
    """Ищет в `word/styles.xml` стили с `<w:name>toc N</w:name>` и возвращает их styleId.

    Это надёжнее, чем угадывать по pStyle первой попавшейся кэш-записи: в
    шаблонах рядом могут лежать кастомные стили вроде «ЛЕТА_Заголовок 3»,
    у которых w:name НЕ начинается с «toc» — но именно их pStyle Word
    использовал в шаблонной TOC. Если их вытянуть, sub-item'ы оглавления
    выглядят как обычные жирные заголовки, без отступов и с гигантским sz.
    """
    style_ids: dict[int, str] = {}
    styles_part = getattr(doc.part, "styles_part", None)
    styles_root = styles_part.element if styles_part is not None else doc.styles.element
    if styles_root is None:
        return style_ids
    for style in styles_root.iter(qn("w:style")):
        name_el = style.find(qn("w:name"))
        if name_el is None:
            continue
        name = (name_el.get(qn("w:val")) or "").strip().lower()
        match = re.match(r"^toc\s*(\d+)$", name)
        if not match:
            continue
        level = int(match.group(1))
        if 1 <= level <= 9:
            sid = style.get(qn("w:styleId"))
            if sid:
                style_ids[level] = sid
    return style_ids


def _detect_toc_pstyles(cached_entries: list) -> dict[int, str]:
    """Fallback: вытаскивает pStyle из существующих записей оглавления.

    Используется только если в `styles.xml` нет именованных «toc N» (что в наших
    шаблонах не встречается, но на всякий случай есть деградация).
    """
    seen: list[str] = []
    for p in cached_entries:
        pPr = p.find(qn("w:pPr"))
        if pPr is None:
            continue
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is None:
            continue
        val = pStyle.get(qn("w:val"))
        if val and val not in seen:
            seen.append(val)
        if len(seen) >= 3:
            break
    return {i + 1: name for i, name in enumerate(seen[:3])}


def _detect_hyperlink_rstyle(cached_entries: list) -> str | None:
    """Находит rStyle гиперссылок в кэше TOC (обычно «ae» или «Hyperlink»)."""
    for p in cached_entries:
        for hyper in p.iter(qn("w:hyperlink")):
            for rPr in hyper.iter(qn("w:rPr")):
                rStyle = rPr.find(qn("w:rStyle"))
                if rStyle is not None:
                    val = rStyle.get(qn("w:val"))
                    if val:
                        return val
    return None


def _xml_escape(text: str) -> str:
    return (
        text.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace('"', "&quot;")
    )


def _build_toc_entry_p(toc_style: str, bookmark: str, text: str, hyperlink_style: str | None):
    """Собирает XML-параграф одной записи оглавления.

    Структура совпадает с тем, что генерирует Word: гиперссылка на бук-марку
    заголовка, текст заголовка, табуляция, вложенное поле PAGEREF, плейсхолдер
    «1» (Word заменит его на реальный номер при открытии). Шрифт — SB Sans
    Display, чтобы оглавление выглядело как остальной текст.
    """
    w_ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    safe_text = _xml_escape(text)
    safe_bookmark = _xml_escape(bookmark)
    style_xml = f'<w:pStyle w:val="{_xml_escape(toc_style)}"/>' if toc_style else ""
    rstyle_xml = f'<w:rStyle w:val="{_xml_escape(hyperlink_style)}"/>' if hyperlink_style else ""
    xml_str = (
        f'<w:p xmlns:w="{w_ns}">'
        f"<w:pPr>{style_xml}</w:pPr>"
        f'<w:hyperlink w:anchor="{safe_bookmark}" w:history="1">'
        "<w:r>"
        f"<w:rPr>{rstyle_xml}<w:rFonts w:ascii=\"SB Sans Display\" w:hAnsi=\"SB Sans Display\" w:cs=\"SB Sans Display\"/></w:rPr>"
        f'<w:t xml:space="preserve">{safe_text}</w:t>'
        "</w:r>"
        "<w:r><w:rPr><w:webHidden/></w:rPr><w:tab/></w:r>"
        '<w:r><w:rPr><w:webHidden/></w:rPr><w:fldChar w:fldCharType="begin"/></w:r>'
        f'<w:r><w:rPr><w:webHidden/></w:rPr><w:instrText xml:space="preserve"> PAGEREF {safe_bookmark} \\h </w:instrText></w:r>'
        '<w:r><w:rPr><w:webHidden/></w:rPr><w:fldChar w:fldCharType="separate"/></w:r>'
        '<w:r><w:rPr><w:webHidden/></w:rPr><w:t>1</w:t></w:r>'
        '<w:r><w:rPr><w:webHidden/></w:rPr><w:fldChar w:fldCharType="end"/></w:r>'
        "</w:hyperlink>"
        "</w:p>"
    )
    return parse_xml(xml_str)


def _assign_unique_heading_bookmarks(doc: DocxDocument) -> list[tuple[int, str, str]]:
    """Каждому H1/H2/H3 в теле документа присваивает уникальную бук-марку.

    Возвращает список `(level, bookmark_name, heading_text)` в порядке
    следования. Полученные имена бук-марок гарантированно уникальны (с префиксом
    `_TocAuto_`) — это важно, потому что после клонирования по-host-секций в
    документе остаются десятки заголовков с одинаковыми `_TocXXXXX`, и PAGEREF
    в кэше TOC резолвится по первому совпадению, давая «все страницы = 4».
    """
    body = doc.element.body
    # Стартуем со значений id выше всех существующих в теле бук-марок.
    max_id = 0
    for bm in body.iter(qn("w:bookmarkStart")):
        try:
            max_id = max(max_id, int(bm.get(qn("w:id"))))
        except (TypeError, ValueError):
            pass
    next_id = max_id + 1

    entries: list[tuple[int, str, str]] = []
    for paragraph in doc.paragraphs:
        level = _heading_level(paragraph)
        if level is None or level > 3:
            continue
        text = paragraph.text.strip()
        if not text:
            continue
        bookmark_name = f"{_TOC_AUTO_BOOKMARK_PREFIX}{next_id:06d}"
        bm_id = str(next_id)
        next_id += 1

        bs = OxmlElement("w:bookmarkStart")
        bs.set(qn("w:id"), bm_id)
        bs.set(qn("w:name"), bookmark_name)
        be = OxmlElement("w:bookmarkEnd")
        be.set(qn("w:id"), bm_id)
        # Бук-марка должна оборачивать видимый текст: bs идёт сразу после pPr,
        # be — последним child параграфа.
        pPr = paragraph._element.find(qn("w:pPr"))
        if pPr is not None:
            pPr.addnext(bs)
        else:
            paragraph._element.insert(0, bs)
        paragraph._element.append(be)

        entries.append((level, bookmark_name, text))
    return entries


def _rebuild_toc(doc: DocxDocument) -> None:
    """Заново собирает «закэшированное» содержимое TOC-поля.

    Решает две UX-проблемы реального открытия отчёта в Word:

    1. **Пустое оглавление при открытии.** Раньше мы вычищали body TOC-поля
       полностью; Word при `<w:updateFields w:val="true"/>` не всегда
       инициирует обновление пустого поля, и пользователь видит чистую
       страницу. Сейчас мы предзаполняем поле полным набором записей —
       один абзац на каждый H1/H2/H3 — поэтому даже без обновления видны
       заголовки и гиперссылки.

    2. **Все номера страниц = 1 после первого F9.** Когда Word обновляет
       PAGEREF за один проход, он берёт позиции бук-марок из ТЕКУЩЕЙ вёрстки.
       Если поле было пустым, добавление записей толкает страницы вниз —
       получаем «всё на странице 1». Если же кэш уже содержит финальное
       количество абзацев (наш случай), вёрстка стабильна и первый же
       перерасчёт даёт правильные номера.

    Дополнительно: каждый заголовок получает уникальную бук-марку
    `_TocAuto_<n>` (после клонирования по-host-секций оригинальные `_Toc...`
    дублируются и PAGEREF указывает не туда).
    """
    toc = _find_outer_toc_field(doc)
    if toc is None:
        return
    begin_node, sep_node, end_node = toc
    sep_p = _ancestor_paragraph(sep_node)
    end_p = _ancestor_paragraph(end_node)
    if sep_p is None or end_p is None or sep_p is end_p:
        return
    parent = sep_p.getparent()
    if parent is None or end_p.getparent() is not parent:
        return

    # 1. Стили TOC ищем по именам «toc 1/2/3» в styles.xml — это правильный
    #    источник; кэш TOC оставлен только как fallback (см. _detect_toc_pstyles).
    sep_idx = parent.index(sep_p)
    end_idx = parent.index(end_p)
    cached_entries = list(parent)[sep_idx + 1:end_idx]
    pstyle_for_level = _find_toc_style_ids_by_name(doc) or _detect_toc_pstyles(cached_entries)
    hyperlink_rstyle = _detect_hyperlink_rstyle(cached_entries)

    # 2. Уникальные бук-марки + список заголовков для оглавления.
    entries = _assign_unique_heading_bookmarks(doc)

    # 3. Сносим старый кэш записей.
    _remove_field_result(sep_node, end_node)

    # 4. Вставляем новые записи между sep_p и end_p.
    sep_idx = parent.index(sep_p)
    insert_at = sep_idx + 1
    for level, bookmark, text in entries:
        toc_style = pstyle_for_level.get(level) or pstyle_for_level.get(1) or ""
        new_p = _build_toc_entry_p(toc_style, bookmark, text, hyperlink_rstyle)
        parent.insert(insert_at, new_p)
        insert_at += 1

    # 5. Помечаем поле dirty: Word при открытии (при включённом updateFields)
    #    или при F9 пересчитает PAGEREF'ы и подставит реальные номера страниц.
    begin_node.set(qn("w:dirty"), "true")

    # 6. На случай документов с простыми TOC-полями <w:fldSimple>, помечаем их тоже.
    for fld_simple in doc.element.body.iter(qn("w:fldSimple")):
        instr = (fld_simple.get(qn("w:instr")) or "").upper()
        if "TOC" in instr:
            fld_simple.set(qn("w:dirty"), "true")


def _remove_field_result(separate_node, end_node) -> None:
    """Удаляет всё, что лежит между `separate` и `end` одного и того же поля.

    Поле может пересекать границы параграфов: сами `separate`/`end` fldChar'ы
    лежат в `<w:r>`, который — в `<w:p>`. «Результат» поля состоит из:
      • остатка параграфа, в котором лежит `separate` (всё после `separate`-run'а)
      • всех ПОЛНОСТЬЮ промежуточных параграфов (между параграфом separate и
        параграфом end)
      • начала параграфа end (всё до `end`-run'а)
    Параграфы separate и end не удаляем целиком, чтобы не сломать структуру
    самого поля.
    """
    sep_run = _ancestor_run(separate_node)
    end_run = _ancestor_run(end_node)
    if sep_run is None or end_run is None:
        return
    sep_p = _ancestor_paragraph(sep_run)
    end_p = _ancestor_paragraph(end_run)
    if sep_p is None or end_p is None:
        return

    # Чистим хвост параграфа separate: всё после sep_run.
    sep_p_kids = list(sep_p)
    try:
        sep_run_idx = sep_p_kids.index(sep_run)
    except ValueError:
        sep_run_idx = -1
    if sep_run_idx >= 0:
        for victim in sep_p_kids[sep_run_idx + 1:]:
            sep_p.remove(victim)

    # Если separate и end в одном параграфе — больше делать нечего: останется
    # тот же sep_run + end_run, между ними чисто.
    if sep_p is end_p:
        # На случай редкой ситуации, когда между sep_run и end_run в одном
        # параграфе ещё что-то осталось (например, end_run шёл раньше
        # хвостовых элементов): тоже удалим.
        kids = list(sep_p)
        try:
            si = kids.index(sep_run)
            ei = kids.index(end_run)
        except ValueError:
            return
        if si < ei:
            for victim in kids[si + 1:ei]:
                sep_p.remove(victim)
        return

    # Удаляем все параграфы строго между sep_p и end_p.
    parent = sep_p.getparent()
    if parent is None or end_p.getparent() is not parent:
        return
    parent_kids = list(parent)
    try:
        sep_p_idx = parent_kids.index(sep_p)
        end_p_idx = parent_kids.index(end_p)
    except ValueError:
        return
    for victim in parent_kids[sep_p_idx + 1:end_p_idx]:
        parent.remove(victim)

    # Чистим начало параграфа end: всё до end_run.
    # Сохраняем pPr (если есть) — это первый дочерний элемент параграфа.
    end_p_kids = list(end_p)
    try:
        end_run_idx = end_p_kids.index(end_run)
    except ValueError:
        return
    pPr_tag = qn("w:pPr")
    for victim in end_p_kids[:end_run_idx]:
        if victim.tag == pPr_tag:
            continue
        end_p.remove(victim)


def _ancestor_paragraph(node):
    p_tag = qn("w:p")
    cur = node.getparent()
    while cur is not None and cur.tag != p_tag:
        cur = cur.getparent()
    return cur


def _ancestor_run(node):
    r_tag = qn("w:r")
    cur = node.getparent()
    while cur is not None and cur.tag != r_tag:
        cur = cur.getparent()
    return cur


def _normalize_heading_sizes(doc: DocxDocument) -> None:
    """Принудительно проставляет размеры шрифта заголовкам по уровням.

    H1-H3 → 14 pt, H4 → 13.5 pt (см. `HEADING_FONT_SIZE_PT`). Размер выставляется
    на каждом run'е заголовка, чтобы перекрыть «застывшие» `<w:sz>` шаблона.
    Бежит по всем параграфам тела документа после полной сборки отчёта.
    """
    for paragraph in _iter_body_paragraphs(doc):
        level = _heading_level(paragraph)
        size_pt = _heading_size_pt_for_level(level)
        if size_pt is None:
            continue
        for run in paragraph.runs:
            _force_run_font_size_pt(run, size_pt)


def _force_h3_non_bold(doc: DocxDocument) -> None:
    """Снимает явный `<w:b/>` со всех H3-параграфов документа.

    Заголовки H3 в шаблоне имеют bold-runs, и Word при обновлении оглавления
    с дефолтным `TOC`-полем наследует это форматирование (т.е. подпункты вида
    «4.4.1 …» отображаются жирным). Сбрасываем bold на уровне run'ов: это
    влияет и на сам H3-текст в теле, и на соответствующую запись в оглавлении
    после автообновления Word'ом.
    """
    for paragraph in _iter_body_paragraphs(doc):
        if _heading_level(paragraph) != 3:
            continue
        for run in paragraph.runs:
            rPr = run._element.find(qn("w:rPr"))
            if rPr is None:
                continue
            for b in rPr.findall(qn("w:b")):
                b.set(qn("w:val"), "0")
            for b in rPr.findall(qn("w:bCs")):
                b.set(qn("w:val"), "0")


def _update_test_stand_paragraphs(doc: DocxDocument, hosts: list[Host]) -> None:
    """Заменяет «Исследование стенда http://XXX» актуальным списком стендов.

    H3-заголовки «Исследование стенда …» в разделах 4.4/4.5/4.6 пропускаем —
    их клонирует под каждый стенд `_replicate_h3_with_cards` (по одному
    H3-разделу на хост). Здесь обрабатываем только обычные параграфы
    (например, лид-абзац списка тестов в 4.1).
    """
    text_value = _hosts_text(hosts)
    if not text_value:
        return
    for paragraph in _iter_body_paragraphs(doc):
        norm = _normalize_label(paragraph.text)
        if not norm.startswith("исследование стенда"):
            continue
        if _heading_level(paragraph) is not None:
            continue
        if "http://XXX" in paragraph.text:
            _set_paragraph_text(paragraph, paragraph.text.replace("http://XXX", text_value))
        elif "XXX" in paragraph.text:
            _set_paragraph_text(paragraph, paragraph.text.replace("XXX", text_value))


def _update_findings_summary_text(
    doc: DocxDocument,
    *,
    vulnerability_count: int,
    weakness_count: int,
) -> None:
    """Подставляет количество уязвимостей в абзац-итог «В ходе исследования было выявлено …».

    PP: «выявлено XXX уязвимостей …» — XXX → vulnerability_count.
    СЗИ: «выявлено X уязвимостей и X недостатков …» — оба X → счётчики.
    """
    pp_pattern = re.compile(r"выявлено\s+XXX\s+уязвимостей", flags=re.IGNORECASE)
    szi_pattern = re.compile(
        r"(выявлено\s+)X(\s+уязвимостей\s+и\s+)X(\s+недостатков)",
        flags=re.IGNORECASE,
    )
    intro_pattern = re.compile(r"в\s+ходе\s+исследования\s+было\s+выявлено", flags=re.IGNORECASE)

    for paragraph in _iter_body_paragraphs(doc):
        text = paragraph.text or ""
        if not intro_pattern.search(text):
            continue

        if pp_pattern.search(text):
            new_text = pp_pattern.sub(f"выявлено {vulnerability_count} уязвимостей", text, count=1)
            _set_paragraph_text(paragraph, new_text)
            continue

        if szi_pattern.search(text):
            new_text = szi_pattern.sub(
                lambda m: f"{m.group(1)}{vulnerability_count}{m.group(2)}{weakness_count}{m.group(3)}",
                text,
                count=1,
            )
            _set_paragraph_text(paragraph, new_text)


_PP_HEADING_PATTERNS = (
    # «ПП «XXX»», «ПП «XXX»» — с кавычками-ёлочками.
    (re.compile(r"ПП\s*«\s*XXX\s*»"), "ПП «{name}»"),
    # «ПП XXX» — без кавычек.
    (re.compile(r"ПП\s+XXX(?![A-Za-zА-Яа-я0-9])"), "ПП {name}"),
    # «ПА «XXX»» — для титульного листа объекта оценки.
    (re.compile(r"ПА\s*«\s*XXX\s*»"), "ПА «{name}»"),
)


def _replace_pp_xxx_in_text(text: str, name: str) -> str:
    """Применяет шаблоны замены «ПП XXX» / «ПП «XXX»» / «ПА «XXX»» → имя проекта."""
    new_text = text
    for pattern, replacement in _PP_HEADING_PATTERNS:
        new_text = pattern.sub(replacement.format(name=name), new_text)
    return new_text


def _update_pp_headings(doc: DocxDocument, project: Project) -> None:
    """Заменяет XXX на название проекта в любых параграфах/ячейках вида «ПП XXX» / «ПП «XXX»»."""
    name = (project.name or "").strip()
    if not name:
        return

    def _maybe_update(paragraph: Paragraph) -> None:
        text = paragraph.text or ""
        new_text = _replace_pp_xxx_in_text(text, name)
        if new_text != text:
            _set_paragraph_text(paragraph, new_text)

    for paragraph in _iter_body_paragraphs(doc):
        _maybe_update(paragraph)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _maybe_update(paragraph)


def _direct_text_elements(paragraph_element) -> list:
    """Возвращает `<w:t>`, которые принадлежат именно этому параграфу.

    `Element.iter(...)` в lxml идёт вглубь всего поддерева, поэтому для внешнего
    `<w:p>`, внутри которого лежат `<w:txbxContent>` с собственными `<w:p>`,
    обычный обход подтянет тексты из вложенных параграфов. Здесь фильтруем только
    те `<w:t>`, у которых ближайший предок-`<w:p>` совпадает с текущим.
    """
    result = []
    p_tag = qn("w:p")
    for t in paragraph_element.iter(qn("w:t")):
        ancestor = t.getparent()
        while ancestor is not None and ancestor.tag != p_tag:
            ancestor = ancestor.getparent()
        if ancestor is paragraph_element:
            result.append(t)
    return result


def _replace_xxx_in_xml_subtree(element, name: str) -> None:
    """Заменяет XXX → name в `<w:t>` подходящих параграфов (футеры/хедеры).

    «Анализ защищенности «XXX»» часто разнесено по нескольким `<w:t>` элементам
    (например `«` и `XXX` и `»` лежат в разных runs). Поэтому работаем на уровне
    параграфа: склеиваем тексты, применяем подстановки на полной строке и
    переписываем результат в первый `<w:t>`, очищая остальные.

    ВАЖНО: учитываем только `<w:t>`, принадлежащие конкретному `<w:p>` (не
    вложенным параграфам внутри `<w:txbxContent>`). Иначе тексты из соседних
    текстовых полей (например «КОНФИДЕНЦИАЛЬНО», номер страницы) сольются в
    первый `<w:t>` и будут обрезаны узким текстовым полем.

    Применяется только к параграфам с маркерами «защищенности»/«ПП»/«ПА», чтобы
    случайные XXX в коде/значениях не затронуть.
    """
    markers = ("защищенности", "ПП", "ПА")
    for paragraph_element in element.iter(qn("w:p")):
        text_elements = _direct_text_elements(paragraph_element)
        if not text_elements:
            continue
        full_text = "".join(t.text or "" for t in text_elements)
        if "XXX" not in full_text:
            continue
        if not any(marker in full_text for marker in markers):
            continue
        new_full = re.sub(r"«\s*XXX\s*»", f"«{name}»", full_text)
        new_full = _replace_pp_xxx_in_text(new_full, name)
        if new_full == full_text:
            continue
        # Записываем результат в первый `<w:t>`, очищая остальные.
        text_elements[0].text = new_full
        for tail in text_elements[1:]:
            tail.text = ""


_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"


def _flatten_alternate_content(element) -> None:
    """Разворачивает все `<mc:AlternateContent>` в поддереве, оставляя только Choice.

    В шаблоне ПП футер построен через `<mc:AlternateContent>` с двумя ветками:
    `<mc:Choice Requires="wps">` (современный `<w:drawing>`) и `<mc:Fallback>`
    (легаси VML `<w:pict>`). MS Word/совместимые рендеры обязаны выбирать ровно
    одну ветку, но некоторые просмотрщики (LibreOffice, превьюшки) выводят обе —
    и пользователь видит дубль текста «Анализ защищенности «…»Анализ защищенности
    «…»» и «КОНФИДЕНЦИАЛЬНОКОНФИДЕНЦИАЛЬНО».

    Чтобы исключить дубль во всех рендерерах, схлопываем `AlternateContent`:
    оставляем содержимое первого подходящего `<mc:Choice>` (или `<mc:Fallback>`,
    если Choice пуст), а сам контейнер удаляем — Word и LibreOffice просто
    увидят обычный `<w:drawing>`/`<w:pict>` без альтернатив.
    """
    ac_tag = f"{{{_MC_NS}}}AlternateContent"
    choice_tag = f"{{{_MC_NS}}}Choice"
    fallback_tag = f"{{{_MC_NS}}}Fallback"
    # Сначала собираем список — модифицировать дерево во время iter() небезопасно.
    targets = [el for el in element.iter(ac_tag)]
    for ac in targets:
        keeper = None
        for child in ac:
            tag = child.tag
            if tag == choice_tag and len(child) > 0:
                keeper = child
                break
        if keeper is None:
            for child in ac:
                if child.tag == fallback_tag and len(child) > 0:
                    keeper = child
                    break
        parent = ac.getparent()
        if parent is None:
            continue
        idx = list(parent).index(ac)
        if keeper is not None:
            for inner in list(keeper):
                parent.insert(idx, inner)
                idx += 1
        parent.remove(ac)


def _update_footers_and_headers(doc: DocxDocument, project: Project) -> None:
    """Подставляет название проекта в футеры/хедеры всех секций.

    Текст футера лежит внутри `<w:txbxContent>` (текстовые поля), поэтому
    обходим XML-дерево соответствующих parts напрямую. Дополнительно схлопываем
    `<mc:AlternateContent>` (см. `_flatten_alternate_content`) — без этого в
    некоторых просмотрщиках текст футера дублируется (Choice + Fallback ветки
    отрисовываются обе).
    """
    name = (project.name or "").strip()
    for section in doc.sections:
        for container in (section.footer, section.header, section.first_page_footer,
                          section.first_page_header, section.even_page_footer,
                          section.even_page_header):
            if container is None:
                continue
            container_element = getattr(container, "_element", None)
            if container_element is None:
                # Для headers/footers `python-docx` использует `.part.element`.
                part = getattr(container, "part", None)
                container_element = getattr(part, "element", None) if part is not None else None
            if container_element is None:
                continue
            # Сначала схлопываем альтернативы — это уменьшит число `<w:t>` и
            # упростит дальнейшую подстановку имени.
            _flatten_alternate_content(container_element)
            if name:
                _replace_xxx_in_xml_subtree(container_element, name)


def _update_object_under_test(doc: DocxDocument, project: Project, hosts: list[Host]) -> None:
    """Подставляет имя продукта и список стендов в раздел «Объект оценки» (СЗИ).

    В шаблоне строка выглядит как «Стенды\xa0тестирования:» с неразрывным пробелом,
    а плейсхолдер «XXX» может быть как в СЛЕДУЮЩЕМ параграфе, так и прямо в той же
    строке после двоеточия/таба. Поэтому нормализуем пробелы перед сравнением и
    обрабатываем оба варианта.

    Стенды выводим как маркированный список (по одному хосту в строке): берём
    параграф-плейсхолдер «XXX» (он уже под bullet-стилем `af5`) и клонируем его
    под каждый дополнительный хост.
    """
    name = (project.name or "").strip()
    host_labels: list[str] = []
    for host in hosts:
        label = host.hostname or host.ip_address
        if label:
            host_labels.append(label)
    paragraphs = _iter_body_paragraphs(doc)

    def _normalize_ws(value: str) -> str:
        return re.sub(r"\s+", " ", (value or "").replace("\xa0", " ")).strip()

    def _append_bullets(template_para: Paragraph, labels: list[str]) -> None:
        """Заменяет содержимое template_para на labels[0] и клонирует его под labels[1:]."""
        _set_paragraph_text(template_para, labels[0])
        anchor = template_para._element
        for label in labels[1:]:
            new_p = copy.deepcopy(template_para._element)
            for r in list(new_p.findall(qn("w:r"))):
                new_p.remove(r)
            anchor.addnext(new_p)
            wrapper = Paragraph(new_p, template_para._parent)
            run = wrapper.add_run(label)
            _apply_run_font(run)
            anchor = new_p

    for index, paragraph in enumerate(paragraphs):
        raw_text = paragraph.text or ""
        normalized = _normalize_ws(raw_text)

        if name and (normalized.startswith("Продукт:") or normalized.startswith("Продукт :")):
            _set_paragraph_text(paragraph, f"Продукт: ПА «{name}»")
            continue

        if host_labels and normalized.lower().startswith("стенды тестирования"):
            # Вариант 1: всё в одной строке «Стенды тестирования: XXX»: меняем XXX
            # на первый хост, остальные — отдельными параграфами-bullet'ами после.
            if "XXX" in raw_text:
                new_text = re.sub(r"XXX\s*$", host_labels[0], raw_text)
                if new_text == raw_text:
                    new_text = raw_text.replace("XXX", host_labels[0], 1)
                _set_paragraph_text(paragraph, new_text)
                if len(host_labels) > 1:
                    # Дополнительные хосты вставляем отдельными параграфами после
                    # «Стенды тестирования: <host1>», копируя стиль текущего параграфа.
                    anchor = paragraph._element
                    for label in host_labels[1:]:
                        new_p = copy.deepcopy(paragraph._element)
                        for r in list(new_p.findall(qn("w:r"))):
                            new_p.remove(r)
                        anchor.addnext(new_p)
                        wrapper = Paragraph(new_p, paragraph._parent)
                        run = wrapper.add_run(label)
                        _apply_run_font(run)
                        anchor = new_p
                continue
            # Вариант 2: следующий параграф «XXX» (или пустой) — это и есть
            # bullet-плейсхолдер. Заменяем содержимое на первый хост, под него
            # клонируем доп. параграфы под каждый дополнительный хост.
            if index + 1 < len(paragraphs):
                next_para = paragraphs[index + 1]
                next_text = (next_para.text or "").strip()
                if "XXX" in next_text or not next_text:
                    _append_bullets(next_para, host_labels)


def _find_research_h3(start_para: Paragraph) -> Paragraph | None:
    """Ищет первый H3 «Исследование стенда …» после start_para до следующего H2."""
    body = start_para._element.getparent()
    children = list(body.iterchildren())
    start_idx = children.index(start_para._element)
    for child in children[start_idx + 1:]:
        if child.tag != qn("w:p"):
            continue
        wrapper = Paragraph(child, start_para._parent)
        level = _heading_level(wrapper)
        if level == 2:
            return None
        if level == 3 and _normalize_label(wrapper.text).startswith("исследование стенда"):
            return wrapper
    return None


def _block_until_heading_lt(start_para: Paragraph, max_level_exclusive: int) -> list:
    """Собирает элементы от start_para до первого заголовка с уровнем < max_level_exclusive.

    Например, при `max_level_exclusive=4` останавливаемся на H1/H2/H3.
    """
    parent = start_para._parent
    body = start_para._element.getparent()
    children = list(body.iterchildren())
    start_idx = children.index(start_para._element)
    elements = [start_para._element]
    for child in children[start_idx + 1:]:
        if child.tag == qn("w:p"):
            tmp = Paragraph(child, parent)
            level = _heading_level(tmp)
            if level is not None and level < max_level_exclusive:
                break
        elements.append(child)
    return elements


def _group_items_by_host(
    items: list[Vulnerability],
    *,
    indexes: dict,
    hosts: list[Host],
) -> list[tuple[str, list[Vulnerability]]]:
    """Группирует уязвимости по хосту в порядке списка `hosts` проекта.

    Если у уязвимости нет привязки к хосту, относим её в группу первого хоста
    проекта (fallback). Хосты, у которых нет ни одной уязвимости, в результат
    не попадают.
    """
    fallback = ""
    if hosts:
        fallback = hosts[0].hostname or hosts[0].ip_address or ""
    if not fallback:
        fallback = "Стенд"

    host_order: list[str] = []
    seen: set[str] = set()
    for host in hosts:
        label = host.hostname or host.ip_address or ""
        if label and label not in seen:
            seen.add(label)
            host_order.append(label)

    groups: dict[str, list[Vulnerability]] = {}
    for vuln in items:
        host = _vuln_host_address(vuln, indexes) or fallback
        groups.setdefault(host, []).append(vuln)
        if host not in seen:
            seen.add(host)
            host_order.append(host)

    return [(host, groups[host]) for host in host_order if host in groups]


def _replicate_h3_with_cards(
    doc: DocxDocument,
    *,
    h2: Paragraph,
    items: list[Vulnerability],
    indexes: dict,
    hosts: list[Host],
) -> list[tuple[Vulnerability, list]] | None:
    """Клонирует H3 «Исследование стенда» и блок-карточку под каждый стенд.

    Возвращает список (vuln, cloned_card_block_elements) в порядке следования
    в документе. Если шаблон H3/карточки не найден — `None` (вызывающий должен
    использовать legacy-fallback). Оригинальные шаблонные элементы (H3 + одна
    карточка) удаляются.
    """
    h3 = _find_research_h3(h2)
    if h3 is None:
        return None
    h4 = _find_first_h4_after(h3, max_stop_level=2)
    if h4 is None:
        return None
    template_card_block = _block_until_heading_lt(h4, max_level_exclusive=4)
    template_h3_element = h3._element

    if not items:
        for el in [template_h3_element] + template_card_block:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)
        return []

    groups = _group_items_by_host(items, indexes=indexes, hosts=hosts)
    last = template_card_block[-1]
    result: list[tuple[Vulnerability, list]] = []
    for host, host_items in groups:
        new_h3_el = copy.deepcopy(template_h3_element)
        last.addnext(new_h3_el)
        last = new_h3_el
        new_h3_para = Paragraph(new_h3_el, doc)
        _set_heading_text(new_h3_para, f"Исследование стенда {host}")

        for vuln in host_items:
            cloned: list = []
            for el in template_card_block:
                new_el = copy.deepcopy(el)
                last.addnext(new_el)
                last = new_el
                cloned.append(new_el)
            result.append((vuln, cloned))

    for el in [template_h3_element] + template_card_block:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
    return result


def _build_section(
    doc: DocxDocument,
    *,
    section_h2_text: str,
    items: list[Vulnerability],
    kind: ReportKind,
    indexes: dict,
    project: Project,
    hosts: list[Host],
    image_bytes_by_id: dict[UUID, bytes],
    image_state: dict,
) -> None:
    """Клонирует пример карточки в указанном H2-разделе под каждый элемент `items`.

    Если в разделе есть H3 «Исследование стенда …», клонируем его под каждый
    стенд (хост) и под каждым размещаем карточки уязвимостей, привязанных к
    этому стенду. Если элементов нет, пример карточки удаляется.
    """
    h2 = _find_h2_paragraph(doc, section_h2_text)
    if h2 is not None:
        per_host = _replicate_h3_with_cards(
            doc, h2=h2, items=items, indexes=indexes, hosts=hosts,
        )
        if per_host is not None:
            for index, (vuln, block) in enumerate(per_host, start=1):
                _fill_card(
                    block,
                    kind=kind,
                    doc=doc,
                    vuln=vuln,
                    sequence=index,
                    indexes=indexes,
                    project=project,
                    image_bytes_by_id=image_bytes_by_id,
                    image_state=image_state,
                )
            return

    anchor = _find_card_anchor(doc, section_h2_text)
    if anchor is None:
        return
    template_block = _card_block(anchor)
    if not items:
        _remove_block(template_block)
        return

    blocks = _replicate_card(template_block, len(items))
    for index, (block, vuln) in enumerate(zip(blocks, items), start=1):
        _fill_card(
            block,
            kind=kind,
            doc=doc,
            vuln=vuln,
            sequence=index,
            indexes=indexes,
            project=project,
            image_bytes_by_id=image_bytes_by_id,
            image_state=image_state,
        )


def _find_h2_paragraph(doc: DocxDocument, section_h2_text: str) -> Paragraph | None:
    """Возвращает H2-параграф с указанным текстом или None."""
    norm_target = _normalize_label(section_h2_text)
    for paragraph in _iter_body_paragraphs(doc):
        if _heading_level(paragraph) == 2 and _normalize_label(paragraph.text) == norm_target:
            return paragraph
    return None


def _is_test_list_item(paragraph: Paragraph) -> bool:
    """Параграф вида «Test–N «Наименование теста»» в разделе 4.1."""
    text = paragraph.text.strip()
    return bool(re.match(r"^test[\s\u00a0]*[–\-]\s*\d+", text, flags=re.IGNORECASE))


def _build_test_procedures_list(doc: DocxDocument, items: list[Vulnerability]) -> None:
    """Раздел 4.1 «Перечень тестовых процедур» — список тестов = названия items."""
    section = _find_h2_paragraph(doc, "Перечень тестовых процедур")
    if section is None:
        return
    body = section._element.getparent()
    children = list(body.iterchildren())
    start_idx = children.index(section._element)

    test_paragraphs: list[Paragraph] = []
    ellipsis_paragraph: Paragraph | None = None

    for child in children[start_idx + 1:]:
        if child.tag != qn("w:p"):
            continue
        wrapper = Paragraph(child, doc)
        level = _heading_level(wrapper)
        if level is not None and level <= 2:
            break
        text_stripped = wrapper.text.strip()
        if _is_test_list_item(wrapper):
            test_paragraphs.append(wrapper)
        elif test_paragraphs and not ellipsis_paragraph and text_stripped in {"…", "...", "..."}:
            ellipsis_paragraph = wrapper

    if not test_paragraphs:
        return

    template_para = test_paragraphs[0]
    template_element = template_para._element

    if not items:
        for para in test_paragraphs:
            para._element.getparent().remove(para._element)
        if ellipsis_paragraph is not None:
            ellipsis_paragraph._element.getparent().remove(ellipsis_paragraph._element)
        return

    # Готовим N клонов на основе шаблона (включая первый — переиспользуем).
    rendered: list[Paragraph] = [template_para]
    last_element = template_element
    for _ in range(len(items) - 1):
        new_element = copy.deepcopy(template_element)
        last_element.addnext(new_element)
        last_element = new_element
        rendered.append(Paragraph(new_element, doc))

    for index, (para, vuln) in enumerate(zip(rendered, items), start=1):
        title = (vuln.title or "").strip() or "(без названия)"
        _set_paragraph_text(para, f"Test-{index} «{title}»")

    # Удаляем «лишние» исходные шаблонные параграфы и многоточие.
    for tail in test_paragraphs[1:]:
        if tail._element not in [p._element for p in rendered]:
            parent = tail._element.getparent()
            if parent is not None:
                parent.remove(tail._element)
    if ellipsis_paragraph is not None:
        parent = ellipsis_paragraph._element.getparent()
        if parent is not None:
            parent.remove(ellipsis_paragraph._element)


# ----------------------------- Раздел 4.4 «Выполнение тестов» -----------------------------

# Подписи разделов внутри карточки теста (проза после H4).
TEST_DESC_HEADER = "Описание теста:"
TEST_STEPS_HEADER = "Шаги при проведении теста:"
TEST_EXPECTED_HEADER = "Ожидаемый результат:"
TEST_ACTUAL_HEADER = "Фактический результат:"
TEST_CONCLUSION_HEADER = "Вывод:"
TEST_HEADERS_ALL = (
    TEST_DESC_HEADER,
    TEST_STEPS_HEADER,
    TEST_EXPECTED_HEADER,
    TEST_ACTUAL_HEADER,
    TEST_CONCLUSION_HEADER,
)


def _identify_test_header(text: str) -> str | None:
    norm = _normalize_label(text)
    for header in TEST_HEADERS_ALL:
        if norm == _normalize_label(header):
            return header
    return None


def _find_first_h4_after(start_para: Paragraph, max_stop_level: int = 2) -> Paragraph | None:
    """Возвращает первый H4-параграф после start_para до заголовка ≤ max_stop_level."""
    body = start_para._element.getparent()
    children = list(body.iterchildren())
    start_idx = children.index(start_para._element)
    for child in children[start_idx + 1:]:
        if child.tag != qn("w:p"):
            continue
        wrapper = Paragraph(child, start_para._parent)
        level = _heading_level(wrapper)
        if level is not None and level <= max_stop_level:
            return None
        if level == 4:
            return wrapper
    return None


def _block_until_heading_lte(start_para: Paragraph, max_stop_level: int) -> list:
    """Собирает XML-элементы от стартового параграфа до заголовка ≤ max_stop_level (не включая)."""
    parent = start_para._parent
    body = start_para._element.getparent()
    children = list(body.iterchildren())
    start_idx = children.index(start_para._element)
    elements = [start_para._element]
    for child in children[start_idx + 1:]:
        if child.tag == qn("w:p"):
            tmp = Paragraph(child, parent)
            level = _heading_level(tmp)
            if level is not None and level <= max_stop_level:
                break
        elements.append(child)
    return elements


def _filter_paragraph_elements(elements: list, doc: DocxDocument) -> list[Paragraph]:
    return [Paragraph(el, doc) for el in elements if el.tag == qn("w:p")]


def _fill_test_card(
    block_elements: list,
    *,
    doc: DocxDocument,
    vuln: Vulnerability,
    sequence: int,
    image_bytes_by_id: dict[UUID, bytes],
    files_by_id: dict[UUID, File],
    files_by_vuln_id: dict[UUID, list[File]],
    image_state: dict,
) -> None:
    """Заполняет один блок-карточку теста в разделе 4.4.

    H4 → «Тест на проникновение Test-N «<title>»».
    «Описание теста:» → vuln.impact (описание воздействия/уязвимости).
    «Шаги при проведении теста:» → этапы без картинок.
    «Фактический результат:» → этапы с картинками.
    «Вывод:» → краткий шаблонный итог.
    """
    paragraphs = _filter_paragraph_elements(block_elements, doc)
    if not paragraphs:
        return

    title = (vuln.title or "").strip() or "(без названия)"
    _set_heading_text(paragraphs[0], f"Тест на проникновение Test-{sequence} «{title}»")

    # Группируем параграфы (после H4) по подписям-разделам.
    sections: dict[str, list[Paragraph]] = {}
    section_headers: dict[str, Paragraph] = {}
    current: str | None = None
    for p in paragraphs[1:]:
        header = _identify_test_header(p.text)
        if header is not None:
            current = header
            sections.setdefault(current, [])
            section_headers[current] = p
            continue
        if current is None:
            continue
        sections[current].append(p)

    impact_text = (vuln.impact or "").strip()
    description_text = (vuln.description or "").strip()
    workflow_steps = list(vuln.workflow_steps or [])

    content_paragraphs: list[Paragraph] = []

    content_paragraphs += _replace_section_with_lines(
        paragraphs=sections.get(TEST_DESC_HEADER, []),
        lines=_split_to_lines(impact_text or description_text),
    )
    content_paragraphs += _replace_section_with_lines(
        paragraphs=sections.get(TEST_STEPS_HEADER, []),
        lines=[_step_one_line(idx, step) for idx, step in enumerate(workflow_steps, start=1)] if workflow_steps else [],
        strip_numbering=True,
    )
    # «Ожидаемый результат:» в шаблоне содержит примеры с жёстко прописанным
    # «Test–1». Поскольку реальных данных у нас нет, очищаем содержимое раздела
    # одним «…», чтобы исключить путаницу с шаблонным номером.
    content_paragraphs += _replace_section_with_lines(
        paragraphs=sections.get(TEST_EXPECTED_HEADER, []),
        lines=[],
    )
    content_paragraphs += _fill_actual_result_section(
        paragraphs=sections.get(TEST_ACTUAL_HEADER, []),
        steps=workflow_steps,
        vuln=vuln,
        image_bytes_by_id=image_bytes_by_id,
        files_by_id=files_by_id,
        files_by_vuln_id=files_by_vuln_id,
        image_state=image_state,
        doc=doc,
    )
    # «Вывод:» — пользователь попросил, чтобы итоговый текст шёл сразу после
    # подписи, в одной строке (а не как отдельный абзац ниже). Поэтому
    # дописываем его run'ом прямо в параграф-подпись и удаляем все
    # автогенерированные контентные параграфы под ним.
    conclusion_text = (
        f" Сценарий атаки Test-{sequence} «{title}» может быть использован для реализации "
        "угроз безопасности информации."
    )
    conclusion_header = section_headers.get(TEST_CONCLUSION_HEADER)
    if conclusion_header is not None:
        conclusion_run = conclusion_header.add_run(conclusion_text)
        _apply_run_font(conclusion_run, bold=False)
    for tail in sections.get(TEST_CONCLUSION_HEADER, []):
        parent = tail._element.getparent()
        if parent is not None:
            parent.remove(tail._element)

    # Контентным абзацам разделов теста удваиваем левый отступ — пользователь
    # попросил, чтобы текст *под* жирной подписью имел «отступ от начала
    # строки два раза» (сами подписи остаются на своих местах).
    for content_p in content_paragraphs:
        _double_paragraph_left_indent(content_p)


def _split_to_lines(value: str) -> list[str]:
    """Разбивает текст на непустые строки (без обрезки внутри строк)."""
    return [line for line in (value or "").splitlines() if line.strip()] or ([value] if value.strip() else [])


def _replace_section_with_lines(
    *,
    paragraphs: list[Paragraph],
    lines: list[str],
    strip_numbering: bool = False,
) -> list[Paragraph]:
    """Заменяет содержимое раздела (после подписи) на список строк.

    Если данных нет — ставим единственное многоточие «…».
    Лишние шаблонные параграфы удаляем. Опция `strip_numbering=True`
    отключает Word'овскую авто-нумерацию (для шагов/этапов теста).

    Возвращает список фактически использованных контентных параграфов —
    чтобы вызывающий код мог, например, удвоить им левый отступ.
    """
    if not paragraphs:
        return []
    first = paragraphs[0]
    if not lines:
        _set_paragraph_text(first, "…")
        for tail in paragraphs[1:]:
            tail._element.getparent().remove(tail._element)
        return [first]
    _set_paragraph_text(first, lines[0])
    if strip_numbering:
        _strip_paragraph_numbering(first._element)
    used: list[Paragraph] = [first]
    anchor = first._element
    for line in lines[1:]:
        new_p = copy.deepcopy(first._element)
        for r in list(new_p.findall(qn("w:r"))):
            new_p.remove(r)
        anchor.addnext(new_p)
        wrapper = Paragraph(new_p, first._parent)
        run = wrapper.add_run(line)
        _apply_run_font(run)
        if strip_numbering:
            _strip_paragraph_numbering(new_p)
        used.append(wrapper)
        anchor = new_p
    for tail in paragraphs[1:]:
        parent = tail._element.getparent()
        if parent is not None:
            parent.remove(tail._element)
    return used


def _fill_actual_result_section(
    *,
    paragraphs: list[Paragraph],
    steps: list[dict],
    vuln: Vulnerability,
    image_bytes_by_id: dict[UUID, bytes],
    files_by_id: dict[UUID, File],
    files_by_vuln_id: dict[UUID, list[File]],
    image_state: dict,
    doc: DocxDocument,
) -> list[Paragraph]:
    """«Фактический результат:» — для каждого этапа сначала строка-описание, затем его картинки.

    В отчёт идут только картинки, явно привязанные к этапам уязвимости через
    `image_file_ids` или markdown в описании этапа. «Свободные» файлы-доказательства
    в отчёт не попадают. Возвращает список контентных строк-описаний этапов
    (без параграфов с картинками/подписями), чтобы вызывающий код мог удвоить
    им левый отступ.
    """
    del files_by_vuln_id, vuln  # extras-картинки больше не вставляем
    if not paragraphs:
        return []
    has_any_step = bool(steps)
    if not has_any_step:
        return _replace_section_with_lines(paragraphs=paragraphs, lines=[])

    first = paragraphs[0]
    # Удаляем все шаблонные параграфы, кроме первого: его переиспользуем как первую запись.
    for tail in paragraphs[1:]:
        parent = tail._element.getparent()
        if parent is not None:
            parent.remove(tail._element)

    anchor_element = first._element
    first_used = False
    text_paragraphs: list[Paragraph] = []
    for index, step in enumerate(steps, start=1):
        line = _step_one_line(index, step)
        if not first_used:
            _set_paragraph_text(first, line)
            _strip_paragraph_numbering(first._element)
            first_used = True
            text_paragraphs.append(first)
        else:
            new_p = copy.deepcopy(first._element)
            for r in list(new_p.findall(qn("w:r"))):
                new_p.remove(r)
            anchor_element.addnext(new_p)
            wrapper = Paragraph(new_p, first._parent)
            run = wrapper.add_run(line)
            _apply_run_font(run)
            _strip_paragraph_numbering(new_p)
            text_paragraphs.append(wrapper)
            anchor_element = new_p

        for file_id in _step_image_file_ids(step):
            data = image_bytes_by_id.get(file_id)
            if not data:
                continue
            file_meta = files_by_id.get(file_id)
            caption = _figure_caption_text(
                step=step,
                file_name=file_meta.original_name if file_meta is not None else None,
                fallback_index=index,
            )
            anchor_element = _insert_image_with_caption(
                anchor_element=anchor_element,
                doc=doc,
                image_bytes=data,
                caption_text=caption,
                image_state=image_state,
            )
    return text_paragraphs


def _build_test_executions(
    doc: DocxDocument,
    items: list[Vulnerability],
    image_bytes_by_id: dict[UUID, bytes],
    *,
    files_by_id: dict[UUID, File],
    files_by_vuln_id: dict[UUID, list[File]],
    indexes: dict,
    hosts: list[Host],
    image_state: dict,
) -> None:
    """Раздел 4.4 «Выполнение тестов» — клонирует карточку H4 на каждый item.

    Если в шаблоне присутствует H3 «Исследование стенда …», клонируем его под
    каждый стенд и под каждым размещаем карточки тестов соответствующих
    уязвимостей.
    """
    section = _find_h2_paragraph(doc, "Выполнение тестов")
    if section is None:
        return

    per_host = _replicate_h3_with_cards(
        doc, h2=section, items=items, indexes=indexes, hosts=hosts,
    )
    if per_host is not None:
        for index, (vuln, block) in enumerate(per_host, start=1):
            _fill_test_card(
                block,
                doc=doc,
                vuln=vuln,
                sequence=index,
                image_bytes_by_id=image_bytes_by_id,
                files_by_id=files_by_id,
                files_by_vuln_id=files_by_vuln_id,
                image_state=image_state,
            )
        return

    template_h4 = _find_first_h4_after(section, max_stop_level=2)
    if template_h4 is None:
        return
    template_block = _block_until_heading_lte(template_h4, max_stop_level=3)

    if not items:
        for element in template_block:
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
        return

    blocks = _replicate_card(template_block, len(items))
    for index, (block, vuln) in enumerate(zip(blocks, items), start=1):
        _fill_test_card(
            block,
            doc=doc,
            vuln=vuln,
            sequence=index,
            image_bytes_by_id=image_bytes_by_id,
            files_by_id=files_by_id,
            files_by_vuln_id=files_by_vuln_id,
            image_state=image_state,
        )


def _build_common(
    kind: ReportKind,
    data: dict,
    indexes: dict,
    image_bytes_by_id: dict[UUID, bytes],
) -> bytes:
    project: Project = data["project"]
    hosts: list[Host] = data["hosts"]
    vulnerabilities: list[Vulnerability] = data["vulnerabilities"]
    doc = load_template(kind)

    files_by_id: dict[UUID, File] = indexes.get("files_by_id", {}) or {}
    files_by_vuln_id: dict[UUID, list[File]] = indexes.get("files_by_vuln_id", {}) or {}
    image_state: dict = {_FIGURE_COUNTER_KEY: 0}

    _update_cover_page(doc, project)
    _update_cover_date(doc, project)
    _update_project_intro(doc, project)
    if kind == "szi":
        _update_object_under_test(doc, project, hosts)
    _update_test_stand_paragraphs(doc, hosts)
    _update_pp_headings(doc, project)
    _update_footers_and_headers(doc, project)

    severity_counts: dict[Severity, int] = {sev: 0 for sev in Severity}
    for vuln in vulnerabilities:
        severity_counts[vuln.severity] = severity_counts.get(vuln.severity, 0) + 1
    _update_severity_summary_table(doc, severity_counts)

    main_vulns, weakness_vulns = _split_vulnerabilities(vulnerabilities)
    all_items = main_vulns + weakness_vulns
    _update_findings_summary_text(doc, vulnerability_count=len(main_vulns), weakness_count=len(weakness_vulns))

    if kind == "szi":
        _build_test_procedures_list(doc, all_items)
        _build_test_executions(
            doc,
            all_items,
            image_bytes_by_id,
            files_by_id=files_by_id,
            files_by_vuln_id=files_by_vuln_id,
            indexes=indexes,
            hosts=hosts,
            image_state=image_state,
        )

    _build_section(
        doc,
        section_h2_text="Информация по выявленным уязвимостям",
        items=main_vulns,
        kind=kind,
        indexes=indexes,
        project=project,
        hosts=hosts,
        image_bytes_by_id=image_bytes_by_id,
        image_state=image_state,
    )
    _build_section(
        doc,
        section_h2_text="Информация по выявленным слабостям",
        items=weakness_vulns,
        kind=kind,
        indexes=indexes,
        project=project,
        hosts=hosts,
        image_bytes_by_id=image_bytes_by_id,
        image_state=image_state,
    )

    # Финальный проход: сбрасываем жирный со всех H3 — иначе Word при обновлении
    # оглавления отображает подпункты типа «4.4.1 …» жирным.
    _force_h3_non_bold(doc)

    # Жёстко фиксируем размеры заголовков (H1-H3 = 14 pt, H4 = 13.5 pt) на
    # уровне run'ов, чтобы перекрыть «застывшие» `<w:sz>` из шаблона.
    _normalize_heading_sizes(doc)

    # Просим Word/совместимые редакторы автоматически обновить все поля
    # (включая TOC) при открытии документа. Без этого пользователь видит
    # «застывшее» оглавление шаблона и должен вручную нажимать «Update field».
    _enable_auto_field_update(doc)
    _rebuild_toc(doc)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_szi(data: dict, indexes: dict, image_bytes_by_id: dict[UUID, bytes]) -> bytes:
    """Собирает СЗИ-отчёт (сертификация)."""
    return _build_common("szi", data, indexes, image_bytes_by_id)


def build_pp(data: dict, indexes: dict, image_bytes_by_id: dict[UUID, bytes]) -> bytes:
    """Собирает ПП-отчёт (внутренняя приёмка)."""
    return _build_common("pp", data, indexes, image_bytes_by_id)
