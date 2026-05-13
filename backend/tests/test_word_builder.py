"""Смоук-тесты сборщика Word-отчётов СЗИ/ПП."""
from __future__ import annotations

from datetime import date
from io import BytesIO
from uuid import uuid4

import pytest
from docx import Document

from app.enums import AssetType, CvssVersion, ProjectStatus, Severity, VulnerabilityStatus
from app.models import Host, Project, Vulnerability, VulnerabilityAsset
from app.reports import build_pp, build_szi


def _make_project(name: str = "DemoApp") -> Project:
    project = Project()
    project.id = uuid4()
    project.name = name
    project.folder = ""
    project.description = "Описание"
    project.start_date = date(2026, 1, 10)
    project.end_date = date(2026, 1, 20)
    project.status = ProjectStatus.ACTIVE
    project.created_by = uuid4()
    return project


def _make_host(project: Project, *, hostname: str | None = None, ip: str | None = None) -> Host:
    host = Host()
    host.id = uuid4()
    host.project_id = project.id
    host.hostname = hostname
    host.ip_address = ip
    return host


def _make_vuln(
    project: Project,
    *,
    title: str,
    severity: Severity,
    description: str | None = None,
    impact: str | None = None,
    recommendations: str | None = None,
    cvss_vector: str | None = None,
    cvss_score: float | None = None,
    cwe_id: str | None = None,
) -> Vulnerability:
    vuln = Vulnerability()
    vuln.id = uuid4()
    vuln.project_id = project.id
    vuln.title = title
    vuln.severity = severity
    vuln.description = description
    vuln.impact = impact
    vuln.recommendations = recommendations
    vuln.status = VulnerabilityStatus.OPEN
    vuln.cvss_version = CvssVersion.V40 if cvss_vector else None
    vuln.cvss_score = cvss_score
    vuln.cvss_vector = cvss_vector
    vuln.cwe_id = cwe_id
    vuln.workflow_steps = None
    vuln.steps_to_reproduce = None
    vuln.created_by = uuid4()
    return vuln


def _make_link(vuln: Vulnerability, host: Host) -> VulnerabilityAsset:
    link = VulnerabilityAsset()
    link.id = uuid4()
    link.vulnerability_id = vuln.id
    link.asset_type = AssetType.HOST
    link.asset_id = host.id
    return link


def _build_payload() -> tuple[dict, dict]:
    project = _make_project()
    host_a = _make_host(project, hostname="api.demo.local", ip="10.0.0.1")
    host_b = _make_host(project, ip="10.0.0.2")
    vuln_high = _make_vuln(
        project,
        title="SQL injection in login form",
        severity=Severity.HIGH,
        description="Параметр username не экранируется.",
        impact="Возможна полная компрометация БД.",
        recommendations="Использовать параметризованные запросы.",
        cvss_vector="CVSS:4.0/AV:N/AC:L/AT:N/PR:N/UI:N/VC:H/VI:H/VA:H/SC:N/SI:N/SA:N",
        cvss_score=9.3,
        cwe_id="CWE-89",
    )
    vuln_info = _make_vuln(
        project,
        title="Verbose error messages",
        severity=Severity.INFO,
        description=None,
        impact=None,
        recommendations=None,
    )
    link = _make_link(vuln_high, host_a)
    data = {
        "project": project,
        "hosts": [host_a, host_b],
        "vulnerabilities": [vuln_high, vuln_info],
        "files": [],
        "members": [],
    }
    indexes = {
        "host_by_id": {host_a.id: host_a, host_b.id: host_b},
        "assets_by_vuln_id": {vuln_high.id: [link], vuln_info.id: []},
    }
    return data, indexes


def _doc_text(content: bytes) -> str:
    document = Document(BytesIO(content))
    parts: list[str] = []
    for paragraph in document.paragraphs:
        parts.append(paragraph.text)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


@pytest.mark.parametrize("builder", [build_szi, build_pp])
def test_word_builder_returns_non_empty_docx(builder) -> None:
    data, indexes = _build_payload()
    content = builder(data, indexes, image_bytes_by_id={})

    assert isinstance(content, bytes) and len(content) > 0
    assert content[:2] == b"PK"  # zip-сигнатура DOCX

    text = _doc_text(content)
    assert "DemoApp" in text
    assert "SQL injection in login form" in text
    assert "Verbose error messages" in text
    assert "10.0.0.1" in text or "api.demo.local" in text


def test_word_builder_splits_vulnerabilities_and_weaknesses() -> None:
    data, indexes = _build_payload()
    content = build_szi(data, indexes, image_bytes_by_id={})
    document = Document(BytesIO(content))

    section_heads: list[tuple[str, list[str]]] = []
    current: tuple[str, list[str]] | None = None
    for paragraph in document.paragraphs:
        style = (paragraph.style.name or "") if paragraph.style else ""
        if style in {"Heading 2", "Заголовок 2"}:
            current = (paragraph.text.strip(), [])
            section_heads.append(current)
        elif current is not None:
            current[1].append(paragraph.text)

    sections = {head: "\n".join(body) for head, body in section_heads}
    vuln_section = next((v for k, v in sections.items() if "выявленным уязвимостям" in k.lower()), "")
    weakness_section = next((v for k, v in sections.items() if "выявленным слабостям" in k.lower()), "")

    assert "SQL injection in login form" in vuln_section
    assert "Verbose error messages" in weakness_section
